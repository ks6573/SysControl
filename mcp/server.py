#!/usr/bin/env python3
"""
MCP Server: System Activity Monitor
Exposes tools for querying CPU, RAM, GPU, disk, network, and process info.
"""

import ast
import base64
import contextlib
import datetime
import heapq
import io
import ipaddress
import json
import os
import pathlib

try:
    import fcntl as _fcntl
    _HAS_FCNTL = True
except ImportError:
    _HAS_FCNTL = False
import platform
import plistlib
import re
import shutil
import socket
import subprocess
import sys
import threading
import time
import urllib.error
import urllib.parse
import urllib.request
import uuid
import webbrowser
from collections.abc import Callable
from concurrent.futures import ThreadPoolExecutor
from functools import lru_cache

import matplotlib
from openai import OpenAI

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker

try:
    import psutil
except ImportError:
    print("psutil not found. Install with: pip install psutil", file=sys.stderr)
    sys.exit(1)

# Optional GPU support via pynvml (nvidia-ml-py).
GPU_BACKEND = None
try:
    import pynvml
    pynvml.nvmlInit()
    GPU_BACKEND = "pynvml"
except Exception as exc:
    sys.stderr.write(f"[syscontrol] GPU backend init failed: {exc}\n")
    GPU_BACKEND = None

GPU_AVAILABLE = GPU_BACKEND is not None

# Optional spreadsheet support via openpyxl.
try:
    import openpyxl
    _HAS_OPENPYXL = True
except ImportError:
    _HAS_OPENPYXL = False

# Optional Word document support via python-docx.
try:
    from docx import Document as _DocxDocument
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

# Optional PDF support via pypdf.
try:
    from pypdf import PdfReader as _PdfReader
    _HAS_PYPDF = True
except ImportError:
    _HAS_PYPDF = False

# ── Platform constants (computed once at startup) ─────────────────────────────
_SYSTEM  = platform.system()
_MACHINE = platform.machine()
IS_MACOS = _SYSTEM == "Darwin"
IS_LINUX = _SYSTEM == "Linux"
IS_WIN   = _SYSTEM == "Windows"

DEFAULT_API_KEY = "ollama"
DEFAULT_BASE_URL = "http://localhost:11434/v1"
DEFAULT_LOCAL_MODEL = "qwen3:30b"

# ── Shared thread pool for parallel metric collection ─────────────────────────
# Reused across calls — avoids per-call thread creation/teardown overhead.
_METRICS_EXECUTOR = ThreadPoolExecutor(max_workers=8, thread_name_prefix="syscontrol-metrics")

# ── pynvml handle cache (handles are stable for the process lifetime) ─────────
_NVML_HANDLES: list = []
_NVML_HANDLES_READY = False  # sentinel — safe to read without lock
_NVML_LOCK = threading.Lock()


def _get_nvml_handles() -> list:
    """Return cached pynvml device handles; populated lazily on first call."""
    global _NVML_HANDLES, _NVML_HANDLES_READY
    if _NVML_HANDLES_READY:
        return _NVML_HANDLES
    with _NVML_LOCK:
        if _NVML_HANDLES_READY:
            return _NVML_HANDLES
        try:
            count = pynvml.nvmlDeviceGetCount()
            _NVML_HANDLES = [pynvml.nvmlDeviceGetHandleByIndex(i) for i in range(count)]
        except Exception:
            _NVML_HANDLES = []
        _NVML_HANDLES_READY = True
    return _NVML_HANDLES


# ── Reminder storage ──────────────────────────────────────────────────────────

_REMINDER_LOCK = threading.Lock()
_REMINDER_DIR  = pathlib.Path.home() / ".syscontrol"
_REMINDER_FILE = _REMINDER_DIR / "reminders.json"
# Create the config directory once at server startup, not on every read/write.
_REMINDER_DIR.mkdir(parents=True, exist_ok=True)

# ── Tool self-extension constants ─────────────────────────────────────────────
_FROZEN = getattr(sys, "frozen", False)

if _FROZEN:
    # Inside a PyInstaller bundle — data files live under sys._MEIPASS
    _BUNDLE_DIR   = pathlib.Path(sys._MEIPASS)  # type: ignore[attr-defined]
    _SERVER_FILE  = _BUNDLE_DIR / "mcp" / "server.py"
    _PROMPT_FILE  = _BUNDLE_DIR / "mcp" / "prompt.json"
else:
    _SERVER_FILE  = pathlib.Path(__file__)
    _PROMPT_FILE  = pathlib.Path(__file__).parent / "prompt.json"

# Persistent memory file — always in the writable user-data directory.
_USER_DATA_DIR = pathlib.Path.home() / ".syscontrol"
_USER_DATA_DIR.mkdir(parents=True, exist_ok=True)
_MEMORY_FILE = _USER_DATA_DIR / "SysControl_Memory.md"
_MEMORY_LOCK  = threading.Lock()
# Marker prepended to each user-defined function block in this file.
_USER_TOOL_FN_MARKER  = "# ── User-Defined Tool:"
# Anchor comment inside the TOOLS dict where new entries are inserted.
_USER_TOOL_REG_MARKER = "# ── User-Defined Tools (registry) ──────────────────────────────────────────"
_REMINDER_START_LOCK = threading.Lock()
_REMINDER_STARTED = False


def _load_reminders() -> list:
    """Load reminders from disk. Creates file if missing. Must be called under _REMINDER_LOCK."""
    if not _REMINDER_FILE.exists():
        return []
    try:
        return json.loads(_REMINDER_FILE.read_text())
    except (json.JSONDecodeError, OSError):
        return []


def _save_reminders(reminders: list) -> None:
    """Write reminders to disk. Must be called under _REMINDER_LOCK."""
    _REMINDER_FILE.write_text(json.dumps(reminders, indent=2))


class ReminderChecker:
    """Background daemon thread that fires due reminders via macOS notifications."""

    def __init__(self) -> None:
        self._stop = threading.Event()
        self._thread = threading.Thread(
            target=self._loop, daemon=True, name="syscontrol-reminders"
        )

    def start(self) -> None:
        self._thread.start()

    def _loop(self) -> None:
        while not self._stop.is_set():
            next_due = self._check()
            # Wait until the next reminder is due, capped at 15 s so new
            # reminders set by other tools are noticed quickly.
            self._stop.wait(timeout=min(15.0, max(1.0, next_due)))

    def _check(self) -> float:
        """Check and fire due reminders. Returns seconds until the next unfired reminder."""
        now = datetime.datetime.now()
        cutoff = now - datetime.timedelta(days=7)
        to_fire = []
        next_due = float("inf")
        with _REMINDER_LOCK:
            reminders = _load_reminders()
            changed = False
            survivors = []
            for r in reminders:
                try:
                    fire_at = datetime.datetime.fromisoformat(r["fire_at"])
                except (ValueError, KeyError, TypeError):
                    changed = True  # drop malformed entry
                    continue
                if r.get("fired"):
                    # Prune fired reminders older than 7 days
                    if fire_at >= cutoff:
                        survivors.append(r)
                    else:
                        changed = True
                    continue
                if now >= fire_at:
                    to_fire.append(r["message"])
                    r["fired"] = True
                    changed = True
                else:
                    secs = (fire_at - now).total_seconds()
                    if secs < next_due:
                        next_due = secs
                survivors.append(r)
            if changed:
                _save_reminders(survivors)
        # Fire notifications outside the lock to avoid blocking set/list/cancel
        for msg in to_fire:
            self._fire(msg)
        return next_due

    @staticmethod
    def _fire(message: str) -> None:
        script = (
            f'display notification {json.dumps(message)} '
            f'with title "SysControl Reminder" sound name "default"'
        )
        log_path = pathlib.Path.home() / ".syscontrol" / "reminder_log.txt"
        try:
            proc = subprocess.run(
                ["osascript", "-e", script],
                capture_output=True, text=True, timeout=5,
            )
            if proc.returncode != 0:
                log_path.parent.mkdir(parents=True, exist_ok=True)
                with log_path.open("a") as f:
                    ts = datetime.datetime.now().isoformat(timespec="seconds")
                    f.write(f"[{ts}] osascript failed (rc={proc.returncode}): {proc.stderr.strip()}\n")
        except Exception as exc:
            try:
                log_path.parent.mkdir(parents=True, exist_ok=True)
                with log_path.open("a") as f:
                    ts = datetime.datetime.now().isoformat(timespec="seconds")
                    f.write(f"[{ts}] _fire exception: {exc}\n")
            except Exception:
                pass


def _start_reminder_checker_once() -> None:
    """Start the reminder checker for this process if it is not already running."""
    global _REMINDER_STARTED
    with _REMINDER_START_LOCK:
        if _REMINDER_STARTED:
            return
        ReminderChecker().start()
        _REMINDER_STARTED = True


# ── MCP helpers ──────────────────────────────────────────────────────────────

def _classify_pressure(percent: float) -> str:
    """Classify a resource-usage percentage into a severity label."""
    if percent >= 90:
        return "critical"
    if percent >= 75:
        return "high"
    if percent >= 50:
        return "moderate"
    return "low"


_PROTECTED_PIDS  = {0, 1}
_PROTECTED_NAMES = frozenset({
    "launchd", "systemd", "init", "kernel_task",
    "svchost.exe", "winlogon.exe", "csrss.exe",
    "smss.exe", "wininit.exe", "lsass.exe", "services.exe",
})

# Directories skipped by find_large_files — defined once at module level
# so the set is not re-created on every call.
_FIND_SKIP_DIRS = frozenset({
    ".git", "__pycache__", "node_modules", ".venv", "venv",
    ".Trash", "Library",
})

# FedEx tracking numbers are exactly 12, 15, or 22 digits.
_FEDEX_RE = re.compile(r"^\d{12}$|^\d{15}$|^\d{22}$")


def _detect_cpu_oc(cpu_brand: str, system: str, machine: str) -> dict:
    if machine == "arm64" and system == "Darwin":
        return {"supported": False, "reason": "Apple Silicon CPUs have fixed clock speeds and cannot be overclocked.", "tools": []}
    if system == "Darwin":
        return {"supported": False, "reason": "Intel Macs lack BIOS access required for overclocking.", "tools": []}
    if re.search(r'\bintel\b', cpu_brand, re.I):
        unlocked = bool(re.search(r'\b\d{3,5}[kK][sS]?\b', cpu_brand))
        return {
            "supported": unlocked,
            "reason": ("K/KF/KS-series Intel CPUs support overclocking via BIOS multiplier adjustment."
                       if unlocked else "Non-K Intel CPUs have locked multipliers and cannot be overclocked."),
            "tools": ["Intel Extreme Tuning Utility (XTU)", "BIOS/UEFI"] if unlocked else [],
        }
    if re.search(r'\b(amd|ryzen)\b', cpu_brand, re.I):
        return {
            "supported": True,
            "reason": "AMD Ryzen CPUs support Precision Boost Overdrive (PBO) for automated overclocking and manual clock adjustments.",
            "tools": ["AMD Ryzen Master", "BIOS/UEFI PBO settings"],
        }
    return {"supported": False, "reason": "Could not determine OC capability from CPU brand string.", "tools": []}


def _detect_gpu_oc(system: str, machine: str, gpu_data: dict) -> dict:
    if machine == "arm64":
        return {"supported": False, "reason": "Apple Silicon GPU is integrated into the SoC and cannot be overclocked.", "tools": []}
    if system == "Darwin":
        return {"supported": False, "reason": "macOS does not expose GPU overclocking controls.", "tools": []}
    if "error" in gpu_data:
        return {"supported": False, "reason": "No discrete GPU detected.", "tools": []}
    return {
        "supported": True,
        "reason": "Discrete GPUs on Windows/Linux support overclocking via third-party tools.",
        "tools": ["MSI Afterburner", "EVGA Precision X1", "AMD Radeon Software Adrenalin"],
    }


def _get_upgrade_feasibility(system: str, machine: str) -> dict:
    if machine == "arm64" and system == "Darwin":
        return {
            "ram":     {"upgradeable": False, "note": "Unified memory is soldered to the Apple Silicon SoC — cannot be upgraded."},
            "cpu":     {"upgradeable": False, "note": "CPU is part of the Apple Silicon SoC — cannot be swapped."},
            "gpu":     {"upgradeable": False, "note": "GPU is integrated into the SoC. eGPU support was removed in macOS 14."},
            "storage": {"upgradeable": False, "note": "Internal SSD is proprietary and soldered. External Thunderbolt 4 drives are the only capacity expansion option."},
        }
    if system == "Darwin":
        return {
            "ram":     {"upgradeable": "model-dependent", "note": "Pre-2019 MacBook Pros and some Mac Pros have user-upgradeable RAM — check your exact model."},
            "cpu":     {"upgradeable": False, "note": "Intel Mac CPUs are soldered on most models since 2012."},
            "gpu":     {"upgradeable": "eGPU-only", "note": "Internal GPU not upgradeable. eGPU via Thunderbolt 3 supported on Intel Macs running macOS 13 or earlier."},
            "storage": {"upgradeable": "model-dependent", "note": "Some 2013–2017 MacBook Pro models accept third-party NVMe SSDs via adapters."},
        }
    return {
        "ram":     {"upgradeable": "likely", "note": "Most desktops/laptops support RAM upgrades. Check your motherboard or laptop spec for max supported speed and slot count."},
        "cpu":     {"upgradeable": "varies", "note": "Desktop CPUs are upgradeable if the socket matches. Laptop CPUs are usually soldered — verify your model."},
        "gpu":     {"upgradeable": "likely-desktop", "note": "Desktop PCIe GPUs are freely swappable. Laptop GPUs are typically soldered or MXM (rarely swappable)."},
        "storage": {"upgradeable": "likely", "note": "M.2 NVMe and 2.5-inch SATA slots are widely available. Check how many free slots your system has."},
    }


_USE_CASE_PROFILES = [
    (["lightroom", "photo editing", "photo", "capture one", "darktable"],
     "gpu", "ram",
     "Lightroom's AI features (Denoise, Select Subject, Masking) are GPU-accelerated. Export speed is CPU+GPU bound. Smart Previews and cache performance improve significantly with a fast NVMe SSD."),
    (["premiere", "video editing", "video", "davinci", "resolve", "final cut", "fcpx", "after effects"],
     "gpu", "ram",
     "Video editing benefits most from GPU acceleration (H.264/HEVC decode, effects rendering). RAM is critical for 4K+ multicam timelines. Fast NVMe SSD dramatically improves media cache and scratch disk performance."),
    (["gaming", "games", "game"],
     "gpu", "cpu",
     "Most games are GPU-bound. CPU matters for games with many entities (open-world, RTS). Fast NVMe storage reduces load times. RAM speed (frequency) affects frame pacing on AMD platforms."),
    (["blender", "3d render", "rendering", "maya", "cinema 4d", "c4d", "houdini"],
     "gpu", "ram",
     "GPU rendering (CUDA/OptiX/Metal) is fastest for most 3D renders. VRAM limits scene and texture complexity. CPU rendering uses all physical cores. RAM capacity affects how large a scene can be loaded."),
    (["compile", "compiling", "build", "xcode", "make", "cmake", "gradle", "rust", "go", "code", "coding", "development", "developer"],
     "cpu", "ram",
     "Compilation is highly CPU-bound — more physical cores and higher clock speed both help. RAM limits parallel compile jobs. A fast NVMe SSD dramatically reduces incremental build times via faster cache reads."),
    (["docker", "containers", "kubernetes", "vm", "virtual machine", "virtualbox", "vmware", "parallels"],
     "ram", "cpu",
     "Containers and VMs are RAM-limited first — each VM needs dedicated memory. CPU core count determines how many can run in parallel. Fast storage reduces image pull and disk I/O latency."),
    (["machine learning", "ml", "ai training", "training", "pytorch", "tensorflow", "cuda"],
     "gpu", "ram",
     "ML training is GPU-bound; VRAM limits batch size and model size. CPU handles data loading pipelines. RAM caches the dataset between epochs. Fast NVMe reduces I/O bottlenecks during data loading."),
    (["streaming", "obs", "twitch", "youtube live", "recording"],
     "gpu", "cpu",
     "Streaming with GPU encoding (NVENC/AMF/VideoToolbox) offloads work from the CPU. CPU encoding (x264) produces better quality but is CPU-intensive. RAM and fast storage handle replay buffers and recordings."),
]


def _use_case_analysis(use_case: str, cpu_pct: float, ram_pct: float) -> dict:
    uc = use_case.lower()
    primary, secondary, note = "unknown", "unknown", ""

    for keywords, p, s, n in _USE_CASE_PROFILES:
        if any(k in uc for k in keywords):
            primary, secondary, note = p, s, n
            break

    constraints = []
    if cpu_pct >= 75:
        constraints.append(f"cpu_pressure_{_classify_pressure(cpu_pct)}")
    if ram_pct >= 75:
        constraints.append(f"ram_pressure_{_classify_pressure(ram_pct)}")

    if primary == "unknown":
        note = "Use-case not recognized. Specify a workload (e.g. 'lightroom', 'gaming', 'video editing') for targeted bottleneck analysis."

    return {
        "primary_bottleneck": primary,
        "secondary_bottleneck": secondary,
        "current_constraints": constraints,
        "note": note,
    }


_CHART_TEXT_COLOR = "#cccccc"
_CHART_SPINE_COLOR = "#555555"


def _style_chart_dark(fig: plt.Figure, ax: plt.Axes) -> None:
    """Apply dark-mode styling: transparent background, light text."""
    fig.patch.set_alpha(0)
    ax.set_facecolor("none")
    ax.tick_params(colors=_CHART_TEXT_COLOR)
    ax.xaxis.label.set_color(_CHART_TEXT_COLOR)
    ax.yaxis.label.set_color(_CHART_TEXT_COLOR)
    ax.title.set_color("#ffffff")
    for spine in ax.spines.values():
        spine.set_color(_CHART_SPINE_COLOR)


def _fig_to_b64(fig: plt.Figure) -> str:
    """Serialize a matplotlib figure to a base64 PNG string and close it."""
    try:
        buf = io.BytesIO()
        fig.savefig(buf, format="png", bbox_inches="tight", dpi=144,
                    transparent=True)
        buf.seek(0)
        return base64.b64encode(buf.read()).decode()
    finally:
        plt.close(fig)


def _safe(fn: Callable[[], object]) -> object | None:
    """Call *fn* and swallow psutil access/stale-PID errors, returning None."""
    try:
        return fn()
    except (psutil.AccessDenied, psutil.NoSuchProcess):
        return None


def make_error(id_: int | None, code: int, message: str) -> dict:
    """Build a JSON-RPC 2.0 error response envelope."""
    return {
        "jsonrpc": "2.0",
        "id": id_,
        "error": {"code": code, "message": message}
    }


# ── Tool implementations ──────────────────────────────────────────────────────

def get_cpu_usage() -> dict:
    per_core = psutil.cpu_percent(interval=0.5, percpu=True)
    total = round(sum(per_core) / len(per_core), 1) if per_core else 0.0
    freq = psutil.cpu_freq()
    return {
        "total_percent": total,
        "per_core_percent": per_core,
        "core_count_logical": psutil.cpu_count(logical=True),
        "core_count_physical": psutil.cpu_count(logical=False),
        "frequency_mhz": {
            "current": round(freq.current, 1) if freq else None,
            "min": round(freq.min, 1) if freq else None,
            "max": round(freq.max, 1) if freq else None,
        }
    }


def get_ram_usage() -> dict:
    vm = psutil.virtual_memory()
    sw = psutil.swap_memory()
    return {
        "ram": {
            "total_gb": round(vm.total / 1e9, 2),
            "available_gb": round(vm.available / 1e9, 2),
            "used_gb": round(vm.used / 1e9, 2),
            "percent_used": vm.percent,
        },
        "swap": {
            "total_gb": round(sw.total / 1e9, 2),
            "used_gb": round(sw.used / 1e9, 2),
            "percent_used": sw.percent,
        }
    }


def get_gpu_usage() -> dict:
    if not GPU_AVAILABLE:
        return {"error": "No supported GPU backend found. Install nvidia-ml-py to enable GPU monitoring."}

    try:
        handles = _get_nvml_handles()
        if not handles:
            return {"error": "No GPUs detected"}
        gpus = []
        for i, h in enumerate(handles):
            util = pynvml.nvmlDeviceGetUtilizationRates(h)
            mem  = pynvml.nvmlDeviceGetMemoryInfo(h)
            try:
                temp = pynvml.nvmlDeviceGetTemperature(h, pynvml.NVML_TEMPERATURE_GPU)
            except pynvml.NVMLError:
                temp = None
            name = pynvml.nvmlDeviceGetName(h)
            if isinstance(name, bytes):
                name = name.decode()
            mem_total_mb = mem.total / 1024 / 1024
            mem_used_mb  = mem.used  / 1024 / 1024
            gpus.append({
                "id": i,
                "name": name,
                "load_percent": util.gpu,
                "memory_used_mb": round(mem_used_mb, 1),
                "memory_total_mb": round(mem_total_mb, 1),
                "memory_percent": round(mem_used_mb / mem_total_mb * 100, 1) if mem_total_mb else None,
                "temperature_c": temp,
            })
        return {"gpus": gpus}
    except pynvml.NVMLError as e:
        global _NVML_HANDLES, _NVML_HANDLES_READY
        with _NVML_LOCK:
            _NVML_HANDLES = []  # invalidate cache on NVML error so next call retries
            _NVML_HANDLES_READY = False
        return {"error": f"NVML error: {e}"}


def get_disk_usage() -> dict:
    partitions = []
    for part in psutil.disk_partitions(all=False):
        try:
            usage = psutil.disk_usage(part.mountpoint)
            partitions.append({
                "device": part.device,
                "mountpoint": part.mountpoint,
                "fstype": part.fstype,
                "total_gb": round(usage.total / 1e9, 2),
                "used_gb": round(usage.used / 1e9, 2),
                "free_gb": round(usage.free / 1e9, 2),
                "percent_used": usage.percent,
            })
        except (PermissionError, OSError):
            continue
    disk_io = psutil.disk_io_counters()
    return {
        "partitions": partitions,
        "io_counters": {
            "read_mb": round(disk_io.read_bytes / 1e6, 2) if disk_io else None,
            "write_mb": round(disk_io.write_bytes / 1e6, 2) if disk_io else None,
        }
    }


def get_network_usage() -> dict:
    net_io = psutil.net_io_counters()
    interfaces = {}
    for iface, stats in psutil.net_if_stats().items():
        interfaces[iface] = {
            "is_up": stats.isup,
            "speed_mbps": stats.speed,
        }
    total_io: dict[str, float | int | None]
    if net_io is not None:
        total_io = {
            "bytes_sent_mb": round(net_io.bytes_sent / 1e6, 2),
            "bytes_recv_mb": round(net_io.bytes_recv / 1e6, 2),
            "packets_sent": net_io.packets_sent,
            "packets_recv": net_io.packets_recv,
        }
    else:
        total_io = {
            "bytes_sent_mb": None, "bytes_recv_mb": None,
            "packets_sent": None, "packets_recv": None,
        }
    return {"total_io": total_io, "interfaces": interfaces}


def get_realtime_io(interval: int = 1) -> dict:
    interval = max(1, min(interval, 3))
    d1 = psutil.disk_io_counters()
    n1 = psutil.net_io_counters()
    time.sleep(interval)
    d2 = psutil.disk_io_counters()
    n2 = psutil.net_io_counters()
    dt = float(interval)

    if d1 is not None and d2 is not None:
        read_mbs = round((d2.read_bytes - d1.read_bytes) / 1e6 / dt, 3)
        write_mbs = round((d2.write_bytes - d1.write_bytes) / 1e6 / dt, 3)
        disk_ok = True
    else:
        read_mbs = write_mbs = None
        disk_ok = False

    if n1 is not None and n2 is not None:
        dl_mbs = round((n2.bytes_recv - n1.bytes_recv) / 1e6 / dt, 3)
        ul_mbs = round((n2.bytes_sent - n1.bytes_sent) / 1e6 / dt, 3)
        net_info: dict = {
            "download_mbs": dl_mbs,
            "upload_mbs": ul_mbs,
            "download_mbps": round(dl_mbs * 8, 3),
            "upload_mbps": round(ul_mbs * 8, 3),
        }
    else:
        net_info = {
            "download_mbs": None, "upload_mbs": None,
            "download_mbps": None, "upload_mbps": None,
        }

    return {
        "interval_seconds": interval,
        "disk": {"available": disk_ok, "read_mbs": read_mbs, "write_mbs": write_mbs},
        "network": net_info,
    }


def get_top_processes(n: int = 10, sort_by: str = "cpu") -> dict:
    """Return top N processes sorted by cpu or memory."""
    n = max(1, min(n, 100))
    procs = []
    for p in psutil.process_iter(['pid', 'name', 'cpu_percent', 'memory_percent', 'status', 'num_threads']):
        try:
            info = p.info
            procs.append(info)
        except (psutil.NoSuchProcess, psutil.AccessDenied):
            continue

    key = "memory_percent" if sort_by == "memory" else "cpu_percent"
    procs.sort(key=lambda x: x.get(key) or 0, reverse=True)

    return {
        "sort_by": sort_by,
        "top_processes": [
            {
                "pid": p["pid"],
                "name": p["name"],
                "cpu_percent": round(p.get("cpu_percent") or 0, 2),
                "memory_percent": round(p.get("memory_percent") or 0, 2),
                "status": p.get("status"),
                "threads": p.get("num_threads"),
            }
            for p in procs[:n]
        ]
    }


def _cpu_with_chart() -> dict | tuple[dict, str]:
    data = get_cpu_usage()
    cores = data["per_core_percent"]
    n = len(cores)

    fig, ax = plt.subplots(figsize=(7, max(3, n * 0.4)))
    try:
        colors = ["#e74c3c" if v >= 80 else "#e67e22" if v >= 60 else "#2ecc71" for v in cores]
        labels = [f"Core {i}" for i in range(n)]
        ax.barh(labels, cores, color=colors, height=0.6)
        ax.axvline(data["total_percent"], color="#3498db", linestyle="--", linewidth=1.5,
                   label=f'Total: {data["total_percent"]}%')
        ax.set_xlim(0, 100)
        ax.set_xlabel("Usage %")
        ax.set_title("CPU Usage per Core")
        ax.legend(loc="lower right", fontsize=8, facecolor="#2a2a2a", edgecolor="#555",
                  labelcolor=_CHART_TEXT_COLOR)
        ax.xaxis.set_major_formatter(mticker.PercentFormatter())
        ax.set_yticklabels(labels, color=_CHART_TEXT_COLOR)
        _style_chart_dark(fig, ax)
        fig.tight_layout()
        return data, _fig_to_b64(fig)
    except Exception as exc:
        sys.stderr.write(f"[syscontrol] _cpu_with_chart rendering failed: {exc}\n")
        plt.close(fig)
        return data


def _ram_with_chart() -> dict | tuple[dict, str]:
    data = get_ram_usage()
    ram = data["ram"]
    swap = data["swap"]

    fig, ax = plt.subplots(figsize=(7, 2.5))
    try:
        labels = ["RAM", "Swap"]
        ax.barh(["RAM"],  [ram["used_gb"]],                                    color="#e74c3c", label="Used")
        ax.barh(["RAM"],  [ram["available_gb"]], left=[ram["used_gb"]],         color="#2ecc71", label="Available")
        ax.barh(["Swap"], [swap["used_gb"]],                                    color="#e67e22")
        ax.barh(["Swap"], [swap["total_gb"] - swap["used_gb"]], left=[swap["used_gb"]], color="#95a5a6")
        ax.set_xlabel("GB")
        ax.set_title("Memory Usage")
        ax.legend(loc="lower right", fontsize=8, facecolor="#2a2a2a", edgecolor="#555",
                  labelcolor=_CHART_TEXT_COLOR)
        ax.set_yticklabels(labels, color=_CHART_TEXT_COLOR)
        for bar in ax.patches:
            w = bar.get_width()
            if w > 0.3:
                ax.text(bar.get_x() + w / 2, bar.get_y() + bar.get_height() / 2,
                        f"{w:.1f} GB", ha="center", va="center", fontsize=7, color="white")
        _style_chart_dark(fig, ax)
        fig.tight_layout()
        return data, _fig_to_b64(fig)
    except Exception as exc:
        sys.stderr.write(f"[syscontrol] _ram_with_chart rendering failed: {exc}\n")
        plt.close(fig)
        return data


def _gpu_with_chart() -> dict | tuple[dict, str]:
    data = get_gpu_usage()
    if "error" in data or not data.get("gpus"):
        return data

    gpus = data["gpus"]
    x = list(range(len(gpus)))
    w = 0.25

    fig, ax = plt.subplots(figsize=(7, 3.5))
    try:
        ax.bar([i - w for i in x], [g.get("load_percent") or 0 for g in gpus], width=w, label="Load %",  color="#3498db")
        ax.bar([i      for i in x], [g.get("memory_percent") or 0 for g in gpus], width=w, label="VRAM %",  color="#9b59b6")
        ax.bar([i + w  for i in x], [g.get("temperature_c") or 0  for g in gpus], width=w, label="Temp °C", color="#e74c3c")
        ax.set_xticks(x)
        ax.set_xticklabels([g["name"] for g in gpus], fontsize=8, color=_CHART_TEXT_COLOR)
        ax.set_ylim(0, 110)
        ax.set_ylabel("% / °C")
        ax.set_title("GPU Metrics")
        ax.legend(fontsize=8, facecolor="#2a2a2a", edgecolor="#555",
                  labelcolor=_CHART_TEXT_COLOR)
        _style_chart_dark(fig, ax)
        fig.tight_layout()
        return data, _fig_to_b64(fig)
    except Exception as exc:
        sys.stderr.write(f"[syscontrol] _gpu_with_chart rendering failed: {exc}\n")
        plt.close(fig)
        return data


def get_hardware_profile(use_case: str = "") -> dict:
    """Aggregate hardware specs, live pressure, OC capability, upgrade feasibility, and use-case bottleneck analysis."""
    # Run all four independent data-source calls in parallel.
    f_specs = _METRICS_EXECUTOR.submit(get_device_specs)
    f_cpu   = _METRICS_EXECUTOR.submit(get_cpu_usage)
    f_ram   = _METRICS_EXECUTOR.submit(get_ram_usage)
    f_gpu   = _METRICS_EXECUTOR.submit(get_gpu_usage)
    specs    = f_specs.result()
    cpu_live = f_cpu.result()
    ram_live = f_ram.result()
    gpu_data = f_gpu.result()

    system    = specs["os"]["system"]
    machine   = specs["os"]["machine"]
    cpu_brand = specs["cpu"]["brand"]
    cpu_pct   = cpu_live["total_percent"]
    ram_pct   = ram_live["ram"]["percent_used"]

    return {
        "use_case": use_case,
        "hardware": {
            "cpu":    specs["cpu"],
            "ram":    {"total_gb": specs["ram"]["total_gb"]},
            "gpu":    specs["gpus"],
            "disks":  specs["disks"],
        },
        "current_pressure": {
            "cpu": {"percent": cpu_pct, "level": _classify_pressure(cpu_pct)},
            "ram": {"percent": ram_pct, "level": _classify_pressure(ram_pct)},
        },
        "platform": {
            "system":           system,
            "machine":          machine,
            "is_apple_silicon": machine == "arm64" and system == "Darwin",
        },
        "overclocking": {
            "cpu": _detect_cpu_oc(cpu_brand, system, machine),
            "gpu": _detect_gpu_oc(system, machine, gpu_data),
        },
        "upgrade_feasibility": _get_upgrade_feasibility(system, machine),
        "use_case_analysis":   _use_case_analysis(use_case, cpu_pct, ram_pct),
    }


def get_battery_status() -> dict:
    batt = psutil.sensors_battery()
    if batt is None:
        return {"error": "No battery detected (desktop or unsupported platform)"}
    return {
        "percent": round(batt.percent, 1),
        "plugged_in": batt.power_plugged,
        "time_remaining_min": round(batt.secsleft / 60, 1) if batt.secsleft > 0 else None,
    }


def get_temperature_sensors() -> dict:
    if IS_MACOS:
        return {
            "platform": "macOS",
            "available": False,
            "sensors": {},
            "message": (
                "psutil cannot access CPU/motherboard sensors on macOS. "
                "Alternatives: (1) GPU temp via get_gpu_usage if discrete GPU present. "
                "(2) iStatMenus or HWMonitor for full sensor access. "
                "(3) On Apple Silicon, thermal throttling shows as current_mhz << max_mhz in get_cpu_usage."
            ),
        }
    if not hasattr(psutil, "sensors_temperatures"):
        return {
            "platform": _SYSTEM,
            "available": False,
            "sensors": {},
            "message": "psutil.sensors_temperatures() not available on this platform/version.",
        }
    try:
        raw = psutil.sensors_temperatures()
    except Exception as e:
        return {"platform": _SYSTEM, "available": False, "sensors": {}, "message": f"Failed to read sensors: {e}"}
    if not raw:
        return {
            "platform": _SYSTEM,
            "available": True,
            "sensors": {},
            "message": "No sensors detected (may require elevated privileges on Linux).",
        }
    sensors = {}
    for chip, entries in raw.items():
        sensors[chip] = [
            {
                "label": e.label or chip,
                "current_c": round(e.current, 1) if e.current is not None else None,
                "high_c": round(e.high, 1) if e.high is not None else None,
                "critical_c": round(e.critical, 1) if e.critical is not None else None,
            }
            for e in entries
        ]
    return {"platform": _SYSTEM, "available": True, "message": "", "sensors": sensors}


def get_system_uptime() -> dict:
    boot = psutil.boot_time()
    elapsed = int(datetime.datetime.now().timestamp() - boot)
    return {
        "boot_time": datetime.datetime.fromtimestamp(boot).isoformat(),
        "uptime": {
            "days": elapsed // 86400,
            "hours": (elapsed % 86400) // 3600,
            "minutes": (elapsed % 3600) // 60,
        },
        "load_avg_1_5_15min": list(psutil.getloadavg()),
    }


def _check_gpu_alerts(alerts: list[dict]) -> None:
    """Append GPU load and temperature alerts to *alerts* if pynvml is available."""
    if not GPU_AVAILABLE or GPU_BACKEND != "pynvml":
        return
    try:
        for i, h in enumerate(_get_nvml_handles()):
            util = pynvml.nvmlDeviceGetUtilizationRates(h)
            name = pynvml.nvmlDeviceGetName(h)
            if isinstance(name, bytes):
                name = name.decode()
            load_pct = util.gpu
            if load_pct >= 95:
                alerts.append({"severity": "critical", "resource": f"gpu:{i}",
                    "message": f"GPU {name} load critically high at {load_pct}%", "value": load_pct})
            try:
                temp = pynvml.nvmlDeviceGetTemperature(h, pynvml.NVML_TEMPERATURE_GPU)
                if temp >= 85:
                    alerts.append({"severity": "critical", "resource": f"gpu:{i}",
                        "message": f"GPU {name} temp critically high at {temp}°C", "value": temp})
                elif temp >= 75:
                    alerts.append({"severity": "warning", "resource": f"gpu:{i}",
                        "message": f"GPU {name} temp elevated at {temp}°C", "value": temp})
            except pynvml.NVMLError:
                pass
    except Exception as exc:
        sys.stderr.write(f"[syscontrol] GPU alert check failed: {exc}\n")


def _build_alert_summary(alerts: list[dict]) -> str:
    """Return a human-readable summary string for the given alert list."""
    if not alerts:
        return "All systems nominal — no alerts detected."
    critical_n = sum(1 for a in alerts if a["severity"] == "critical")
    warning_n  = sum(1 for a in alerts if a["severity"] == "warning")
    if critical_n:
        return (f"{critical_n} critical and {warning_n} warning alert(s) detected."
                " Immediate attention recommended.")
    return f"{len(alerts)} warning(s) detected. System under stress but not critical."


def get_system_alerts() -> dict:
    """Return resource alerts for CPU, RAM, swap, disk, GPU, and battery."""
    alerts: list[dict] = []

    cpu_pct = psutil.cpu_percent(interval=0.5)
    if cpu_pct >= 90:
        alerts.append({"severity": "critical", "resource": "cpu",
            "message": f"CPU usage critically high at {cpu_pct}%", "value": cpu_pct})
    elif cpu_pct >= 75:
        alerts.append({"severity": "warning", "resource": "cpu",
            "message": f"CPU usage elevated at {cpu_pct}%", "value": cpu_pct})

    vm = psutil.virtual_memory()
    if vm.percent >= 90:
        alerts.append({"severity": "critical", "resource": "ram",
            "message": f"RAM critically high at {vm.percent}%", "value": vm.percent})
    elif vm.percent >= 75:
        alerts.append({"severity": "warning", "resource": "ram",
            "message": f"RAM elevated at {vm.percent}%", "value": vm.percent})

    try:
        sw = psutil.swap_memory()
        if sw.total > 0 and sw.percent >= 80:
            alerts.append({"severity": "warning", "resource": "swap",
                "message": f"Swap high at {sw.percent}% — system may be memory-constrained",
                "value": sw.percent})
    except Exception as exc:
        sys.stderr.write(f"[syscontrol] swap read failed: {exc}\n")

    for part in psutil.disk_partitions(all=False):
        try:
            usage = psutil.disk_usage(part.mountpoint)
            if usage.percent >= 95:
                alerts.append({"severity": "critical", "resource": f"disk:{part.mountpoint}",
                    "message": f"Disk {part.mountpoint} almost full at {usage.percent}%",
                    "value": usage.percent})
            elif usage.percent >= 85:
                alerts.append({"severity": "warning", "resource": f"disk:{part.mountpoint}",
                    "message": f"Disk {part.mountpoint} getting full at {usage.percent}%",
                    "value": usage.percent})
        except (PermissionError, OSError):
            continue

    _check_gpu_alerts(alerts)

    batt = psutil.sensors_battery()
    if batt is not None and not batt.power_plugged and batt.percent <= 10:
        alerts.append({"severity": "critical", "resource": "battery",
            "message": f"Battery critically low at {batt.percent}% and not plugged in",
            "value": batt.percent})

    summary    = _build_alert_summary(alerts)
    has_critical = any(a["severity"] == "critical" for a in alerts)
    return {
        "alerts": alerts,
        "alert_count": len(alerts),
        "has_critical": has_critical,
        "summary": summary,
    }


def get_network_connections() -> dict:
    try:
        raw_connections = psutil.net_connections(kind="inet")
    except psutil.AccessDenied:
        return {"error": "Access denied. Network connection listing may require elevated privileges.", "connections": [], "total": 0}

    # Build a PID→name map once from process_iter instead of constructing
    # a new psutil.Process object for every connection (O(n) not O(n·k)).
    pid_to_name: dict[int, str] = {}
    for p in psutil.process_iter(["pid", "name"]):
        with contextlib.suppress(psutil.NoSuchProcess, psutil.AccessDenied):
            pid_to_name[p.info["pid"]] = p.info["name"] or ""

    connections = [
        {
            "proto":   "tcp" if conn.type == socket.SOCK_STREAM else "udp",
            "local":   f"{conn.laddr.ip}:{conn.laddr.port}" if conn.laddr else None,
            "remote":  f"{conn.raddr.ip}:{conn.raddr.port}" if conn.raddr else None,
            "status":  conn.status,
            "pid":     conn.pid,
            "process": pid_to_name.get(conn.pid) if conn.pid else None,
        }
        for conn in raw_connections
    ]
    return {"connections": connections, "total": len(connections)}


def _get_startup_items_macos() -> dict:
    """Return launch agents / daemons from standard macOS plist directories."""
    scan_dirs = [
        (pathlib.Path.home() / "Library" / "LaunchAgents", "user"),
        (pathlib.Path("/Library/LaunchAgents"),              "system"),
        (pathlib.Path("/Library/LaunchDaemons"),             "system-daemon"),
    ]
    items: list[dict] = []
    for directory, scope in scan_dirs:
        if not directory.exists():
            continue
        for plist_path in sorted(directory.glob("*.plist")):
            try:
                with open(plist_path, "rb") as fh:
                    data = plistlib.load(fh)
                prog_args = data.get("ProgramArguments", [])
                command   = (
                    " ".join(str(a) for a in prog_args)
                    if prog_args else data.get("Program", "")
                )
                items.append({
                    "name":        data.get("Label") or plist_path.stem,
                    "command":     command,
                    "path":        str(plist_path),
                    "scope":       scope,
                    "run_at_load": bool(data.get("RunAtLoad", False)),
                })
            except (plistlib.InvalidFileException, OSError, KeyError, TypeError):
                items.append({
                    "name":        plist_path.stem,
                    "command":     "",
                    "path":        str(plist_path),
                    "scope":       scope,
                    "run_at_load": None,
                    "parse_error": True,
                })
    return {"platform": "macOS", "items": items, "count": len(items)}


def _get_startup_items_windows() -> dict:
    """Return startup entries from the Windows registry Run keys."""
    try:
        import winreg
    except ImportError:
        return {"platform": "Windows", "error": "winreg not available", "items": [], "count": 0}

    items: list[dict] = []
    run_keys = [
        (winreg.HKEY_CURRENT_USER,
         r"SOFTWARE\Microsoft\Windows\CurrentVersion\Run", "user"),
        (winreg.HKEY_LOCAL_MACHINE,
         r"SOFTWARE\Microsoft\Windows\CurrentVersion\Run", "system"),
    ]
    for hive, reg_path, scope in run_keys:
        try:
            key = winreg.OpenKey(hive, reg_path, 0, winreg.KEY_READ)
            i = 0
            while True:
                try:
                    name, value, _ = winreg.EnumValue(key, i)
                    items.append({"name": name, "command": value, "scope": scope})
                    i += 1
                except OSError:
                    break
            winreg.CloseKey(key)
        except OSError:
            continue
    return {"platform": "Windows", "items": items, "count": len(items)}


def _get_startup_items_linux() -> dict:
    """Return XDG autostart entries from ``~/.config/autostart/*.desktop``."""
    autostart = pathlib.Path.home() / ".config" / "autostart"
    items: list[dict] = []
    if autostart.exists():
        for dp in sorted(autostart.glob("*.desktop")):
            try:
                text    = dp.read_text(encoding="utf-8", errors="replace")
                name    = ""
                command = ""
                hidden  = False
                for line in text.splitlines():
                    if line.startswith("Name="):
                        name = line[5:].strip()
                    elif line.startswith("Exec="):
                        command = line[5:].strip()
                    elif line.startswith("Hidden="):
                        hidden = line[7:].strip().lower() == "true"
                items.append({
                    "name":    name or dp.stem,
                    "command": command,
                    "path":    str(dp),
                    "scope":   "user",
                    "hidden":  hidden,
                })
            except OSError:
                continue
    return {"platform": "Linux", "items": items, "count": len(items)}


def get_startup_items() -> dict:
    """Return platform startup items (launchagents, registry run keys, XDG autostart)."""
    if IS_MACOS:
        return _get_startup_items_macos()
    if IS_WIN:
        return _get_startup_items_windows()
    if IS_LINUX:
        return _get_startup_items_linux()
    return {"platform": _SYSTEM, "error": f"Not supported on {_SYSTEM}", "items": [], "count": 0}


def get_process_details(pid: int) -> dict:
    if pid <= 0:
        return {"error": f"Invalid PID {pid}: must be a positive integer"}
    try:
        p = psutil.Process(pid)
        with p.oneshot():
            mem = p.memory_info()
            return {
                "pid": pid,
                "name": p.name(),
                "exe": _safe(p.exe),
                "cmdline": _safe(lambda: " ".join(p.cmdline())),
                "user": _safe(p.username),
                "status": p.status(),
                "created": datetime.datetime.fromtimestamp(p.create_time()).isoformat(),
                "cpu_percent": p.cpu_percent(interval=0.2),
                "memory": {
                    "rss_mb": round(mem.rss / 1e6, 2),
                    "vms_mb": round(mem.vms / 1e6, 2),
                    "percent": round(p.memory_percent(), 2),
                },
                "threads": p.num_threads(),
                "open_files": _safe(lambda: len(p.open_files())),
            }
    except psutil.NoSuchProcess:
        return {"error": f"No process with PID {pid}"}
    except psutil.AccessDenied:
        return {"error": f"Access denied reading process details for PID {pid}"}


def search_process(name: str) -> dict:
    if not name or not name.strip():
        return {
            "error": "Search query cannot be empty",
            "query": name,
            "matches": [],
            "count": 0,
        }
    name = name.strip()
    name_lower = name.lower()
    matches = []
    for p in psutil.process_iter(['pid', 'name', 'cpu_percent', 'memory_percent', 'status']):
        try:
            if name_lower in (p.info['name'] or '').lower():
                matches.append({
                    "pid": p.info['pid'],
                    "name": p.info['name'],
                    "cpu_percent": round(p.info['cpu_percent'] or 0, 2),
                    "memory_percent": round(p.info['memory_percent'] or 0, 2),
                    "status": p.info['status'],
                })
        except (psutil.NoSuchProcess, psutil.AccessDenied):
            continue
    return {"query": name, "matches": matches, "count": len(matches)}


def kill_process(pid: int, force: bool = False) -> dict:
    if not isinstance(pid, int):
        return {"success": False, "error": f"PID must be an integer, got {type(pid).__name__}"}
    if pid <= 0:
        return {"success": False, "error": f"Invalid PID {pid}: must be a positive integer"}
    if pid in _PROTECTED_PIDS:
        return {"success": False, "error": f"Refusing to kill PID {pid}: protected system process"}
    try:
        p = psutil.Process(pid)
        proc_name = p.name()
    except psutil.NoSuchProcess:
        return {"success": False, "error": f"No process with PID {pid}"}
    except psutil.AccessDenied:
        return {"success": False, "error": f"Access denied reading PID {pid}"}

    if proc_name.lower() in _PROTECTED_NAMES:
        return {
            "success": False,
            "error": f"Refusing to kill '{proc_name}' (PID {pid}): critical system process",
        }

    try:
        if force:
            p.kill()
            method = "SIGKILL"
        else:
            p.terminate()
            method = "SIGTERM"
        return {
            "success": True,
            "pid": pid,
            "name": proc_name,
            "signal": method,
            "message": f"Sent {method} to '{proc_name}' (PID {pid})",
        }
    except psutil.NoSuchProcess:
        return {"success": False, "error": f"Process {pid} exited before signal could be sent"}
    except psutil.AccessDenied:
        return {
            "success": False,
            "error": f"Access denied killing '{proc_name}' (PID {pid}). May require elevated privileges.",
        }


@lru_cache(maxsize=1)
def get_device_specs() -> dict:
    """Return static hardware and OS specifications."""
    vm = psutil.virtual_memory()
    freq = psutil.cpu_freq()

    disks = []
    for part in psutil.disk_partitions(all=False):
        try:
            usage = psutil.disk_usage(part.mountpoint)
            disks.append({
                "device": part.device,
                "mountpoint": part.mountpoint,
                "fstype": part.fstype,
                "total_gb": round(usage.total / 1e9, 2),
            })
        except PermissionError:
            continue

    gpu_specs = []
    if GPU_AVAILABLE:
        try:
            if GPU_BACKEND == "pynvml":
                for h in _get_nvml_handles():
                    name = pynvml.nvmlDeviceGetName(h)
                    if isinstance(name, bytes):
                        name = name.decode()
                    mem = pynvml.nvmlDeviceGetMemoryInfo(h)
                    gpu_specs.append({
                        "name": name,
                        "vram_total_mb": round(mem.total / 1024 / 1024, 1),
                    })
        except Exception:
            pass

    return {
        "os": {
            "system": _SYSTEM,
            "release": platform.release(),
            "version": platform.version(),
            "machine": _MACHINE,
            "hostname": platform.node(),
        },
        "cpu": {
            "brand": platform.processor(),
            "physical_cores": psutil.cpu_count(logical=False),
            "logical_cores": psutil.cpu_count(logical=True),
            "max_frequency_mhz": round(freq.max, 1) if freq else None,
        },
        "ram": {
            "total_gb": round(vm.total / 1e9, 2),
        },
        "gpus": gpu_specs or [{"error": "pynvml not available or no NVIDIA GPUs detected"}],
        "disks": disks,
    }


def get_full_snapshot() -> dict:
    """Aggregate snapshot of all metrics — all sources fetched in parallel."""
    f_cpu     = _METRICS_EXECUTOR.submit(get_cpu_usage)
    f_ram     = _METRICS_EXECUTOR.submit(get_ram_usage)
    f_gpu     = _METRICS_EXECUTOR.submit(get_gpu_usage)
    f_disk    = _METRICS_EXECUTOR.submit(get_disk_usage)
    f_net     = _METRICS_EXECUTOR.submit(get_network_usage)
    f_top_cpu = _METRICS_EXECUTOR.submit(get_top_processes, 5, "cpu")
    f_top_mem = _METRICS_EXECUTOR.submit(get_top_processes, 5, "memory")
    return {
        "cpu":                    f_cpu.result(),
        "ram":                    f_ram.result(),
        "gpu":                    f_gpu.result(),
        "disk":                   f_disk.result(),
        "network":                f_net.result(),
        "top_processes_by_cpu":    f_top_cpu.result()["top_processes"],
        "top_processes_by_memory": f_top_mem.result()["top_processes"],
    }


# ── Agentic tool helpers ───────────────────────────────────────────────────────

# Pre-compiled regex patterns for _parse_reminder_time (compiled once at module load).
_RE_COMPOUND = re.compile(r"in\s+(\d+)\s+hours?\s+(?:and\s+)?(\d+)\s+minutes?")
_RE_RELATIVE = re.compile(r"in\s+(\d+)\s+(\w+)")
_RE_TOMORROW = re.compile(r"tomorrow\s+at\s+(\d{1,2})(?::(\d{2}))?\s*(am|pm)?")
_RE_AT_TIME  = re.compile(r"at\s+(\d{1,2})(?::(\d{2}))?\s*(am|pm)?")

_RELATIVE_UNITS = {
    "second": 1, "seconds": 1,
    "minute": 60, "minutes": 60,
    "hour": 3600, "hours": 3600,
    "day": 86400, "days": 86400,
    "week": 604800, "weeks": 604800,
}


def _parse_reminder_time(s: str) -> datetime.datetime | None:
    """Parse natural-language time string into a datetime. Returns None on failure."""
    s = s.strip().lower()
    now = datetime.datetime.now()

    # "in 2 hours 30 minutes" (compound)
    m = _RE_COMPOUND.match(s)
    if m:
        return now + datetime.timedelta(hours=int(m.group(1)), minutes=int(m.group(2)))

    # "in 2 hours" / "in 30 minutes" / "in 1 day"
    m = _RE_RELATIVE.match(s)
    if m:
        unit = _RELATIVE_UNITS.get(m.group(2))
        if unit:
            return now + datetime.timedelta(seconds=int(m.group(1)) * unit)

    # "tomorrow at 9:00 am" / "tomorrow at 3pm"
    m = _RE_TOMORROW.match(s)
    if m:
        hour, minute = int(m.group(1)), int(m.group(2) or 0)
        period = m.group(3)
        if period == "pm" and hour < 12:
            hour += 12
        if period == "am" and hour == 12:
            hour = 0
        return (now + datetime.timedelta(days=1)).replace(
            hour=hour, minute=minute, second=0, microsecond=0
        )

    # "at 9:00 am" / "at 14:30" / "at 3pm"
    m = _RE_AT_TIME.match(s)
    if m:
        hour, minute = int(m.group(1)), int(m.group(2) or 0)
        period = m.group(3)
        if period == "pm" and hour < 12:
            hour += 12
        if period == "am" and hour == 12:
            hour = 0
        target = now.replace(hour=hour, minute=minute, second=0, microsecond=0)
        if target <= now:
            target += datetime.timedelta(days=1)
        return target

    return None


def _human_timedelta(delta: datetime.timedelta) -> str:
    """Format a timedelta as a short human-readable string."""
    secs = int(delta.total_seconds())
    if secs < 0:
        return "overdue"
    if secs < 60:
        return f"{secs} seconds"
    if secs < 3600:
        return f"{secs // 60} minutes"
    if secs < 86400:
        return f"{secs // 3600} hours {(secs % 3600) // 60} minutes"
    return f"{secs // 86400} days"


# ── Reminder tools ────────────────────────────────────────────────────────────

def set_reminder(message: str, time_str: str) -> dict:
    fire_at = _parse_reminder_time(time_str)
    if fire_at is None:
        return {
            "success": False,
            "error": (
                f"Could not parse time '{time_str}'. "
                "Try: 'in 2 hours', 'in 30 minutes', 'at 9:00 am', 'at 3pm', 'tomorrow at 8am'."
            ),
        }
    reminder_id = uuid.uuid4().hex[:8]
    entry = {
        "id": reminder_id,
        "message": message,
        "fire_at": fire_at.isoformat(),
        "created_at": datetime.datetime.now().isoformat(),
        "fired": False,
    }
    with _REMINDER_LOCK:
        reminders = _load_reminders()
        reminders.append(entry)
        _save_reminders(reminders)
    return {
        "success": True,
        "id": reminder_id,
        "message": message,
        "fires_at": fire_at.strftime("%Y-%m-%d %I:%M %p"),
        "fires_in": _human_timedelta(fire_at - datetime.datetime.now()),
    }


def list_reminders() -> dict:
    with _REMINDER_LOCK:
        reminders = _load_reminders()
    now = datetime.datetime.now()
    pending = [r for r in reminders if not r["fired"]]
    return {
        "count": len(pending),
        "reminders": [
            {
                "id": r["id"],
                "message": r["message"],
                "fires_at": r["fire_at"],
                "fires_in": _human_timedelta(
                    datetime.datetime.fromisoformat(r["fire_at"]) - now
                ),
            }
            for r in pending
        ],
    }


def cancel_reminder(reminder_id: str) -> dict:
    with _REMINDER_LOCK:
        reminders = _load_reminders()
        original_len = len(reminders)
        reminders = [r for r in reminders if not (r["id"] == reminder_id and not r["fired"])]
        if len(reminders) == original_len:
            return {"success": False, "error": f"No active reminder with id '{reminder_id}'"}
        _save_reminders(reminders)
    return {"success": True, "cancelled_id": reminder_id}


# ── Weather tool ──────────────────────────────────────────────────────────────

_WMO_DESCRIPTIONS = {
    0: "Clear sky",
    1: "Mainly clear", 2: "Partly cloudy", 3: "Overcast",
    45: "Fog", 48: "Freezing fog",
    51: "Light drizzle", 53: "Moderate drizzle", 55: "Dense drizzle",
    56: "Freezing drizzle (light)", 57: "Freezing drizzle (heavy)",
    61: "Light rain", 63: "Moderate rain", 65: "Heavy rain",
    66: "Freezing rain (light)", 67: "Freezing rain (heavy)",
    71: "Light snow", 73: "Moderate snow", 75: "Heavy snow",
    77: "Snow grains",
    80: "Light rain showers", 81: "Moderate rain showers", 82: "Violent rain showers",
    85: "Light snow showers", 86: "Heavy snow showers",
    95: "Thunderstorm", 96: "Thunderstorm with hail", 99: "Thunderstorm with heavy hail",
}

_SNOW_CODES = {71, 73, 75, 77, 85, 86}
_RAIN_CODES = {51, 53, 55, 56, 57, 61, 63, 65, 66, 67, 80, 81, 82}
_FOG_CODES  = {45, 48}


def _clothing_suggestions(temp_f: float, code: int, wind_mph: float, humidity_pct: float) -> list:
    suggestions = []
    if temp_f < 10:
        suggestions.append("Extreme cold: insulated parka, thermal underlayers, insulated waterproof boots, face mask, and thick gloves")
    elif temp_f < 25:
        suggestions.append("Heavy winter coat, thermal underlayers, warm hat, insulated gloves, and winter boots")
    elif temp_f < 40:
        suggestions.append("Winter coat, warm sweater or fleece, gloves, and a hat")
    elif temp_f < 55:
        suggestions.append("Medium jacket or fleece and long pants")
    elif temp_f < 68:
        suggestions.append("Light jacket or cardigan and long pants or jeans")
    elif temp_f < 80:
        suggestions.append("T-shirt or light long-sleeve and comfortable pants or shorts")
    else:
        suggestions.append("Light, breathable clothing — stay hydrated")

    if code in _SNOW_CODES:
        suggestions.append("Snow expected: wear waterproof boots and a snow-resistant outer layer")
    elif code in _RAIN_CODES:
        suggestions.append("Rain expected: bring a rain jacket or umbrella and waterproof footwear")
    elif code in _FOG_CODES:
        suggestions.append("Foggy conditions: drive carefully and use low-beam headlights")

    if wind_mph >= 25:
        suggestions.append("Strong winds: a windproof outer layer is important")
    elif wind_mph >= 15:
        suggestions.append("Breezy: a windbreaker helps")

    if temp_f >= 75 and humidity_pct >= 70:
        suggestions.append("High humidity: moisture-wicking fabrics recommended")

    return suggestions


def _resolve_location_coords(
    location: str,
) -> tuple[float, float, str, str, str, str] | dict:
    """Resolve *location* string to GPS coordinates and place labels.

    Args:
        location: City/address string to geocode, or empty string for IP auto-detect.

    Returns:
        On success: ``(lat, lon, city_name, region, country, source)`` tuple.
        On failure: an error dict suitable for returning directly from the tool.
    """
    if location.strip():
        encoded  = urllib.parse.quote(location.strip())
        url      = (
            f"https://nominatim.openstreetmap.org/search"
            f"?q={encoded}&format=json&limit=1"
        )
        req = urllib.request.Request(url, headers={"User-Agent": "syscontrol-mcp/0.1"})
        with urllib.request.urlopen(req, timeout=8) as r:
            geo_data = json.loads(r.read().decode())
        if not geo_data:
            return {"error": f"Location '{location}' not found. Try a different city name."}
        lat     = float(geo_data[0]["lat"])
        lon     = float(geo_data[0]["lon"])
        display = geo_data[0].get("display_name", location)
        parts   = [p.strip() for p in display.split(",")]
        city    = parts[0]
        country = parts[-1] if len(parts) > 1 else ""
        region  = parts[1]  if len(parts) > 2 else ""
        return lat, lon, city, region, country, "geocode"

    # Auto-detect from IP via ipinfo.io
    with urllib.request.urlopen("https://ipinfo.io/json", timeout=8) as r:
        ip_data = json.loads(r.read().decode())
    loc_str = ip_data.get("loc", "0,0")
    lat, lon = map(float, loc_str.split(","))
    return (
        lat, lon,
        ip_data.get("city", "Unknown"),
        ip_data.get("region", ""),
        ip_data.get("country", ""),
        "ip_geolocation",
    )


def _fetch_openmeteo(
    lat: float, lon: float, temp_unit: str, wind_unit: str, units: str,
) -> dict:
    """Fetch current conditions from Open-Meteo (free, no API key required)."""
    params = (
        f"latitude={lat}&longitude={lon}"
        f"&current=temperature_2m,apparent_temperature,relative_humidity_2m,"
        f"precipitation,weathercode,windspeed_10m,is_day"
        f"&temperature_unit={temp_unit}&wind_speed_unit={wind_unit}"
        f"&precipitation_unit={'inch' if units == 'imperial' else 'mm'}"
        f"&forecast_days=1"
    )
    with urllib.request.urlopen(
        f"https://api.open-meteo.com/v1/forecast?{params}", timeout=10
    ) as r:
        return json.loads(r.read().decode())


def get_weather(location: str = "", units: str = "imperial") -> dict:
    """Return current weather conditions for *location* (or the machine's IP city)."""
    units       = units if units in ("imperial", "metric") else "imperial"
    temp_unit   = "fahrenheit" if units == "imperial" else "celsius"
    wind_unit   = "mph"        if units == "imperial" else "kmh"
    temp_symbol = "°F"         if units == "imperial" else "°C"
    speed_label = "mph"        if units == "imperial" else "km/h"

    try:
        coords = _resolve_location_coords(location)
        if isinstance(coords, dict):
            return coords  # error dict from geocoding
        lat, lon, city_name, region, country, location_source = coords

        weather_data = _fetch_openmeteo(lat, lon, temp_unit, wind_unit, units)
        current      = weather_data["current"]
        temp         = current["temperature_2m"]
        feels_like   = current["apparent_temperature"]
        humidity     = current["relative_humidity_2m"]
        wind         = current["windspeed_10m"]
        precip       = current["precipitation"]
        code         = current["weathercode"]
        is_day       = bool(current["is_day"])

        temp_f    = temp if units == "imperial" else (temp * 9 / 5 + 32)
        wind_mph  = wind if units == "imperial" else wind * 0.621371
        condition = _WMO_DESCRIPTIONS.get(code, f"Weather code {code}")
        clothing  = _clothing_suggestions(temp_f, code, wind_mph, humidity)

        return {
            "location": {
                "city":        city_name,
                "region":      region,
                "country":     country,
                "coordinates": {"lat": round(lat, 4), "lon": round(lon, 4)},
                "source":      location_source,
            },
            "current": {
                "temperature":      {"value": round(temp, 1),       "unit": temp_symbol},
                "feels_like":       {"value": round(feels_like, 1), "unit": temp_symbol},
                "humidity_percent": humidity,
                "wind_speed":       {"value": round(wind, 1),       "unit": speed_label},
                "precipitation": {
                    "value": round(precip, 2),
                    "unit":  "in" if units == "imperial" else "mm",
                },
                "condition":      condition,
                "condition_code": code,
                "is_day":         is_day,
            },
            "clothing_suggestions": clothing,
        }
    except (TimeoutError, urllib.error.URLError, OSError) as exc:
        return {"error": f"Network error: {exc}. Check your internet connection."}
    except (KeyError, ValueError, json.JSONDecodeError) as exc:
        return {"error": f"Failed to parse weather data: {exc}"}


# ── App update checker ────────────────────────────────────────────────────────

def _check_brew(results: dict, lock: threading.Lock) -> None:
    """Populate *results* with outdated Homebrew formulae and casks."""
    if not shutil.which("brew"):
        with lock:
            results["errors"].append(
                "Homebrew not installed — install from https://brew.sh"
            )
        return
    try:
        proc = subprocess.run(
            ["brew", "outdated", "--json=v2"],
            capture_output=True, text=True, timeout=120,
            env={**os.environ, "HOMEBREW_NO_AUTO_UPDATE": "1"},
        )
        if proc.returncode in (0, 1) and proc.stdout.strip():
            data     = json.loads(proc.stdout)
            formulae = [
                {
                    "name":      f["name"],
                    "installed": (
                        f["installed_versions"][0]
                        if f.get("installed_versions") else "?"
                    ),
                    "available": f.get("current_version", "?"),
                }
                for f in data.get("formulae", [])
            ]
            casks = [
                {
                    "name":      c["name"],
                    "installed": c.get("installed_versions", ["?"])[0],
                    "available": c.get("current_version", "?"),
                }
                for c in data.get("casks", [])
            ]
            with lock:
                results["brew_formulae"] = formulae
                results["brew_casks"]    = casks
        elif proc.returncode not in (0, 1):
            with lock:
                results["errors"].append(f"brew error: {proc.stderr.strip()[:200]}")
    except subprocess.TimeoutExpired:
        with lock:
            results["errors"].append("brew outdated timed out (>120s)")
    except (json.JSONDecodeError, OSError) as exc:
        with lock:
            results["errors"].append(f"brew parse error: {exc}")


def _check_mas(results: dict, lock: threading.Lock) -> None:
    """Populate *results* with outdated Mac App Store apps via ``mas outdated``."""
    if not shutil.which("mas"):
        with lock:
            results["errors"].append(
                "mas not installed — install with 'brew install mas' to check App Store updates"
            )
        return
    try:
        proc = subprocess.run(
            ["mas", "outdated"], capture_output=True, text=True, timeout=60,
        )
        apps = []
        for line in proc.stdout.splitlines():
            m = re.match(r"(\d+)\s+(.+?)\s+\((.+?)\)", line.strip())
            if m:
                apps.append({
                    "app_id":            m.group(1),
                    "name":              m.group(2).strip(),
                    "available_version": m.group(3),
                })
        with lock:
            results["mac_app_store"] = apps
    except subprocess.TimeoutExpired:
        with lock:
            results["errors"].append("mas outdated timed out (>60s)")
    except OSError as exc:
        with lock:
            results["errors"].append(f"mas error: {exc}")


def _check_softwareupdate(results: dict, lock: threading.Lock) -> None:
    """Populate *results* with pending macOS system updates via ``softwareupdate -l``."""
    if not shutil.which("softwareupdate"):
        return
    try:
        proc = subprocess.run(
            ["softwareupdate", "-l"], capture_output=True, text=True, timeout=60,
        )
        combined      = proc.stdout + proc.stderr
        current_label = None
        updates: list = []
        for line in combined.splitlines():
            stripped = line.strip()
            if stripped.startswith("* Label:"):
                current_label = stripped.split(":", 1)[1].strip()
            elif current_label and "Title:" in stripped:
                m = re.search(r"Title:\s*(.+?),\s*Version:\s*([\d.]+)", stripped)
                if m:
                    updates.append({
                        "label":   current_label,
                        "title":   m.group(1).strip(),
                        "version": m.group(2),
                    })
                current_label = None
        with lock:
            results["system_updates"] = updates
    except subprocess.TimeoutExpired:
        with lock:
            results["errors"].append("softwareupdate timed out (>60s)")
    except OSError as exc:
        with lock:
            results["errors"].append(f"softwareupdate error: {exc}")


def check_app_updates() -> dict:
    """Check for outdated Homebrew, Mac App Store, and macOS system updates."""
    if not IS_MACOS:
        return {"error": "check_app_updates is currently macOS-only."}

    results: dict = {
        "brew_formulae":  [],
        "brew_casks":     [],
        "mac_app_store":  [],
        "system_updates": [],
        "errors":         [],
        "summary":        "",
    }
    lock = threading.Lock()

    # Run all three checks concurrently via the shared executor.
    futures = [
        ("brew",          _METRICS_EXECUTOR.submit(_check_brew,           results, lock)),
        ("mas",           _METRICS_EXECUTOR.submit(_check_mas,            results, lock)),
        ("softwareupdate", _METRICS_EXECUTOR.submit(_check_softwareupdate, results, lock)),
    ]
    for label, f in futures:
        try:
            f.result(timeout=130)   # brew timeout is 120s; add a small buffer
        except Exception as exc:
            with lock:
                results["errors"].append(f"{label} worker failed: {exc}")

    total = (
        len(results["brew_formulae"]) + len(results["brew_casks"])
        + len(results["mac_app_store"]) + len(results["system_updates"])
    )
    if total == 0:
        results["summary"] = "All apps are up to date."
    else:
        parts: list[str] = []
        if results["brew_formulae"]:
            n = len(results["brew_formulae"])
            parts.append(f"{n} Homebrew formula{'e' if n != 1 else ''}")
        if results["brew_casks"]:
            n = len(results["brew_casks"])
            parts.append(f"{n} Homebrew cask{'s' if n != 1 else ''}")
        if results["mac_app_store"]:
            n = len(results["mac_app_store"])
            parts.append(f"{n} App Store app{'s' if n != 1 else ''}")
        if results["system_updates"]:
            n = len(results["system_updates"])
            parts.append(f"{n} system update{'s' if n != 1 else ''}")
        results["summary"] = (
            f"{total} update{'s' if total != 1 else ''} available: "
            + ", ".join(parts)
        )

    return results


# ── Homebrew management tools ─────────────────────────────────────────────────

_BREW_ENV = {**os.environ, "HOMEBREW_NO_AUTO_UPDATE": "1"}
_BREW_PKG_RE = re.compile(r"^[\w\-\./@]+$")


def brew_list(kind: str = "all") -> dict:
    """List installed Homebrew packages.

    Args:
        kind: ``"formulae"``, ``"casks"``, or ``"all"`` (default).

    Returns:
        ``{"formulae": [...], "casks": [...], "total": int}``
    """
    denied = _permission_check("allow_brew", "brew_list")
    if denied:
        return denied
    if not IS_MACOS:
        return {"error": "Homebrew is macOS only."}
    if not shutil.which("brew"):
        return {"error": "Homebrew is not installed. Get it at https://brew.sh"}
    formulae: list[str] = []
    casks:    list[str] = []
    try:
        if kind in ("formulae", "all"):
            proc = subprocess.run(
                ["brew", "list", "--formula"],
                capture_output=True, text=True, timeout=30, env=_BREW_ENV,
            )
            if proc.returncode == 0:
                formulae = [ln for ln in proc.stdout.splitlines() if ln.strip()]
        if kind in ("casks", "all"):
            proc = subprocess.run(
                ["brew", "list", "--cask"],
                capture_output=True, text=True, timeout=30, env=_BREW_ENV,
            )
            if proc.returncode == 0:
                casks = [ln for ln in proc.stdout.splitlines() if ln.strip()]
    except subprocess.TimeoutExpired:
        return {"error": "brew list timed out."}
    except Exception as e:
        return {"error": str(e)}
    return {"formulae": formulae, "casks": casks, "total": len(formulae) + len(casks)}


def brew_install(package: str) -> dict:
    """Install a Homebrew formula or cask.

    Args:
        package: Formula or cask name (e.g. ``"ripgrep"``).
                 Prefix with ``--cask`` for casks (e.g. ``"--cask firefox"``).

    Returns:
        ``{"status": "ok"|"already_installed", "package", "output"}``
    """
    denied = _permission_check("allow_brew", "brew_install")
    if denied:
        return denied
    if not IS_MACOS:
        return {"error": "Homebrew is macOS only."}
    if not package or not package.strip():
        return {"error": "package is required."}
    if not shutil.which("brew"):
        return {"error": "Homebrew is not installed. Get it at https://brew.sh"}
    pkg = package.strip()
    # Allow --cask prefix; validate the package name portion.
    name_part = pkg.removeprefix("--cask").strip() if pkg.startswith("--cask") else pkg
    if not _BREW_PKG_RE.match(name_part):
        return {"error": f"Invalid package name: {pkg!r}"}
    cmd = ["brew", "install"] + pkg.split()
    try:
        proc = subprocess.run(
            cmd, capture_output=True, text=True, timeout=300, env=_BREW_ENV,
        )
        combined = (proc.stdout + proc.stderr).strip()
        if proc.returncode == 0:
            return {"status": "ok", "package": pkg, "output": combined}
        if "already installed" in combined.lower():
            return {"status": "already_installed", "package": pkg, "output": combined}
        return {"error": combined or f"brew install {pkg} failed.", "package": pkg}
    except subprocess.TimeoutExpired:
        return {"error": f"brew install {pkg} timed out (>300s)."}
    except Exception as e:
        return {"error": str(e)}


def brew_upgrade(package: str = "") -> dict:
    """Upgrade one or all installed Homebrew packages.

    Args:
        package: Package name to upgrade, or empty to upgrade everything.

    Returns:
        ``{"status", "package": "all"|name, "output"}``
    """
    denied = _permission_check("allow_brew", "brew_upgrade")
    if denied:
        return denied
    if not IS_MACOS:
        return {"error": "Homebrew is macOS only."}
    if not shutil.which("brew"):
        return {"error": "Homebrew is not installed. Get it at https://brew.sh"}
    pkg = package.strip()
    if pkg and not _BREW_PKG_RE.match(pkg):
        return {"error": f"Invalid package name: {pkg!r}"}
    cmd = ["brew", "upgrade"] + ([pkg] if pkg else [])
    try:
        proc = subprocess.run(
            cmd, capture_output=True, text=True, timeout=300, env=_BREW_ENV,
        )
        combined = (proc.stdout + proc.stderr).strip()
        if proc.returncode == 0 or "already up-to-date" in combined.lower():
            return {"status": "ok", "package": pkg or "all", "output": combined}
        return {"error": combined or "brew upgrade failed.", "package": pkg or "all"}
    except subprocess.TimeoutExpired:
        return {"error": "brew upgrade timed out (>300s)."}
    except Exception as e:
        return {"error": str(e)}


def brew_uninstall(package: str) -> dict:
    """Uninstall a Homebrew formula or cask.

    Args:
        package: Formula or cask name to remove.

    Returns:
        ``{"status": "ok", "package", "output"}``
    """
    denied = _permission_check("allow_brew", "brew_uninstall")
    if denied:
        return denied
    if not IS_MACOS:
        return {"error": "Homebrew is macOS only."}
    if not package or not package.strip():
        return {"error": "package is required."}
    if not shutil.which("brew"):
        return {"error": "Homebrew is not installed. Get it at https://brew.sh"}
    pkg = package.strip()
    if not _BREW_PKG_RE.match(pkg):
        return {"error": f"Invalid package name: {pkg!r}"}
    try:
        proc = subprocess.run(
            ["brew", "uninstall", pkg],
            capture_output=True, text=True, timeout=60, env=_BREW_ENV,
        )
        combined = (proc.stdout + proc.stderr).strip()
        if proc.returncode == 0:
            return {"status": "ok", "package": pkg, "output": combined}
        return {"error": combined or f"brew uninstall {pkg} failed.", "package": pkg}
    except subprocess.TimeoutExpired:
        return {"error": f"brew uninstall {pkg} timed out."}
    except Exception as e:
        return {"error": str(e)}


# ── Package tracking ──────────────────────────────────────────────────────────

def _detect_carrier(tn: str) -> str:
    """Guess the shipping carrier from a tracking number pattern."""
    tn = re.sub(r"\s+", "", tn).upper()
    if tn.startswith("TBA"):
        return "amazon_logistics"
    if re.match(r"^1Z[A-Z0-9]{16}$", tn):
        return "ups"
    if re.match(r"^(94|93|92|91|90)\d{18,20}$", tn):
        return "usps"
    if re.match(r"^[A-Z]{2}\d{9}[A-Z]{2}$", tn):
        return "usps"
    if _FEDEX_RE.match(tn):
        return "fedex"
    if re.match(r"^\d{20,21}$", tn):
        return "usps"
    if re.match(r"^\d{10,11}$", tn):
        return "dhl"
    if re.match(r"^(JD|GM)\d{14,20}$", tn):
        return "dhl"
    return "unknown"


_17TRACK_STATUS_MAP = {
    10: "Not found / No information",
    20: "In transit",
    30: "Out for delivery",
    40: "Delivered",
    50: "Exception / Alert",
}

_17TRACK_CARRIER_NAMES = {
    100001: "UPS", 100002: "USPS", 100003: "FedEx",
    100004: "DHL", 100007: "Amazon Logistics", 100008: "DHL Express",
    100010: "Canada Post", 100012: "Australia Post", 100016: "La Poste",
}


def _parse_17track_response(
    resp: dict, tracking_number: str, carrier: str,
) -> dict:
    """Convert a 17track API response dict into a normalised tracking result."""
    if not resp.get("shipments"):
        return {
            "tracking_number": tracking_number,
            "detected_carrier": carrier,
            "status": "Not found",
            "note": "No tracking information found. The package may not yet be in the system.",
        }

    shipment = resp["shipments"][0]
    carrier_code     = shipment.get("carrier")
    reported_carrier = _17TRACK_CARRIER_NAMES.get(carrier_code, f"Carrier #{carrier_code}")

    track = shipment.get("track", {})
    w1    = track.get("w1", {})
    if not isinstance(w1, dict):
        return {
            "tracking_number": tracking_number,
            "detected_carrier": carrier,
            "status": "Unexpected response structure from 17track — their internal API may have changed.",
        }
    latest      = w1.get("z0", {})
    history_raw = w1.get("z1", [])
    status_code = latest.get("c", 10)
    status      = _17TRACK_STATUS_MAP.get(status_code, f"Status code {status_code}")

    latest_event = {
        "description": latest.get("b", latest.get("a", "")),
        "location":    latest.get("e", ""),
        "timestamp":   latest.get("d", ""),
    }
    history = [
        {"timestamp": e.get("a", ""), "description": e.get("b", ""), "location": e.get("c", "")}
        for e in history_raw[:10]
    ]
    return {
        "tracking_number":  tracking_number,
        "detected_carrier": carrier,
        "reported_carrier": reported_carrier,
        "status":           status,
        "status_code":      status_code,
        "latest_event":     latest_event,
        "history":          history,
    }


def track_package(tracking_number: str) -> dict:
    """Look up package tracking information via the 17track public API."""
    tn_clean = re.sub(r"\s+", "", tracking_number).upper()
    carrier  = _detect_carrier(tn_clean)

    if carrier == "amazon_logistics":
        return {
            "tracking_number": tracking_number,
            "detected_carrier": "Amazon Logistics",
            "status": "Cannot track via this tool",
            "note": (
                "Amazon Logistics (TBA tracking numbers) can only be tracked at "
                "amazon.com/orders. Standard carrier tracking is not available for these."
            ),
        }

    try:
        payload = json.dumps({"number": tn_clean}).encode()
        req = urllib.request.Request(
            "https://t.17track.net/restapi/track",
            data=payload,
            headers={
                "Content-Type": "application/json",
                "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36",
            },
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=12) as r:
            resp = json.loads(r.read().decode())
        return _parse_17track_response(resp, tracking_number, carrier)
    except (TimeoutError, urllib.error.URLError, OSError) as e:
        return {"tracking_number": tracking_number, "error": f"Network error: {str(e)}"}
    except (json.JSONDecodeError, KeyError, ValueError) as e:
        return {"tracking_number": tracking_number, "error": f"Failed to parse tracking response: {str(e)}"}


# ── New tool implementations ──────────────────────────────────────────────────

def find_large_files(path: str = "", n: int = 10) -> dict:
    """Find the top N largest files under path (default: home directory)."""
    root = pathlib.Path(path).expanduser().resolve() if path else pathlib.Path.home()
    if not root.exists():
        return {"error": f"Path '{path}' does not exist."}
    if not root.is_dir():
        return {"error": f"'{path}' is not a directory."}

    files: list[tuple[int, str]] = []
    scanned = 0
    for dirpath, dirnames, filenames in os.walk(root, onerror=None):
        # Prune noisy / hidden dirs in-place so os.walk skips them entirely.
        # Uses the module-level _FIND_SKIP_DIRS constant (not recreated per call).
        dirnames[:] = [
            d for d in dirnames
            if d not in _FIND_SKIP_DIRS and not d.startswith(".")
        ]
        for fname in filenames:
            fpath = pathlib.Path(dirpath) / fname
            try:
                sz = fpath.stat().st_size
                files.append((sz, str(fpath)))
                scanned += 1
            except OSError:
                continue

    top = heapq.nlargest(n, files)

    return {
        "search_root": str(root),
        "files_scanned": scanned,
        "top_files": [
            {"path": p, "size_mb": round(s / 1e6, 2), "size_bytes": s}
            for s, p in top
        ],
    }


def _ping_host(
    label: str, host: str, results: dict, lock: threading.Lock
) -> None:
    """Ping *host* (4 packets) and store latency info under ``results[label]``.

    Designed to run concurrently in a thread pool; *lock* serialises writes to
    the shared *results* dict.
    """
    try:
        cmd = (
            ["ping", "-n", "4", "-w", "2000", host]
            if IS_WIN
            else ["ping", "-c", "4", "-W", "2", host]
        )
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=15)
        out  = proc.stdout + proc.stderr

        avg_ms: float | None = None
        # macOS/Linux: min/avg/max/stddev = x/y/z/w ms
        m = re.search(r"min/avg/max(?:/(?:mdev|stddev))?\s*=\s*[\d.]+/([\d.]+)/", out)
        if m:
            avg_ms = float(m.group(1))
        # Windows: Average = Xms
        if avg_ms is None:
            m = re.search(r"Average\s*=\s*([\d.]+)\s*ms", out, re.I)
            if m:
                avg_ms = float(m.group(1))

        with lock:
            results[label] = {
                "host":           host,
                "reachable":      proc.returncode == 0,
                "avg_latency_ms": avg_ms,
            }
    except subprocess.TimeoutExpired:
        with lock:
            results[label] = {"host": host, "reachable": False, "error": "ping timed out"}
    except Exception as exc:
        with lock:
            results[label] = {"host": host, "reachable": False, "error": str(exc)}


def _diagnose_latency(
    gateway: str | None, results: dict
) -> list[str]:
    """Interpret ping results and return a list of human-readable diagnosis strings."""
    gw = results.get("gateway",        {})
    cf = results.get("cloudflare_dns", {})
    gd = results.get("google_dns",     {})

    if gateway and not gw.get("reachable"):
        return ["Cannot reach your local gateway — likely a router/Wi-Fi issue."]
    if not cf.get("reachable") and not gd.get("reachable"):
        return ["Gateway reachable but public DNS is not — likely an ISP or WAN issue."]

    lat = cf.get("avg_latency_ms") or gd.get("avg_latency_ms")
    if lat and lat > 100:
        return [f"High latency ({lat} ms) to public DNS — possible ISP congestion."]
    if lat and lat > 50:
        return [f"Moderate latency ({lat} ms) — network is functional but not ideal."]
    return ["Network connectivity looks normal."]


def network_latency_check() -> dict:
    """
    Pings the local gateway, Cloudflare (1.1.1.1), and Google DNS (8.8.8.8)
    CONCURRENTLY using threads, then diagnoses where latency is introduced.
    Async: YES — all pings run in parallel via threading.
    """
    gateway: str | None = None
    try:
        nr = subprocess.run(["netstat", "-nr"], capture_output=True, text=True, timeout=5)
        for line in nr.stdout.splitlines():
            parts = line.split()
            if parts and parts[0] in ("default", "0.0.0.0") and len(parts) >= 2:
                gateway = parts[1]
                break
    except Exception:
        pass

    targets: dict[str, str] = {}
    if gateway:
        targets["gateway"] = gateway
    targets["cloudflare_dns"] = "1.1.1.1"
    targets["google_dns"]     = "8.8.8.8"
    targets["cloudflare.com"] = "cloudflare.com"

    results: dict = {}
    lock = threading.Lock()

    futures = [
        _METRICS_EXECUTOR.submit(_ping_host, lbl, h, results, lock)
        for lbl, h in targets.items()
    ]
    for f in futures:
        with contextlib.suppress(Exception):
            f.result(timeout=20)

    return {"targets": results, "diagnosis": _diagnose_latency(gateway, results)}


def _attach_docker_stats(containers: list[dict]) -> None:
    """Fetch one-shot container stats and merge CPU/memory data into *containers* in-place."""
    if not containers:
        return
    stats = subprocess.run(
        ["docker", "stats", "--no-stream", "--format",
         "{{.Name}}\t{{.CPUPerc}}\t{{.MemUsage}}\t{{.MemPerc}}"],
        capture_output=True, text=True, timeout=20,
    )
    stat_map: dict[str, dict] = {}
    for line in stats.stdout.strip().splitlines():
        parts = line.split("\t")
        if len(parts) >= 4:
            stat_map[parts[0]] = {
                "cpu_percent":    parts[1],
                "memory_usage":   parts[2],
                "memory_percent": parts[3],
            }
    for c in containers:
        c.update(stat_map.get(c["name"], {}))


def get_docker_status() -> dict:
    """Return running Docker containers with CPU and memory stats."""
    if not shutil.which("docker"):
        return {"error": "Docker is not installed or not in PATH."}

    try:
        ping = subprocess.run(
            ["docker", "info", "--format", "{{.ServerVersion}}"],
            capture_output=True, text=True, timeout=10,
        )
        if ping.returncode != 0:
            return {"error": "Docker daemon is not running. Start Docker Desktop first."}
        server_version = ping.stdout.strip()
    except subprocess.TimeoutExpired:
        return {"error": "Docker daemon did not respond in time."}

    try:
        ps = subprocess.run(
            ["docker", "ps", "--format",
             "{{.ID}}\t{{.Names}}\t{{.Image}}\t{{.Status}}\t{{.Ports}}"],
            capture_output=True, text=True, timeout=10,
        )
        containers: list[dict] = []
        for line in ps.stdout.strip().splitlines():
            parts = line.split("\t")
            if len(parts) >= 4:
                containers.append({
                    "id":     parts[0],
                    "name":   parts[1],
                    "image":  parts[2],
                    "status": parts[3],
                    "ports":  parts[4] if len(parts) > 4 else "",
                })

        _attach_docker_stats(containers)

        all_ps = subprocess.run(
            ["docker", "ps", "-a", "--format", "{{.Status}}"],
            capture_output=True, text=True, timeout=10,
        )
        total = len(all_ps.stdout.strip().splitlines()) if all_ps.stdout.strip() else 0

        return {
            "docker_version":     server_version,
            "running_count":      len(containers),
            "total_containers":   total,
            "running_containers": containers,
        }
    except subprocess.TimeoutExpired:
        return {"error": "Docker command timed out."}
    except Exception as exc:
        return {"error": f"Failed to query Docker: {exc}"}


def _tmutil_status(result: dict, lock: threading.Lock) -> None:
    """Populate *result* with current Time Machine backup status (running, phase, progress)."""
    try:
        proc = subprocess.run(
            ["tmutil", "status"], capture_output=True, text=True, timeout=10
        )
        out  = proc.stdout
        data: dict = {"running": "Running = 1" in out}
        m = re.search(r'BackupPhase\s*=\s*"?([^";\n]+)"?', out)
        if m:
            data["phase"] = m.group(1).strip()
        m = re.search(r'Percent\s*=\s*([\d.]+)', out)
        if m:
            data["progress_percent"] = round(float(m.group(1)) * 100, 1)
        m = re.search(r'_raw_Percent\s*=\s*([\d.]+)', out)
        if m:
            data["progress_percent"] = round(float(m.group(1)) * 100, 1)
        with lock:
            result.update(data)
    except Exception as exc:
        with lock:
            result["status_error"] = str(exc)


def _tmutil_latest(result: dict, lock: threading.Lock) -> None:
    """Populate *result* with the path and human-readable age of the latest backup."""
    try:
        proc = subprocess.run(
            ["tmutil", "latestbackup"], capture_output=True, text=True, timeout=10
        )
        bp = proc.stdout.strip()
        if bp and "No backups" not in bp:
            with lock:
                result["last_backup_path"] = bp
            m = re.search(r"(\d{4}-\d{2}-\d{2}-\d{6})", bp)
            if m:
                try:
                    dt    = datetime.datetime.strptime(m.group(1), "%Y-%m-%d-%H%M%S")
                    delta = datetime.datetime.now() - dt
                    hours = int(delta.total_seconds() // 3600)
                    age   = f"{hours} hours ago" if hours < 48 else f"{delta.days} days ago"
                    with lock:
                        result["last_backup"]     = dt.isoformat()
                        result["last_backup_age"] = age
                except ValueError:
                    with lock:
                        result["last_backup"] = m.group(1)
        else:
            with lock:
                result["last_backup"] = "No backups found"
    except Exception as exc:
        with lock:
            result["last_backup_error"] = str(exc)


def _tmutil_dest(result: dict, lock: threading.Lock) -> None:
    """Populate *result* with Time Machine destination name and kind."""
    try:
        proc = subprocess.run(
            ["tmutil", "destinationinfo"], capture_output=True, text=True, timeout=10
        )
        m = re.search(r"Name\s*:\s*(.+)", proc.stdout)
        if m:
            with lock:
                result["destination"] = m.group(1).strip()
        m = re.search(r"Kind\s*:\s*(.+)", proc.stdout)
        if m:
            with lock:
                result["destination_kind"] = m.group(1).strip()
    except Exception:
        pass


def get_time_machine_status() -> dict:
    """
    Return macOS Time Machine backup status, last backup time, and destination.
    Async: YES — tmutil status, latestbackup, and destinationinfo run in parallel.
    """
    if not IS_MACOS:
        return {"error": "Time Machine is macOS-only."}
    if not shutil.which("tmutil"):
        return {"error": "tmutil not found."}

    result: dict = {}
    lock = threading.Lock()

    futures = [
        _METRICS_EXECUTOR.submit(_tmutil_status, result, lock),
        _METRICS_EXECUTOR.submit(_tmutil_latest, result, lock),
        _METRICS_EXECUTOR.submit(_tmutil_dest,   result, lock),
    ]
    for f in futures:
        with contextlib.suppress(Exception):
            f.result(timeout=15)

    return result


def _tail_macos_log(lines: int, filter_str: str) -> dict:
    """Tail the macOS unified log (last 5 minutes)."""
    cmd = ["log", "show", "--last", "5m", "--style", "compact"]
    if filter_str:
        # Restrict to a conservative character set so the value cannot break
        # out of the NSPredicate string literal (newlines, quotes, etc. would
        # alter the predicate's logical structure).
        if not re.fullmatch(r"[A-Za-z0-9 _\-.:/=]+", filter_str):
            return {
                "error": (
                    "filter_str contains characters not allowed in a log predicate. "
                    "Allowed: letters, digits, space, and _-.:/="
                )
            }
        cmd += ["--predicate", f'eventMessage CONTAINS[c] "{filter_str}"']
    try:
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        all_lines = [ln for ln in proc.stdout.splitlines() if ln.strip()]
        tail = all_lines[-lines:]
        return {
            "platform": "macOS",
            "source":   "unified system log (last 5 minutes)",
            "filter":   filter_str or None,
            "line_count": len(tail),
            "lines":    tail,
        }
    except subprocess.TimeoutExpired:
        return {"error": "Log command timed out — try reducing lines or adding a filter."}
    except Exception as exc:
        return {"error": f"Failed to read logs: {exc}"}


def _tail_linux_log(lines: int, filter_str: str) -> dict:
    """Tail Linux system logs via journalctl or /var/log/syslog fallback."""
    if shutil.which("journalctl"):
        cmd = ["journalctl", "-n", str(lines), "--no-pager", "-o", "short"]
        if filter_str:
            try:
                re.compile(filter_str)
            except re.error:
                return {"error": "Invalid regex in filter_str."}
            cmd += ["-g", filter_str]
        try:
            proc = subprocess.run(cmd, capture_output=True, text=True, timeout=15)
            log_lines = proc.stdout.splitlines()
            return {
                "platform": "Linux", "source": "journalctl",
                "filter": filter_str or None,
                "line_count": len(log_lines), "lines": log_lines,
            }
        except Exception as exc:
            return {"error": f"journalctl failed: {exc}"}
    syslog = pathlib.Path("/var/log/syslog")
    if syslog.exists():
        if syslog.stat().st_size > 50_000_000:  # 50 MB guard
            return {"error": "System log is too large to read safely (>50 MB)."}
        try:
            all_lines = syslog.read_text(errors="replace").splitlines()
            tail = [
                ln for ln in all_lines
                if not filter_str or filter_str.lower() in ln.lower()
            ][-lines:]
            return {
                "platform": "Linux", "source": "/var/log/syslog",
                "filter": filter_str or None,
                "line_count": len(tail), "lines": tail,
            }
        except PermissionError:
            return {"error": "Permission denied reading /var/log/syslog. Try sudo."}
    return {"error": "No supported log source found (journalctl or /var/log/syslog)."}


def tail_system_logs(lines: int = 50, filter_str: str = "") -> dict:
    """Tail recent system logs. macOS: unified log (last 5 min). Linux: journalctl."""
    lines = max(10, min(lines, 500))

    # Sanitise filter_str: cap length and strip control characters.
    filter_str = filter_str[:200]
    filter_str = re.sub(r"[\x00-\x1f]", "", filter_str)

    if IS_MACOS:
        return _tail_macos_log(lines, filter_str)
    if IS_LINUX:
        return _tail_linux_log(lines, filter_str)
    return {"error": f"tail_system_logs is not supported on {_SYSTEM}."}


# ── Browser / Web tools ──────────────────────────────────────────────────────

_BROWSER_PERMISSION_FILE = pathlib.Path.home() / ".syscontrol" / "browser_permission"

# Browsers the AppleScript helpers know how to talk to, in preference order.
# Arc, Brave, and Edge all use the Chrome AppleScript dictionary.
_CHROMIUM_APPS = ["Arc", "Google Chrome", "Brave Browser", "Microsoft Edge", "Chromium"]
_SAFARI_APP    = "Safari"


def _browser_permission_granted() -> bool:
    return _BROWSER_PERMISSION_FILE.exists()


def _browser_permission_required() -> dict:
    return {
        "error": "browser_access_not_granted",
        "message": (
            "Browser access has not been granted yet. "
            "Ask the user to confirm, then call grant_browser_access() to enable it."
        ),
    }


def _running_browser() -> str | None:
    """Return the name of the first recognised browser that is currently running."""
    if not IS_MACOS:
        return None
    # Single AppleScript call to get all running process names, then match
    # against known browsers — avoids spawning one subprocess per browser.
    try:
        r = subprocess.run(
            ["osascript", "-e",
             'tell application "System Events" to get name of every process'],
            capture_output=True, text=True, timeout=5,
        )
        if r.returncode != 0:
            return None
        running = {n.strip() for n in r.stdout.split(",")}
        for app in _CHROMIUM_APPS + [_SAFARI_APP]:
            if app in running:
                return app
    except Exception:
        pass
    return None


def _osa(script: str, timeout: int = 10) -> tuple[str, str, int]:
    """Run an AppleScript snippet and return (stdout, stderr, returncode)."""
    r = subprocess.run(["osascript", "-e", script], capture_output=True, text=True, timeout=timeout)
    return r.stdout.strip(), r.stderr.strip(), r.returncode


def _chromium_script(app: str, js_or_cmd: str) -> str:
    """Wrap a Chrome-protocol AppleScript command for the given app."""
    return f'tell application "{app}" to {js_or_cmd}'


def _safari_script(cmd: str) -> str:
    return f'tell application "Safari" to {cmd}'


# ─────────────────────────────────────────────────────────────────────────────

def grant_browser_access() -> dict:
    """
    Writes the browser permission flag so that browser control tools can run.
    ONLY call this after the user has explicitly said yes.
    """
    try:
        _BROWSER_PERMISSION_FILE.parent.mkdir(parents=True, exist_ok=True)
        _BROWSER_PERMISSION_FILE.write_text("granted")
        browser = _running_browser()
        return {
            "success": True,
            "message": "Browser access granted.",
            "detected_browser": browser or "none running — open a browser and try again",
        }
    except Exception as exc:
        return {"success": False, "error": str(exc)}


_RE_SCRIPT_STYLE = re.compile(r"<(script|style)[^>]*>.*?</\1>", re.DOTALL | re.IGNORECASE)
_RE_HTML_TAG      = re.compile(r"<[^>]+>")
_RE_WHITESPACE    = re.compile(r"\s+")


def _strip_html(html: str, max_chars: int) -> str:
    """Very fast HTML → plain-text: strip tags, collapse whitespace."""
    text = _RE_SCRIPT_STYLE.sub(" ", html)
    text = _RE_HTML_TAG.sub(" ", text)
    for entity, char in (("&amp;", "&"), ("&lt;", "<"), ("&gt;", ">"),
                          ("&quot;", '"'), ("&#39;", "'"), ("&nbsp;", " ")):
        text = text.replace(entity, char)
    text = _RE_WHITESPACE.sub(" ", text).strip()
    return text[:max_chars]


def web_fetch(url: str, max_chars: int = 8000) -> dict:
    """
    Fetch a web page and return plain-text content (no browser needed).
    HTML tags are stripped. Does NOT require browser permission.
    """
    if not url.startswith(("http://", "https://")):
        url = "https://" + url

    parsed = urllib.parse.urlparse(url)
    host = parsed.hostname or ""
    _BLOCKED_HOSTS = {"localhost", "127.0.0.1", "::1", "0.0.0.0", "169.254.169.254"}
    if host in _BLOCKED_HOSTS:
        return {"url": url, "error": "Access to private/internal URLs is blocked."}
    try:
        ip = ipaddress.ip_address(host)
        if ip.is_private or ip.is_loopback or ip.is_link_local:
            return {"url": url, "error": "Access to private/internal IP addresses is blocked."}
    except ValueError:
        pass

    max_chars = max(500, min(max_chars, 32000))
    try:
        req = urllib.request.Request(
            url,
            headers={
                "User-Agent": (
                    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/123.0 Safari/537.36"
                ),
                "Accept": "text/html,application/xhtml+xml,*/*;q=0.9",
                "Accept-Language": "en-US,en;q=0.9",
            },
        )
        with urllib.request.urlopen(req, timeout=15) as r:
            raw = r.read()
            charset = "utf-8"
            ctype = r.headers.get("Content-Type", "")
            m = re.search(r"charset=([^\s;]+)", ctype)
            if m:
                charset = m.group(1)
            html = raw.decode(charset, errors="replace")
        text = _strip_html(html, max_chars)
        return {
            "url": url,
            "status": r.status,   # type: ignore[possibly-undefined]
            "content_length": len(text),
            "text": text,
            "truncated": len(text) == max_chars,
        }
    except urllib.error.HTTPError as e:
        return {"url": url, "error": f"HTTP {e.code}: {e.reason}"}
    except (TimeoutError, urllib.error.URLError, OSError) as e:
        return {"url": url, "error": f"Network error: {e}"}


def web_search(query: str, num_results: int = 5) -> dict:
    """
    Search DuckDuckGo and return the top results (title, URL, snippet).
    No API key needed. Does NOT require browser permission.
    """
    num_results = max(1, min(num_results, 10))
    encoded = urllib.parse.quote_plus(query)
    url = f"https://html.duckduckgo.com/html/?q={encoded}"
    try:
        req = urllib.request.Request(
            url,
            headers={
                "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36",
                "Accept-Language": "en-US,en;q=0.9",
            },
        )
        with urllib.request.urlopen(req, timeout=12) as r:
            html = r.read().decode("utf-8", errors="replace")
    except Exception as e:
        return {"query": query, "error": f"Search request failed: {e}"}

    # Parse DuckDuckGo HTML results — structure is stable enough for parsing
    results = []
    # Each result block: <a class="result__a" href="...">Title</a>
    #                    <a class="result__snippet">Snippet</a>
    title_pattern   = re.compile(r'<a[^>]+class="result__a"[^>]*href="([^"]+)"[^>]*>(.*?)</a>', re.DOTALL)
    snippet_pattern = re.compile(r'<a[^>]+class="result__snippet"[^>]*>(.*?)</a>', re.DOTALL)

    titles   = title_pattern.findall(html)
    snippets = [re.sub(r"<[^>]+>", "", s).strip() for s in snippet_pattern.findall(html)]

    for i, (href, title_raw) in enumerate(titles[:num_results]):
        title = re.sub(r"<[^>]+>", "", title_raw).strip()
        # DDG wraps URLs — extract the actual destination
        m = re.search(r"uddg=([^&]+)", href)
        real_url = urllib.parse.unquote(m.group(1)) if m else href
        results.append({
            "rank":    i + 1,
            "title":   title,
            "url":     real_url,
            "snippet": snippets[i] if i < len(snippets) else "",
        })

    return {
        "query":       query,
        "result_count": len(results),
        "results":     results,
        **({
            "warning": "No results parsed — DuckDuckGo HTML structure may have changed."
        } if not results else {}),
    }


def browser_open_url(url: str) -> dict:
    """Open a URL in the user's default browser. Requires browser permission."""
    if not _browser_permission_granted():
        return _browser_permission_required()
    if not url.startswith(("http://", "https://", "file://")):
        url = "https://" + url
    if IS_MACOS:
        try:
            subprocess.run(["open", url], check=True, timeout=10)
            return {"success": True, "url": url, "action": "opened in default browser"}
        except Exception as e:
            return {"success": False, "error": str(e)}
    # Linux / Windows fallback
    try:
        webbrowser.open(url)
        return {"success": True, "url": url, "action": "opened in default browser"}
    except Exception as e:
        return {"success": False, "error": str(e)}


# Compiled once at module load — used by browser_navigate to validate URLs
# before embedding them in AppleScript string literals.
_SAFE_URL_RE = re.compile(r'^[\x20-\x7E]+$')


def browser_navigate(url: str) -> dict:
    """
    Navigate the currently active browser tab to a URL via AppleScript (macOS).
    Requires browser permission.
    """
    if not _browser_permission_granted():
        return _browser_permission_required()
    if not IS_MACOS:
        return browser_open_url(url)   # graceful fallback on non-macOS

    # Normalise scheme
    if not url.startswith(("http://", "https://", "file://")):
        url = "https://" + url

    # Reject URLs with characters that could break out of the AppleScript string
    # literal. _SAFE_URL_RE is a module-level constant compiled once at import.
    if not _SAFE_URL_RE.match(url):
        return {"success": False, "error": "URL contains non-printable or non-ASCII characters."}
    if any(c in url for c in ('"', "'", '`', '\\', '\r', '\n')):
        return {"success": False, "error": "URL contains characters that are not safe for AppleScript."}

    browser = _running_browser()
    if not browser:
        # No known browser running — just open the URL
        return browser_open_url(url)

    try:
        if browser == _SAFARI_APP:
            script = _safari_script(f'set URL of current tab of front window to "{url}"')
        else:
            script = _chromium_script(browser, f'set URL of active tab of front window to "{url}"')
        stdout, stderr, rc = _osa(script)
        if rc != 0 and stderr:
            # Fallback: just open it
            return browser_open_url(url)
        return {"success": True, "url": url, "browser": browser, "action": "navigated"}
    except subprocess.TimeoutExpired:
        return {"success": False, "error": "AppleScript timed out — browser may be busy"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def browser_get_page() -> dict:
    """
    Return the URL, title, and visible text of the current active browser tab
    via AppleScript (macOS only). Requires browser permission.
    """
    if not _browser_permission_granted():
        return _browser_permission_required()
    if not IS_MACOS:
        return {"error": "browser_get_page requires macOS (uses AppleScript)."}

    browser = _running_browser()
    if not browser:
        return {
            "error": "No supported browser is running.",
            "supported": _CHROMIUM_APPS + [_SAFARI_APP],
        }

    try:
        if browser == _SAFARI_APP:
            url_out, _, rc1 = _osa(_safari_script("URL of current tab of front window"))
            title_out, _, rc2 = _osa(_safari_script("name of current tab of front window"))
            # Get visible text via JavaScript
            js_script = _safari_script(
                'do JavaScript "document.body ? document.body.innerText.substring(0,12000) : \'\'" '
                'in current tab of front window'
            )
            text_out, _, _ = _osa(js_script)
        else:
            url_out,   _, rc1 = _osa(_chromium_script(browser, "URL of active tab of front window"))
            title_out, _, rc2 = _osa(_chromium_script(browser, "title of active tab of front window"))
            js_script = _chromium_script(
                browser,
                'execute active tab of front window javascript '
                '"document.body ? document.body.innerText.substring(0,12000) : \'\'"'
            )
            text_out, _, _ = _osa(js_script)

        if rc1 != 0 or rc2 != 0:
            return {
                "error": "Could not read browser tab — make sure a window is open and focused.",
                "browser": browser,
            }

        # Strip excessive whitespace from innerText
        clean_text = re.sub(r"\n{3,}", "\n\n", text_out).strip()

        return {
            "browser":  browser,
            "url":      url_out,
            "title":    title_out,
            "text":     clean_text,
            "text_length": len(clean_text),
        }

    except subprocess.TimeoutExpired:
        return {"error": "AppleScript timed out — browser may be unresponsive.", "browser": browser}
    except Exception as e:
        return {"error": f"Failed to read browser page: {e}", "browser": browser}


# ── iMessage tools ───────────────────────────────────────────────────────────

# Recipient must be an E.164 phone number or a simple email address.
_IMESSAGE_RECIPIENT_RE = re.compile(
    r"^\+?[0-9]{7,15}$"          # phone: optional +, 7-15 digits
    r"|^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}$"  # email
)


def _escape_applescript(s: str) -> str:
    """Escape a string for safe embedding in an AppleScript double-quoted literal.

    Replaces backslashes first, then double quotes, then converts CR/LF/TAB into
    AppleScript escape sequences so they cannot terminate the surrounding literal.
    Strips remaining control characters.
    """
    s = (
        s.replace("\\", "\\\\")
        .replace('"', '\\"')
        .replace("\r", "\\r")
        .replace("\n", "\\n")
        .replace("\t", "\\t")
    )
    s = re.sub(r"[\x00-\x1f]", "", s)
    return s


def _imessage_fallback(safe_recipient: str, safe_message: str, first_stderr: str) -> dict | None:
    """Try sending an iMessage without specifying the service (SMS relay compatible).

    Returns an error dict on failure, or ``None`` on success.
    """
    script = (
        f'tell application "Messages"\n'
        f'  send "{safe_message}" to buddy "{safe_recipient}"\n'
        f'end tell'
    )
    proc = subprocess.run(
        ["osascript", "-e", script],
        capture_output=True, text=True, timeout=15,
    )
    if proc.returncode != 0:
        return {
            "error": proc.stderr.strip() or first_stderr,
            "hint": (
                "Make sure Messages.app is signed in and you have granted "
                "Automation permission to Terminal/iTerm in System Settings → "
                "Privacy & Security → Automation."
            ),
        }
    return None


def send_imessage(recipient: str, message: str) -> dict:
    """Send an iMessage or SMS via macOS Messages.app using AppleScript."""
    denied = _permission_check("allow_messaging", "send_imessage")
    if denied:
        return denied
    if not IS_MACOS:
        return {"error": "send_imessage requires macOS."}
    if not recipient or not message:
        return {"error": "recipient and message are required."}

    # Validate recipient format to prevent AppleScript injection.
    recipient_clean = recipient.strip()
    if not _IMESSAGE_RECIPIENT_RE.match(recipient_clean):
        return {
            "error": (
                f"Invalid recipient format: {recipient!r}. "
                "Must be an E.164 phone number (e.g. +14155551234) or email address."
            ),
        }

    safe_recipient = _escape_applescript(recipient_clean)
    safe_message = _escape_applescript(message)

    script = (
        f'tell application "Messages"\n'
        f'  set targetService to 1st service whose service type = iMessage\n'
        f'  set targetBuddy to buddy "{safe_recipient}" of targetService\n'
        f'  send "{safe_message}" to targetBuddy\n'
        f'end tell'
    )
    try:
        proc = subprocess.run(
            ["osascript", "-e", script],
            capture_output=True, text=True, timeout=15,
        )
        if proc.returncode != 0:
            err = _imessage_fallback(safe_recipient, safe_message, proc.stderr.strip())
            if err is not None:
                return err
        return {"status": "sent", "recipient": recipient_clean, "message": message}
    except subprocess.TimeoutExpired:
        return {"error": "AppleScript timed out sending iMessage."}
    except Exception as e:
        return {"error": str(e)}


def _query_imessage_db(db_path: "pathlib.Path", contact_q: str, limit: int) -> list[dict]:
    """Query chat.db read-only and return a list of message dicts (newest first)."""
    import sqlite3 as _sqlite3
    conn = _sqlite3.connect(f"file:{db_path}?mode=ro", uri=True, timeout=5)
    conn.row_factory = _sqlite3.Row
    cur = conn.cursor()
    cur.execute(
        """
        SELECT
            m.text,
            m.is_from_me,
            datetime(m.date / 1000000000 + strftime('%s','2001-01-01'), 'unixepoch', 'localtime') AS sent_at,
            h.id AS handle
        FROM message m
        JOIN chat_message_join cmj ON cmj.message_id = m.rowid
        JOIN chat c ON c.rowid = cmj.chat_id
        JOIN chat_handle_join chj ON chj.chat_id = c.rowid
        JOIN handle h ON h.rowid = chj.handle_id
        WHERE h.id LIKE ?
          AND m.text IS NOT NULL AND m.text != ''
        ORDER BY m.date DESC
        LIMIT ?
        """,
        (contact_q, limit),
    )
    rows = cur.fetchall()
    conn.close()
    return [
        {"from": "me" if r["is_from_me"] else r["handle"], "text": r["text"], "sent_at": r["sent_at"]}
        for r in rows
    ]


def get_imessage_history(contact: str, limit: int = 20) -> dict:
    """Return recent iMessage/SMS messages for a contact from chat.db."""
    denied = _permission_check("allow_message_history", "get_imessage_history")
    if denied:
        return denied
    if not IS_MACOS:
        return {"error": "get_imessage_history requires macOS."}

    db_path = pathlib.Path.home() / "Library" / "Messages" / "chat.db"
    if not db_path.exists():
        return {"error": f"chat.db not found at {db_path}. Full Disk Access may be required."}

    limit = max(1, min(limit, 200))
    try:
        messages = _query_imessage_db(db_path, f"%{contact}%", limit)
        return {
            "contact_filter": contact,
            "count": len(messages),
            "messages": list(reversed(messages)),  # chronological order
        }
    except Exception as e:
        return {
            "error": str(e),
            "hint": "Full Disk Access for Terminal is required in System Settings → Privacy & Security → Full Disk Access.",
        }


# ── Email tools (Mail.app) ────────────────────────────────────────────────────

_EMAIL_RE = re.compile(r"^[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}$")


def _parse_mail_records(raw: str) -> list[dict]:
    """Split pipe-delimited Mail.app AppleScript output into message dicts."""
    messages: list[dict] = []
    for record in raw.split("||"):
        record = record.strip()
        if not record:
            continue
        parts = record.split("|")
        messages.append({
            "subject": parts[0] if len(parts) > 0 else "",
            "sender":  parts[1] if len(parts) > 1 else "",
            "date":    parts[2] if len(parts) > 2 else "",
            "preview": parts[3] if len(parts) > 3 else "",
            "is_read": (parts[4] if len(parts) > 4 else "false") == "true",
        })
    return messages


def read_emails(folder: str = "INBOX", count: int = 10) -> dict:
    """Return the most recent messages from a Mail.app mailbox.

    Args:
        folder: Mailbox name (default: ``"INBOX"``).
        count:  Maximum messages to return (1–50, default 10).

    Returns:
        ``{"folder", "count", "messages"}`` — each message has ``subject``,
        ``sender``, ``date``, ``preview`` (first 200 chars), and ``is_read``.
    """
    denied = _permission_check("allow_email", "read_emails")
    if denied:
        return denied
    if not IS_MACOS:
        return {"error": "read_emails requires macOS (Mail.app)."}
    try:
        count = max(1, min(int(count), 50))
    except (TypeError, ValueError):
        count = 10
    safe_folder = _escape_applescript(folder)
    script = f"""
set output to ""
tell application "Mail"
    set msgList to {{}}
    repeat with acct in every account
        try
            set mbox to mailbox "{safe_folder}" of acct
            set msgList to msgList & (messages of mbox)
        end try
    end repeat
    set total to count of msgList
    if total = 0 then return ""
    set lim to {count}
    if total < lim then set lim to total
    repeat with i from 1 to lim
        set msg to item i of msgList
        set preview to content of msg
        if (count of characters of preview) > 200 then
            set preview to (characters 1 through 200 of preview as string) & "..."
        end if
        set output to output & (subject of msg) & "|" & (sender of msg) & "|" & ((date received of msg) as string) & "|" & preview & "|" & ((read status of msg) as string) & "||"
    end repeat
end tell
return output
"""
    try:
        proc = subprocess.run(
            ["osascript", "-e", script],
            capture_output=True, text=True, timeout=30,
        )
        if proc.returncode != 0:
            return {
                "error": proc.stderr.strip() or "Mail access denied.",
                "hint": "Grant Mail access in System Settings → Privacy & Security → Automation.",
            }
        messages = _parse_mail_records(proc.stdout.strip())
        return {"folder": folder, "count": len(messages), "messages": messages}
    except subprocess.TimeoutExpired:
        return {"error": "Mail.app timed out loading messages."}
    except Exception as e:
        return {"error": str(e)}


def send_email(to: str, subject: str, body: str, cc: str = "") -> dict:
    """Send an email via Mail.app.

    Args:
        to:      Recipient email address.
        subject: Email subject line.
        body:    Plain-text message body.
        cc:      Optional CC email address.

    Returns:
        ``{"status": "sent", "to", "subject"}``
    """
    denied = _permission_check("allow_email", "send_email")
    if denied:
        return denied
    if not IS_MACOS:
        return {"error": "send_email requires macOS (Mail.app)."}
    if not to or not subject or not body:
        return {"error": "to, subject, and body are required."}
    to_clean = to.strip()
    if not _EMAIL_RE.match(to_clean):
        return {"error": f"Invalid recipient email address: {to!r}"}
    if cc and not _EMAIL_RE.match(cc.strip()):
        return {"error": f"Invalid CC email address: {cc!r}"}
    safe_to      = _escape_applescript(to_clean)
    safe_subject = _escape_applescript(subject)
    safe_body    = _escape_applescript(body)
    cc_line = ""
    if cc:
        safe_cc = _escape_applescript(cc.strip())
        cc_line = f'        make new cc recipient with properties {{address:"{safe_cc}"}}\n'
    script = (
        f'tell application "Mail"\n'
        f'    set newMsg to make new outgoing message with properties '
        f'{{subject:"{safe_subject}", content:"{safe_body}", visible:false}}\n'
        f'    tell newMsg\n'
        f'        make new to recipient with properties {{address:"{safe_to}"}}\n'
        f'{cc_line}'
        f'        send\n'
        f'    end tell\n'
        f'end tell\n'
    )
    try:
        proc = subprocess.run(
            ["osascript", "-e", script],
            capture_output=True, text=True, timeout=20,
        )
        if proc.returncode != 0:
            return {
                "error": proc.stderr.strip() or "Failed to send email.",
                "hint": "Ensure Mail.app is open and configured with a sending account.",
            }
        return {"status": "sent", "to": to_clean, "subject": subject}
    except subprocess.TimeoutExpired:
        return {"error": "Mail.app timed out while sending."}
    except Exception as e:
        return {"error": str(e)}


def search_emails(query: str, count: int = 20) -> dict:
    """Search emails in Mail.app across all accounts and mailboxes.

    Args:
        query: Search term (subject, sender, or body text).
        count: Maximum results (1–50, default 20).

    Returns:
        ``{"query", "count", "messages"}``
    """
    denied = _permission_check("allow_email", "search_emails")
    if denied:
        return denied
    if not IS_MACOS:
        return {"error": "search_emails requires macOS (Mail.app)."}
    if not query:
        return {"error": "query is required."}
    try:
        count = max(1, min(int(count), 50))
    except (TypeError, ValueError):
        count = 20
    safe_query = _escape_applescript(query)
    script = f"""
set output to ""
tell application "Mail"
    set found to search every mailbox for "{safe_query}"
    set total to count of found
    if total = 0 then return ""
    set lim to {count}
    if total < lim then set lim to total
    repeat with i from 1 to lim
        set msg to item i of found
        set preview to content of msg
        if (count of characters of preview) > 200 then
            set preview to (characters 1 through 200 of preview as string) & "..."
        end if
        set output to output & (subject of msg) & "|" & (sender of msg) & "|" & ((date received of msg) as string) & "|" & preview & "|" & ((read status of msg) as string) & "||"
    end repeat
end tell
return output
"""
    try:
        proc = subprocess.run(
            ["osascript", "-e", script],
            capture_output=True, text=True, timeout=30,
        )
        if proc.returncode != 0:
            return {
                "error": proc.stderr.strip() or "Mail search failed.",
                "hint": "Grant Mail access in System Settings → Privacy & Security → Automation.",
            }
        messages = _parse_mail_records(proc.stdout.strip())
        return {"query": query, "count": len(messages), "messages": messages}
    except subprocess.TimeoutExpired:
        return {"error": "Mail.app search timed out."}
    except Exception as e:
        return {"error": str(e)}


# ── Clipboard tools ───────────────────────────────────────────────────────────

def get_clipboard() -> dict:
    """Return the current contents of the system clipboard."""
    if not IS_MACOS:
        return {"error": "get_clipboard is currently macOS only (uses pbpaste)."}
    try:
        result = subprocess.run(["pbpaste"], capture_output=True, text=True, timeout=5)
        text = result.stdout
        return {
            "text": text,
            "length": len(text),
            "has_content": bool(text.strip()),
        }
    except Exception as e:
        return {"error": str(e)}


def set_clipboard(text: str) -> dict:
    """Write text to the system clipboard."""
    if not IS_MACOS:
        return {"error": "set_clipboard is currently macOS only (uses pbcopy)."}
    try:
        subprocess.run(
            ["pbcopy"],
            input=text, text=True, timeout=5, check=True,
        )
        return {"status": "ok", "length": len(text)}
    except Exception as e:
        return {"error": str(e)}


# ── Screenshot tool ───────────────────────────────────────────────────────────

def take_screenshot(path: str = "") -> tuple:
    """
    Capture the entire screen. Always returns a 2-tuple (metadata_dict, base64_png_string).
    On error, returns ({"error": ...}, "").
    Optionally saves the image to `path` if provided.
    """
    denied = _permission_check("allow_screenshot", "take_screenshot")
    if denied:
        return denied, ""
    import tempfile as _tempfile

    if not IS_MACOS:
        return {"error": "take_screenshot requires macOS (uses screencapture)."}, ""

    with _tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
        tmp_path = tmp.name

    try:
        # -x = no sound, -C = capture cursor
        proc = subprocess.run(
            ["screencapture", "-x", tmp_path],
            capture_output=True, timeout=10,
        )
        if proc.returncode != 0:
            stderr = proc.stderr.decode(errors="replace").strip()
            return {"error": f"screencapture failed: {stderr or 'unknown error'}"}, ""

        img_file = pathlib.Path(tmp_path)
        if not img_file.exists() or img_file.stat().st_size == 0:
            return {"error": "screencapture produced no output (screen may not be accessible)."}, ""

        img_bytes = img_file.read_bytes()
        img_b64 = base64.b64encode(img_bytes).decode()

        saved_to = None
        if path:
            dest = pathlib.Path(path).expanduser().resolve()
            dest.parent.mkdir(parents=True, exist_ok=True)
            dest.write_bytes(img_bytes)
            saved_to = str(dest)

        meta = {
            "size_bytes": len(img_bytes),
            "saved_to": saved_to,
        }
        return meta, img_b64
    except subprocess.TimeoutExpired:
        return {"error": "screencapture timed out."}, ""
    except Exception as e:
        return {"error": str(e)}, ""
    finally:
        with contextlib.suppress(Exception):
            pathlib.Path(tmp_path).unlink(missing_ok=True)


# ── App control tools ─────────────────────────────────────────────────────────

def open_app(name: str) -> dict:
    """Open an application by name using macOS `open -a`."""
    if not IS_MACOS:
        return {"error": "open_app requires macOS."}
    if not name:
        return {"error": "app name is required."}
    try:
        proc = subprocess.run(
            ["open", "-a", name],
            capture_output=True, text=True, timeout=10,
        )
        if proc.returncode != 0:
            return {"error": proc.stderr.strip() or f"Could not open '{name}'."}
        return {"status": "ok", "app": name}
    except Exception as e:
        return {"error": str(e)}


def quit_app(name: str, force: bool = False) -> dict:
    """Gracefully quit an application by name using AppleScript."""
    if not IS_MACOS:
        return {"error": "quit_app requires macOS."}
    if not name:
        return {"error": "app name is required."}
    try:
        if force:
            # Force-quit via kill
            find_proc = subprocess.run(
                ["pgrep", "-ix", name],
                capture_output=True, text=True, timeout=5,
            )
            pids = find_proc.stdout.strip().splitlines()
            if not pids:
                return {"error": f"No process found matching '{name}'."}
            failed = []
            for pid in pids:
                result = subprocess.run(
                    ["kill", "-9", pid], capture_output=True, timeout=5,
                )
                if result.returncode != 0:
                    failed.append(pid)
            if failed:
                return {
                    "status": "partial-failure",
                    "app": name,
                    "pids": pids,
                    "failed_pids": failed,
                }
            return {"status": "force-killed", "app": name, "pids": pids}
        else:
            script = f'tell application "{_escape_applescript(name)}" to quit'
            proc = subprocess.run(
                ["osascript", "-e", script],
                capture_output=True, text=True, timeout=10,
            )
            if proc.returncode != 0:
                return {"error": proc.stderr.strip() or f"Could not quit '{name}'."}
            return {"status": "quit", "app": name}
    except Exception as e:
        return {"error": str(e)}


# ── Volume tools ──────────────────────────────────────────────────────────────

def get_volume() -> dict:
    """Return the current output volume and mute state."""
    if not IS_MACOS:
        return {"error": "get_volume requires macOS."}
    try:
        proc = subprocess.run(
            ["osascript", "-e", "get volume settings"],
            capture_output=True, text=True, timeout=5,
        )
        if proc.returncode != 0:
            return {"error": proc.stderr.strip()}
        # Output format: "output volume:75, input volume:54, alert volume:100, output muted:false"
        raw = proc.stdout.strip()
        result = {}
        for part in raw.split(","):
            part = part.strip()
            if ":" in part:
                k, v = part.split(":", 1)
                key = k.strip().replace(" ", "_")
                val_str = v.strip()
                if val_str.isdigit():
                    result[key] = int(val_str)
                elif val_str in ("true", "false"):
                    result[key] = val_str == "true"
                else:
                    result[key] = val_str
        return result
    except Exception as e:
        return {"error": str(e)}


def set_volume(level: int) -> dict:
    """Set the system output volume (0–100)."""
    if not IS_MACOS:
        return {"error": "set_volume requires macOS."}
    level = max(0, min(100, int(level)))
    try:
        proc = subprocess.run(
            ["osascript", "-e", f"set volume output volume {level}"],
            capture_output=True, text=True, timeout=5,
        )
        if proc.returncode != 0:
            return {"error": proc.stderr.strip()}
        return {"status": "ok", "output_volume": level}
    except Exception as e:
        return {"error": str(e)}


# ── Media control tools ───────────────────────────────────────────────────────

_MEDIA_APPS = ("Music", "Spotify")


def _get_media_state(app: str) -> dict | None:
    """Return playback state for *app* via AppleScript, or None if not running/available."""
    if app == "Music":
        script = """
tell application "Music"
    if it is not running then return "not_running"
    set ps to player state
    if ps is stopped then return "stopped"
    set t to current track
    return (ps as string) & "|" & (name of t) & "|" & (artist of t) & "|" & (album of t) & "|" & (player position as string) & "|" & (duration of t as string)
end tell
"""
    else:  # Spotify
        script = """
tell application "Spotify"
    if it is not running then return "not_running"
    set ps to player state
    if ps is stopped then return "stopped"
    set t to current track
    return (ps as string) & "|" & (name of t) & "|" & (artist of t) & "|" & (album of t) & "|" & (player position as string) & "|" & (duration of t as string)
end tell
"""
    try:
        proc = subprocess.run(
            ["osascript", "-e", script],
            capture_output=True, text=True, timeout=5,
        )
        raw = proc.stdout.strip()
        if not raw or raw == "not_running" or proc.returncode != 0:
            return None
        if raw == "stopped":
            return {"app": app, "state": "stopped"}
        parts = raw.split("|")
        return {
            "app":          app,
            "state":        parts[0] if parts else "unknown",
            "track":        parts[1] if len(parts) > 1 else "",
            "artist":       parts[2] if len(parts) > 2 else "",
            "album":        parts[3] if len(parts) > 3 else "",
            "position_sec": round(float(parts[4]), 1) if len(parts) > 4 else None,
            "duration_sec": round(float(parts[5]), 1) if len(parts) > 5 else None,
        }
    except Exception:
        return None


def get_now_playing() -> dict:
    """Return information about the currently playing media track.

    Checks Music.app then Spotify in order, returning whichever is playing.
    If neither is playing, returns the paused state if found, otherwise
    ``{"state": "nothing_playing"}``.

    Returns:
        A dict with ``app``, ``state``, ``track``, ``artist``, ``album``,
        ``position_sec``, and ``duration_sec``.
    """
    if not IS_MACOS:
        return {"error": "get_now_playing requires macOS."}
    for app in _MEDIA_APPS:
        info = _get_media_state(app)
        if info and info.get("state") == "playing":
            return info
    # Fall back to paused state if any app has a track loaded.
    for app in _MEDIA_APPS:
        info = _get_media_state(app)
        if info:
            return info
    return {"state": "nothing_playing"}


_MEDIA_ACTIONS: dict[str, dict[str, str]] = {
    "play":       {"Music": "play",           "Spotify": "play"},
    "pause":      {"Music": "pause",          "Spotify": "pause"},
    "play_pause": {"Music": "playpause",      "Spotify": "playpause"},
    "next":       {"Music": "next track",     "Spotify": "next track"},
    "previous":   {"Music": "previous track", "Spotify": "previous track"},
    "stop":       {"Music": "stop",           "Spotify": "pause"},
}


def media_control(action: str, app: str = "") -> dict:
    """Control media playback in Music.app or Spotify.

    Args:
        action: One of ``"play"``, ``"pause"``, ``"play_pause"``, ``"next"``,
                ``"previous"``, ``"stop"``.
        app:    ``"Music"`` or ``"Spotify"`` (default: auto-detect active player).

    Returns:
        ``{"status": "ok", "action": ..., "app": ...}``
    """
    if not IS_MACOS:
        return {"error": "media_control requires macOS."}
    action_key = action.lower().strip()
    if action_key not in _MEDIA_ACTIONS:
        return {
            "error": f"Unknown action {action!r}.",
            "valid_actions": list(_MEDIA_ACTIONS),
        }
    if app:
        target = app.strip().title()
        if target not in _MEDIA_APPS:
            return {"error": f"Unknown app {app!r}. Use 'Music' or 'Spotify'."}
    else:
        # Auto-detect: prefer whichever app is currently playing or paused.
        target = "Music"
        for candidate in _MEDIA_APPS:
            info = _get_media_state(candidate)
            if info and info.get("state") in ("playing", "paused"):
                target = candidate
                break
    cmd    = _MEDIA_ACTIONS[action_key][target]
    script = f'tell application "{target}" to {cmd}'
    try:
        proc = subprocess.run(
            ["osascript", "-e", script],
            capture_output=True, text=True, timeout=8,
        )
        if proc.returncode != 0:
            err = proc.stderr.strip()
            return {
                "error": err or f"Could not control {target}.",
                "hint":  f"Make sure {target} is installed and open.",
            }
        return {"status": "ok", "action": action_key, "app": target}
    except subprocess.TimeoutExpired:
        return {"error": f"{target} did not respond in time."}
    except Exception as e:
        return {"error": str(e)}


# ── Wi-Fi tool ────────────────────────────────────────────────────────────────

_AIRPORT_PATH = (
    "/System/Library/PrivateFrameworks/Apple80211.framework"
    "/Versions/Current/Resources/airport"
)

def _parse_rssi(sig_noise_str: str) -> int | None:
    """Parse a ``'-56 dBm / -95 dBm'``-style string from system_profiler to an int RSSI."""
    if not sig_noise_str:
        return None
    try:
        return int(sig_noise_str.split()[0])
    except (ValueError, IndexError):
        return None


def _scan_wifi_airport(airport_path: str) -> dict | None:
    """Scan visible Wi-Fi networks using the ``airport`` CLI (macOS ≤ 13).

    Returns a result dict on success, or ``None`` to signal fall-through to the
    system_profiler path.  Raises ``subprocess.TimeoutExpired`` on timeout so
    the caller can surface a user-facing error.
    """
    proc = subprocess.run(
        [airport_path, "-s"], capture_output=True, text=True, timeout=20,
    )
    if proc.returncode != 0:
        return None

    networks: list[dict] = []
    for line in proc.stdout.splitlines()[1:]:
        if not line.strip():
            continue
        try:
            ssid     = line[:33].strip()
            rest     = line[33:].split()
            bssid    = rest[0] if rest else ""
            rssi     = int(rest[1]) if len(rest) > 1 else None
            channel  = rest[2] if len(rest) > 2 else ""
            security = rest[6] if len(rest) > 6 else (rest[-1] if rest else "")
            networks.append({
                "ssid": ssid, "bssid": bssid,
                "rssi_dbm": rssi, "channel": channel,
                "security": security,
            })
        except (IndexError, ValueError):
            continue

    networks.sort(key=lambda n: n.get("rssi_dbm") or -999, reverse=True)
    return {"source": "airport", "networks": networks, "count": len(networks)}


def _scan_wifi_system_profiler() -> dict:
    """Scan Wi-Fi networks via ``system_profiler SPAirPortDataType`` (macOS 14+)."""
    proc = subprocess.run(
        ["system_profiler", "SPAirPortDataType", "-json"],
        capture_output=True, text=True, timeout=30,
    )
    if proc.returncode != 0:
        return {"error": proc.stderr.strip() or "system_profiler failed."}

    data     = json.loads(proc.stdout)
    sp_wifi  = data.get("SPAirPortDataType", [])
    networks: list[dict] = []

    for entry in sp_wifi:
        interfaces = entry.get("spairport_airport_interfaces", [])
        for iface in interfaces:
            cur = iface.get("spairport_current_network_information", {})
            if cur:
                networks.append({
                    "ssid":     cur.get("_name", ""),
                    "phy_mode": cur.get("spairport_network_phymode", ""),
                    "channel":  str(cur.get("spairport_network_channel", "")),
                    "security": cur.get("spairport_security_mode", ""),
                    "rssi_dbm": _parse_rssi(cur.get("spairport_signal_noise", "")),
                    "connected": True,
                })
            others = iface.get("spairport_airport_other_local_wireless_networks", [])
            if isinstance(others, list):
                for net in others:
                    sn = net.get("spairport_signal_noise", "")
                    networks.append({
                        "ssid":     net.get("_name", ""),
                        "phy_mode": net.get("spairport_network_phymode", ""),
                        "channel":  str(net.get("spairport_network_channel", "")),
                        "security": net.get("spairport_security_mode", ""),
                        "rssi_dbm": _parse_rssi(sn) if sn else None,
                        "connected": False,
                    })

    return {"source": "system_profiler", "networks": networks, "count": len(networks)}


def get_wifi_networks() -> dict:
    """
    Return information about nearby / available Wi-Fi networks.
    Uses the `airport` CLI when available (macOS ≤13), otherwise falls back to
    `system_profiler SPAirPortDataType` which works on macOS 14+.
    """
    if not IS_MACOS:
        return {"error": "get_wifi_networks requires macOS."}

    # ── Try airport (macOS ≤13) ──────────────────────────────────────────────
    airport = pathlib.Path(_AIRPORT_PATH)
    if airport.exists():
        try:
            result = _scan_wifi_airport(str(airport))
            if result is not None:
                return result
        except subprocess.TimeoutExpired:
            return {"error": "Wi-Fi scan timed out (20s). Enable Wi-Fi and try again."}
        except Exception:
            pass  # fall through to system_profiler

    # ── Fallback: system_profiler SPAirPortDataType (macOS 14+) ─────────────
    try:
        return _scan_wifi_system_profiler()
    except subprocess.TimeoutExpired:
        return {"error": "system_profiler timed out (30s)."}
    except Exception as exc:
        return {"error": str(exc)}


# ── File tools ────────────────────────────────────────────────────────────────

_MAX_READ_CHARS = 32_000

def read_file(path: str, max_chars: int = 16_000) -> dict:
    """Read a text file and return its contents."""
    denied = _permission_check("allow_file_read", "read_file")
    if denied:
        return denied
    if not path:
        return {"error": "path is required."}
    max_chars = max(1, min(max_chars, _MAX_READ_CHARS))
    try:
        p = pathlib.Path(path).expanduser().resolve()
        if not p.exists():
            return {"error": f"File not found: {p}"}
        if not p.is_file():
            return {"error": f"Not a file: {p}"}
        size = p.stat().st_size
        content = p.read_text(errors="replace")
        truncated = len(content) > max_chars
        return {
            "path": str(p),
            "size_bytes": size,
            "chars_read": min(len(content), max_chars),
            "truncated": truncated,
            "content": content[:max_chars],
        }
    except PermissionError:
        return {"error": f"Permission denied: {path}"}
    except Exception as e:
        return {"error": str(e)}


def write_file(path: str, content: str, overwrite: bool = True) -> dict:
    """Write text content to a file. Creates parent directories as needed."""
    denied = _permission_check("allow_file_write", "write_file")
    if denied:
        return denied
    if not path:
        return {"error": "path is required."}
    try:
        p = pathlib.Path(path).expanduser().resolve()
        if p.exists() and not overwrite:
            return {"error": f"File already exists: {p}. Pass overwrite=true to replace it."}
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(content)
        return {
            "status": "ok",
            "path": str(p),
            "bytes_written": len(content.encode()),
        }
    except PermissionError:
        return {"error": f"Permission denied: {path}"}
    except Exception as e:
        return {"error": str(e)}


# ── Code editing & navigation tools ──────────────────────────────────────────

_SKIP_DIRS: frozenset[str] = frozenset({
    ".git", "node_modules", ".venv", "__pycache__", ".mypy_cache", ".ruff_cache",
    ".tox", ".eggs", "dist", "build",
})

_MAX_GREP_FILES = 10_000
_MAX_GREP_RESULTS = 500
_MAX_GLOB_RESULTS = 500
_MAX_DIFF_CHARS = 16_000


def read_file_lines(path: str, offset: int = 1, limit: int = 200) -> dict:
    """Read a file with line numbers, supporting offset/limit for large files."""
    denied = _permission_check("allow_file_read", "read_file_lines")
    if denied:
        return denied
    if not path:
        return {"error": "path is required."}
    offset = max(1, offset)
    limit = max(1, min(limit, 2000))
    try:
        p = pathlib.Path(path).expanduser().resolve()
        if not p.exists():
            return {"error": f"File not found: {p}"}
        if not p.is_file():
            return {"error": f"Not a file: {p}"}
        lines: list[str] = []
        total_lines: int | None = None
        with p.open(errors="replace") as f:
            lineno = 0
            for lineno, text in enumerate(f, start=1):
                if lineno < offset:
                    continue
                if lineno >= offset + limit:
                    # Count remaining lines only if file is small enough.
                    if p.stat().st_size < 10_000_000:
                        remaining = sum(1 for _ in f)
                        total_lines = lineno + remaining
                    break
                lines.append(f"{lineno:>6}\t{text}")
            else:
                # Reached EOF — we know the total.
                total_lines = (lineno if lines or offset == 1 else 0)
        content = "".join(lines)
        return {
            "path": str(p),
            "offset": offset,
            "limit": limit,
            "lines_read": len(lines),
            "total_lines": total_lines,
            "content": content,
        }
    except Exception as e:
        return {"error": str(e)}


def glob_files(pattern: str, path: str = ".") -> dict:
    """Find files matching a glob pattern. Skips common non-project directories."""
    if not pattern:
        return {"error": "pattern is required."}
    try:
        base = pathlib.Path(path).expanduser().resolve()
        if not base.exists():
            return {"error": f"Path not found: {base}"}
        files: list[dict] = []
        truncated = False
        for p in sorted(base.glob(pattern)):
            if any(part in _SKIP_DIRS for part in p.relative_to(base).parts):
                continue
            if not p.is_file():
                continue
            if len(files) >= _MAX_GLOB_RESULTS:
                truncated = True
                break
            try:
                st = p.stat()
                files.append({
                    "path": str(p),
                    "size_bytes": st.st_size,
                    "modified": datetime.fromtimestamp(st.st_mtime).isoformat(timespec="seconds"),
                })
            except OSError:
                continue
        return {
            "pattern": pattern,
            "base_path": str(base),
            "count": len(files),
            "truncated": truncated,
            "files": files,
        }
    except Exception as e:
        return {"error": str(e)}


def _grep_single_file(
    p: pathlib.Path,
    regex: re.Pattern,  # type: ignore[type-arg]
    context_lines: int,
) -> list[dict]:
    """Search a single file for regex matches, returning results with context."""
    try:
        # Skip binary files.
        with p.open("rb") as bf:
            chunk = bf.read(8192)
            if b"\x00" in chunk:
                return []
        with p.open(errors="replace") as f:
            all_lines = f.readlines()
    except (OSError, PermissionError):
        return []
    results: list[dict] = []
    for i, line in enumerate(all_lines):
        if regex.search(line):
            before = [ln.rstrip("\n\r") for ln in all_lines[max(0, i - context_lines):i]]
            after = [ln.rstrip("\n\r") for ln in all_lines[i + 1:i + 1 + context_lines]]
            results.append({
                "file": str(p),
                "line": i + 1,
                "text": line.rstrip("\n\r"),
                "context_before": before if context_lines > 0 else [],
                "context_after": after if context_lines > 0 else [],
            })
    return results


def grep_files(
    pattern: str,
    path: str = ".",
    include: str = "",
    max_results: int = 50,
    context_lines: int = 0,
) -> dict:
    """Search file contents for a regex pattern, like grep/ripgrep."""
    if not pattern:
        return {"error": "pattern is required."}
    try:
        regex = re.compile(pattern)
    except re.error as e:
        return {"error": f"Invalid regex: {e}"}
    max_results = max(1, min(max_results, _MAX_GREP_RESULTS))
    context_lines = max(0, min(context_lines, 10))
    try:
        base = pathlib.Path(path).expanduser().resolve()
        if not base.exists():
            return {"error": f"Path not found: {base}"}
        # Single file mode.
        if base.is_file():
            hits = _grep_single_file(base, regex, context_lines)
            return {
                "pattern": pattern, "path": str(base), "include": include,
                "match_count": len(hits), "file_count": 1 if hits else 0,
                "truncated": False, "results": hits[:max_results],
            }
        glob_pattern = include if include else "*"
        results: list[dict] = []
        file_count = 0
        files_scanned = 0
        truncated = False
        for p in sorted(base.rglob(glob_pattern)):
            if files_scanned >= _MAX_GREP_FILES:
                truncated = True
                break
            if not p.is_file():
                continue
            if any(part in _SKIP_DIRS for part in p.relative_to(base).parts):
                continue
            files_scanned += 1
            hits = _grep_single_file(p, regex, context_lines)
            if hits:
                file_count += 1
                results.extend(hits)
                if len(results) >= max_results:
                    truncated = True
                    results = results[:max_results]
                    break
        return {
            "pattern": pattern, "path": str(base), "include": include,
            "match_count": len(results), "file_count": file_count,
            "truncated": truncated, "results": results,
        }
    except Exception as e:
        return {"error": str(e)}


def edit_file(
    path: str, old_string: str, new_string: str, replace_all: bool = False,
) -> dict:
    """Make targeted edits to a file using exact find-and-replace."""
    denied = _permission_check("allow_file_write", "edit_file")
    if denied:
        return denied
    if not path:
        return {"error": "path is required."}
    if not old_string:
        return {"error": "old_string is required (cannot be empty)."}
    if old_string == new_string:
        return {"error": "old_string and new_string are identical."}
    try:
        p = pathlib.Path(path).expanduser().resolve()
        if not p.exists():
            return {"error": f"File not found: {p}"}
        if not p.is_file():
            return {"error": f"Not a file: {p}"}
        content = p.read_text(errors="replace")
        count = content.count(old_string)
        if count == 0:
            return {"error": "old_string not found in file."}
        if count > 1 and not replace_all:
            return {
                "error": (
                    f"old_string found {count} times. Use replace_all=true "
                    "or provide a more specific string that matches exactly once."
                ),
            }
        if replace_all:
            new_content = content.replace(old_string, new_string)
            replacements = count
        else:
            new_content = content.replace(old_string, new_string, 1)
            replacements = 1
        p.write_text(new_content)
        return {"status": "ok", "path": str(p), "replacements": replacements}
    except PermissionError:
        return {"error": f"Permission denied: {path}"}
    except Exception as e:
        return {"error": str(e)}


def _parse_git_porcelain(output: str) -> dict[str, list[str]]:
    """Parse ``git status --porcelain=v1`` into staged, unstaged, and untracked lists."""
    staged: list[str] = []
    unstaged: list[str] = []
    untracked: list[str] = []
    for line in output.splitlines():
        if len(line) < 3:
            continue
        x, y, name = line[0], line[1], line[3:]
        if x == "?" and y == "?":
            untracked.append(name)
        else:
            if x not in (" ", "?"):
                staged.append(name)
            if y not in (" ", "?"):
                unstaged.append(name)
    return {"staged": staged, "unstaged": unstaged, "untracked": untracked}


def _run_git(args: list[str], cwd: str, timeout: int = 10) -> subprocess.CompletedProcess:  # type: ignore[type-arg]
    """Run a git command and return the CompletedProcess."""
    return subprocess.run(
        ["git"] + args, capture_output=True, text=True, timeout=timeout, cwd=cwd,
    )


def git_status(path: str = ".") -> dict:
    """Show git repository status: branch, files, and recent commits."""
    try:
        p = pathlib.Path(path).expanduser().resolve()
        cwd = str(p if p.is_dir() else p.parent)
        # Find repo root.
        r = _run_git(["rev-parse", "--show-toplevel"], cwd=cwd)
        if r.returncode != 0:
            return {"error": "Not a git repository (or git is not installed)."}
        repo_root = r.stdout.strip()
        # Branch name.
        r = _run_git(["branch", "--show-current"], cwd=repo_root)
        branch = r.stdout.strip() or "(detached HEAD)"
        # Porcelain status.
        r = _run_git(["status", "--porcelain=v1"], cwd=repo_root)
        file_lists = _parse_git_porcelain(r.stdout)
        clean = not any(file_lists.values())
        # Recent commits.
        r = _run_git(["log", "--oneline", "-5"], cwd=repo_root)
        commits = []
        for line in r.stdout.strip().splitlines():
            parts = line.split(" ", 1)
            if len(parts) == 2:
                commits.append({"hash": parts[0], "message": parts[1]})
        return {
            "repo_root": repo_root, "branch": branch, "clean": clean,
            **file_lists, "recent_commits": commits,
        }
    except FileNotFoundError:
        return {"error": "git is not installed or not in PATH."}
    except subprocess.TimeoutExpired:
        return {"error": "git command timed out."}
    except Exception as e:
        return {"error": str(e)}


def git_diff(path: str = ".", staged: bool = False) -> dict:
    """Show git diff output for unstaged or staged changes."""
    try:
        p = pathlib.Path(path).expanduser().resolve()
        cwd = str(p if p.is_dir() else p.parent)
        r = _run_git(["rev-parse", "--show-toplevel"], cwd=cwd)
        if r.returncode != 0:
            return {"error": "Not a git repository (or git is not installed)."}
        repo_root = r.stdout.strip()
        cmd = ["diff"]
        if staged:
            cmd.append("--cached")
        if p.is_file():
            cmd.append(str(p))
        r = _run_git(cmd, cwd=repo_root)
        diff_text = r.stdout
        truncated = len(diff_text) > _MAX_DIFF_CHARS
        if truncated:
            diff_text = diff_text[:_MAX_DIFF_CHARS]
        return {
            "repo_root": repo_root, "staged": staged, "path": str(p),
            "diff": diff_text, "truncated": truncated,
        }
    except FileNotFoundError:
        return {"error": "git is not installed or not in PATH."}
    except subprocess.TimeoutExpired:
        return {"error": "git command timed out."}
    except Exception as e:
        return {"error": str(e)}


# ── Spreadsheet tools ─────────────────────────────────────────────────────────

_MAX_SPREADSHEET_ROWS = 200
# Pre-flight size cap to avoid pathological openpyxl/csv parses on hostile inputs.
_MAX_SPREADSHEET_BYTES = 50 * 1024 * 1024  # 50 MB


def read_spreadsheet(
    path: str,
    sheet: str | None = None,
    cell_range: str | None = None,
    max_rows: int = 200,
) -> dict:
    """Read rows from a spreadsheet file (.xlsx or .csv)."""
    denied = _permission_check("allow_file_read", "read_spreadsheet")
    if denied:
        return denied
    if not path:
        return {"error": "path is required."}
    max_rows = max(1, min(max_rows, _MAX_SPREADSHEET_ROWS))
    try:
        p = pathlib.Path(path).expanduser().resolve()
        if not p.exists():
            return {"error": f"File not found: {p}"}
        if not p.is_file():
            return {"error": f"Not a file: {p}"}
        if p.stat().st_size > _MAX_SPREADSHEET_BYTES:
            return {
                "error": (
                    f"File exceeds {_MAX_SPREADSHEET_BYTES // (1024 * 1024)} MB "
                    "size limit for spreadsheet reads."
                )
            }

        suffix = p.suffix.lower()
        if suffix == ".csv":
            import csv as _csv
            with p.open(newline="", encoding="utf-8", errors="replace") as fh:
                reader = _csv.reader(fh)
                all_rows = [row for row in reader]
            rows = all_rows[:max_rows]
            headers = rows[0] if rows else []
            return {
                "path": str(p),
                "sheet": None,
                "headers": headers,
                "rows": rows,
                "row_count": len(rows),
                "truncated": len(all_rows) > max_rows,
            }

        if suffix != ".xlsx":
            return {"error": f"Unsupported file type '{suffix}'. Supported: .xlsx, .csv"}
        if not _HAS_OPENPYXL:
            return {"error": "openpyxl is not installed. Run: uv add openpyxl"}

        wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
        ws = wb[sheet] if sheet else wb.active
        if ws is None:
            return {"error": f"Sheet '{sheet}' not found. Available: {wb.sheetnames}"}

        if cell_range:
            cells = ws[cell_range]
            # ws[range] returns a tuple of tuples
            if not isinstance(cells, tuple):
                cells = ((cells,),)
            rows = [[c.value for c in row] for row in cells]
        else:
            rows = []
            for row in ws.iter_rows(values_only=True):
                rows.append(list(row))
                if len(rows) >= max_rows:
                    break

        wb.close()
        headers = rows[0] if rows else []
        return {
            "path": str(p),
            "sheet": ws.title,
            "headers": headers,
            "rows": rows,
            "row_count": len(rows),
        }
    except KeyError as e:
        return {"error": f"Sheet not found: {e}"}
    except PermissionError:
        return {"error": f"Permission denied: {path}"}
    except Exception as e:
        return {"error": str(e)}


def edit_spreadsheet(
    path: str,
    sheet: str | None = None,
    updates: list | None = None,
    append_rows: list | None = None,
    create_if_missing: bool = False,
) -> dict:
    """Write cells or append rows to a spreadsheet (.xlsx or .csv)."""
    denied = _permission_check("allow_file_write", "edit_spreadsheet")
    if denied:
        return denied
    if not path:
        return {"error": "path is required."}
    if not updates and not append_rows:
        return {"error": "Provide 'updates' (cell changes) or 'append_rows', or both."}

    try:
        p = pathlib.Path(path).expanduser().resolve()
        suffix = p.suffix.lower()

        if suffix == ".csv":
            import csv as _csv
            existing: list = []
            if p.exists():
                with p.open(newline="", encoding="utf-8", errors="replace") as fh:
                    existing = list(_csv.reader(fh))
            if updates:
                for upd in updates:
                    cell_ref = upd.get("cell", "")
                    if not cell_ref:
                        continue
                    col_letter = "".join(c for c in cell_ref if c.isalpha()).upper()
                    row_str = "".join(c for c in cell_ref if c.isdigit())
                    if not col_letter or not row_str:
                        continue
                    col_idx = sum(
                        (ord(ch) - ord("A") + 1) * (26 ** i)
                        for i, ch in enumerate(reversed(col_letter))
                    ) - 1
                    row_idx = int(row_str) - 1
                    if row_idx > 100_000:
                        return {"error": f"Row index {row_idx + 1} exceeds CSV safety limit (100 000)."}
                    while len(existing) <= row_idx:
                        existing.append([])
                    while len(existing[row_idx]) <= col_idx:
                        existing[row_idx].append("")
                    existing[row_idx][col_idx] = upd.get("value", "")
            cells_updated = len(updates) if updates else 0
            rows_appended = 0
            if append_rows:
                existing.extend(list(r) for r in append_rows)
                rows_appended = len(append_rows)
            p.parent.mkdir(parents=True, exist_ok=True)
            with p.open("w", newline="", encoding="utf-8") as fh:
                _csv.writer(fh).writerows(existing)
            return {"status": "ok", "path": str(p), "cells_updated": cells_updated,
                    "rows_appended": rows_appended}

        if suffix != ".xlsx":
            return {"error": f"Unsupported file type '{suffix}'. Supported: .xlsx, .csv"}
        if not _HAS_OPENPYXL:
            return {"error": "openpyxl is not installed. Run: uv add openpyxl"}

        if p.exists():
            wb = openpyxl.load_workbook(str(p))
        elif create_if_missing:
            wb = openpyxl.Workbook()
        else:
            return {"error": f"File not found: {p}. Pass create_if_missing=true to create it."}

        ws = wb[sheet] if sheet and sheet in wb.sheetnames else wb.active
        cells_updated = 0
        if updates:
            for upd in updates:
                cell_ref = upd.get("cell")
                if not cell_ref:
                    continue
                ws[cell_ref] = upd.get("value")
                cells_updated += 1

        rows_appended = 0
        if append_rows:
            for row in append_rows:
                ws.append(list(row))
                rows_appended += 1

        p.parent.mkdir(parents=True, exist_ok=True)
        wb.save(str(p))
        return {"status": "ok", "path": str(p), "sheet": ws.title,
                "cells_updated": cells_updated, "rows_appended": rows_appended}
    except PermissionError:
        return {"error": f"Permission denied: {path}"}
    except Exception as e:
        return {"error": str(e)}


# ── Document tools ─────────────────────────────────────────────────────────────

_MAX_DOC_PARAGRAPHS = 500


def read_document(path: str) -> dict:
    """Read paragraphs from a Word document (.docx) or plain text file."""
    denied = _permission_check("allow_file_read", "read_document")
    if denied:
        return denied
    if not path:
        return {"error": "path is required."}
    try:
        p = pathlib.Path(path).expanduser().resolve()
        if not p.exists():
            return {"error": f"File not found: {p}"}
        if not p.is_file():
            return {"error": f"Not a file: {p}"}

        suffix = p.suffix.lower()
        if suffix == ".docx":
            if not _HAS_DOCX:
                return {"error": "python-docx is not installed. Run: uv add python-docx"}
            doc = _DocxDocument(str(p))
            paragraphs = [
                {"index": i, "text": para.text}
                for i, para in enumerate(doc.paragraphs)
                if para.text.strip()
            ][:_MAX_DOC_PARAGRAPHS]
        elif suffix in {".txt", ".md", ".rst", ".text"}:
            lines = p.read_text(errors="replace").splitlines()
            paragraphs = [
                {"index": i, "text": line}
                for i, line in enumerate(lines)
                if line.strip()
            ][:_MAX_DOC_PARAGRAPHS]
        else:
            return {"error": f"Unsupported file type '{suffix}'. Supported: .docx, .txt, .md"}

        word_count = sum(len(para["text"].split()) for para in paragraphs)
        return {"path": str(p), "paragraphs": paragraphs, "word_count": word_count}
    except PermissionError:
        return {"error": f"Permission denied: {path}"}
    except Exception as e:
        return {"error": str(e)}


def edit_document(
    path: str,
    replacements: list | None = None,
    append_paragraphs: list | None = None,
    set_paragraph: dict | None = None,
) -> dict:
    """Edit a Word document (.docx): find/replace text, append paragraphs, or overwrite a paragraph."""
    denied = _permission_check("allow_file_write", "edit_document")
    if denied:
        return denied
    if not path:
        return {"error": "path is required."}
    if not replacements and not append_paragraphs and not set_paragraph:
        return {"error": "Provide at least one of: replacements, append_paragraphs, set_paragraph."}

    try:
        p = pathlib.Path(path).expanduser().resolve()
        if not p.exists():
            return {"error": f"File not found: {p}"}

        suffix = p.suffix.lower()
        if suffix != ".docx":
            return {"error": f"Unsupported file type '{suffix}'. edit_document supports .docx only."}
        if not _HAS_DOCX:
            return {"error": "python-docx is not installed. Run: uv add python-docx"}

        doc = _DocxDocument(str(p))
        replacements_made = 0

        if replacements:
            for rep in replacements:
                find_text = rep.get("find", "")
                replace_text = rep.get("replace", "")
                if not find_text:
                    continue
                for para in doc.paragraphs:
                    if find_text in para.text:
                        # Preserve runs by doing a full-text replace on the first run
                        for run in para.runs:
                            if find_text in run.text:
                                run.text = run.text.replace(find_text, replace_text, 1)
                                replacements_made += 1
                                break

        if set_paragraph:
            idx = set_paragraph.get("index")
            new_text = set_paragraph.get("text", "")
            if idx is not None and 0 <= idx < len(doc.paragraphs):
                para = doc.paragraphs[idx]
                # Clear all runs and set text on the first run
                if para.runs:
                    para.runs[0].text = new_text
                    for run in para.runs[1:]:
                        run.text = ""
                else:
                    para.add_run(new_text)

        paragraphs_appended = 0
        if append_paragraphs:
            for text in append_paragraphs:
                doc.add_paragraph(str(text))
                paragraphs_appended += 1

        doc.save(str(p))
        return {
            "status": "ok",
            "path": str(p),
            "replacements_made": replacements_made,
            "paragraphs_appended": paragraphs_appended,
        }
    except PermissionError:
        return {"error": f"Permission denied: {path}"}
    except Exception as e:
        return {"error": str(e)}


# ── PDF tools ─────────────────────────────────────────────────────────────────

_MAX_PDF_PAGES = 200
_DEFAULT_PDF_PAGES = 50
_MAX_PDF_BYTES = 200 * 1024 * 1024  # pre-flight cap on .pdf input size


def read_pdf(path: str, max_pages: int = _DEFAULT_PDF_PAGES) -> dict:
    """Extract text from a PDF file, one entry per page.

    Args:
        path: Absolute or home-relative path to the ``.pdf`` file.
        max_pages: Maximum number of pages to extract (1–200).
            Pages beyond this limit are silently skipped.

    Returns:
        A dict with ``path``, ``total_pages``, ``pages_returned``,
        ``word_count``, and ``pages`` (list of ``{page, text}`` dicts).
        Includes a ``note`` key when the PDF was truncated.
    """
    denied = _permission_check("allow_file_read", "read_pdf")
    if denied:
        return denied
    if not path:
        return {"error": "path is required."}
    if not _HAS_PYPDF:
        return {"error": "pypdf is not installed. Run: uv add pypdf"}
    try:
        # C1: coerce max_pages defensively — JSON input may arrive as float/str.
        try:
            max_pages = int(max_pages)
        except (TypeError, ValueError):
            max_pages = _DEFAULT_PDF_PAGES

        p = pathlib.Path(path).expanduser().resolve()
        if not p.exists():
            return {"error": f"File not found: {p}"}
        if not p.is_file():
            return {"error": f"Not a file: {p}"}
        if p.suffix.lower() != ".pdf":
            return {"error": f"Expected a .pdf file, got '{p.suffix}'."}
        if p.stat().st_size > _MAX_PDF_BYTES:
            return {
                "error": (
                    f"PDF exceeds {_MAX_PDF_BYTES // (1024 * 1024)} MB size limit."
                )
            }

        clamped = max(1, min(max_pages, _MAX_PDF_PAGES))
        reader = _PdfReader(str(p))
        if reader.is_encrypted:
            return {"error": "PDF is encrypted/password-protected and cannot be read."}

        total_pages = len(reader.pages)
        pages: list[dict] = []
        for i, page in enumerate(reader.pages[:clamped]):
            # C2: isolate per-page extraction so one corrupt page doesn't
            # abort the entire document.
            try:
                text = (page.extract_text() or "").strip()
            except Exception:
                text = ""
            if text:
                pages.append({"page": i + 1, "text": text})

        word_count = sum(len(pg["text"].split()) for pg in pages)
        result: dict = {
            "path": str(p),
            "total_pages": total_pages,
            "pages_returned": len(pages),
            "word_count": word_count,
            "pages": pages,
        }
        if total_pages > clamped:
            result["note"] = (
                f"PDF has {total_pages} pages; only the first {clamped} were read. "
                f"Pass max_pages to read more (up to {_MAX_PDF_PAGES})."
            )
        return result
    except PermissionError:
        return {"error": f"Permission denied: {path}"}
    except Exception as e:
        return {"error": str(e)}


# ── File-management tools ─────────────────────────────────────────────────────

_HOME_DIR = pathlib.Path.home()


def list_directory(path: str = "", show_hidden: bool = False) -> dict:
    """List the contents of a directory with name, type, size, and modification time.

    Args:
        path:        Directory to list (default: home directory).
        show_hidden: Include entries whose names start with ``'.'``.

    Returns:
        A dict with ``path``, ``count``, and ``entries`` (each entry has
        ``name``, ``type``, ``size_bytes``, and ``modified``).
    """
    denied = _permission_check("allow_file_read", "list_directory")
    if denied:
        return denied
    root = pathlib.Path(path).expanduser().resolve() if path else _HOME_DIR
    if not root.exists():
        return {"error": f"Path not found: {root}"}
    if not root.is_dir():
        return {"error": f"Not a directory: {root}"}
    entries: list[dict] = []
    try:
        for entry in sorted(root.iterdir(), key=lambda e: (e.is_file(), e.name.lower())):
            if not show_hidden and entry.name.startswith("."):
                continue
            try:
                stat = entry.stat()
                entries.append({
                    "name":       entry.name,
                    "type":       "file" if entry.is_file() else "dir",
                    "size_bytes": stat.st_size if entry.is_file() else 0,
                    "modified":   datetime.datetime.fromtimestamp(
                        stat.st_mtime
                    ).isoformat(timespec="seconds"),
                })
            except OSError:
                continue
    except PermissionError:
        return {"error": f"Permission denied: {root}"}
    return {"path": str(root), "count": len(entries), "entries": entries}


def move_file(src: str, dst: str) -> dict:
    """Move or rename a file or directory.

    Args:
        src: Source path (absolute or home-relative).
        dst: Destination path or target directory.

    Returns:
        ``{"status": "ok", "src": ..., "dst": ...}``
    """
    denied = _permission_check("allow_file_write", "move_file")
    if denied:
        return denied
    if not src or not dst:
        return {"error": "src and dst are required."}
    src_path = pathlib.Path(src).expanduser().resolve()
    dst_path = pathlib.Path(dst).expanduser().resolve()
    if not src_path.exists():
        return {"error": f"Source not found: {src_path}"}
    try:
        final = shutil.move(str(src_path), str(dst_path))
        return {"status": "ok", "src": str(src_path), "dst": str(final)}
    except PermissionError:
        return {"error": f"Permission denied moving {src_path}"}
    except Exception as e:
        return {"error": str(e)}


def copy_file(src: str, dst: str, overwrite: bool = False) -> dict:
    """Copy a file to a destination path or directory.

    Args:
        src:       Source file path.
        dst:       Destination file path or directory.
        overwrite: If False (default), error if the destination already exists.

    Returns:
        ``{"status": "ok", "src": ..., "dst": ..., "size_bytes": ...}``
    """
    denied = _permission_check("allow_file_write", "copy_file")
    if denied:
        return denied
    if not src or not dst:
        return {"error": "src and dst are required."}
    src_path = pathlib.Path(src).expanduser().resolve()
    dst_path = pathlib.Path(dst).expanduser().resolve()
    if not src_path.exists():
        return {"error": f"Source not found: {src_path}"}
    if not src_path.is_file():
        return {"error": f"Source is not a file: {src_path}"}
    # If dst is an existing directory, copy into it preserving the filename.
    if dst_path.is_dir():
        dst_path = dst_path / src_path.name
    if dst_path.exists() and not overwrite:
        return {
            "error": f"Destination already exists: {dst_path}. Pass overwrite=true to replace it.",
        }
    try:
        shutil.copy2(str(src_path), str(dst_path))
        return {
            "status":     "ok",
            "src":        str(src_path),
            "dst":        str(dst_path),
            "size_bytes": dst_path.stat().st_size,
        }
    except PermissionError:
        return {"error": f"Permission denied copying to {dst_path}"}
    except Exception as e:
        return {"error": str(e)}


def delete_file(path: str, permanent: bool = False) -> dict:
    """Delete a file or directory, defaulting to Trash on macOS.

    On macOS, ``permanent=False`` (default) moves the item to the Trash via
    Finder so it can be recovered.  Pass ``permanent=True`` to permanently
    remove it.  On non-macOS systems the item is always permanently removed.

    Safety: refuses to delete anything outside the user's home directory.

    Args:
        path:      Path to the file or directory to remove.
        permanent: If True, bypass the Trash and delete immediately.

    Returns:
        ``{"status": "ok", "path": ..., "method": "trash"|"permanent"}``
    """
    denied = _permission_check("allow_file_write", "delete_file")
    if denied:
        return denied
    if not path:
        return {"error": "path is required."}
    target = pathlib.Path(path).expanduser().resolve()
    if not target.exists():
        return {"error": f"Path not found: {target}"}
    # Safety guard: only allow deletion within the home directory.
    try:
        target.relative_to(_HOME_DIR)
    except ValueError:
        return {
            "error": (
                "For safety, delete_file only removes items inside your home directory. "
                f"Got: {target}"
            ),
        }
    try:
        if IS_MACOS and not permanent:
            safe_path = _escape_applescript(str(target))
            script = f'tell application "Finder" to delete POSIX file "{safe_path}"'
            proc = subprocess.run(
                ["osascript", "-e", script],
                capture_output=True, text=True, timeout=10,
            )
            if proc.returncode != 0:
                return {"error": proc.stderr.strip() or "Finder could not move item to Trash."}
            return {"status": "ok", "path": str(target), "method": "trash"}
        # Permanent deletion (or non-macOS).
        if target.is_dir():
            shutil.rmtree(str(target))
        else:
            target.unlink()
        return {"status": "ok", "path": str(target), "method": "permanent"}
    except PermissionError:
        return {"error": f"Permission denied: {target}"}
    except Exception as e:
        return {"error": str(e)}


def create_directory(path: str) -> dict:
    """Create a directory and any missing parents.

    Args:
        path: Absolute or home-relative path for the new directory.

    Returns:
        ``{"status": "ok"|"already_exists", "path": ...}``
    """
    denied = _permission_check("allow_file_write", "create_directory")
    if denied:
        return denied
    if not path:
        return {"error": "path is required."}
    target = pathlib.Path(path).expanduser().resolve()
    if target.exists():
        if target.is_dir():
            return {"status": "already_exists", "path": str(target)}
        return {"error": f"A file already exists at that path: {target}"}
    try:
        target.mkdir(parents=True, exist_ok=True)
        return {"status": "ok", "path": str(target)}
    except PermissionError:
        return {"error": f"Permission denied: {target}"}
    except Exception as e:
        return {"error": str(e)}


# ── Spotlight search ──────────────────────────────────────────────────────────

_SPOTLIGHT_KIND_MAP: dict[str, str] = {
    "document":    "kMDItemContentTypeTree == 'public.text'cd",
    "pdf":         "kMDItemContentType == 'com.adobe.pdf'",
    "image":       "kMDItemContentTypeTree == 'public.image'",
    "video":       "kMDItemContentTypeTree == 'public.movie'",
    "audio":       "kMDItemContentTypeTree == 'public.audio'",
    "folder":      "kMDItemContentType == 'public.folder'",
    "application": "kMDItemContentType == 'com.apple.application-bundle'",
    "code":        "kMDItemContentTypeTree == 'public.source-code'",
}


def search_files(
    query: str,
    scope: str = "~",
    kind: str = "",
    limit: int = 50,
) -> dict:
    """Search for files using macOS Spotlight (mdfind).

    Args:
        query: File-name fragment or Spotlight metadata query string.
        scope: Directory to restrict the search to (default: home directory).
        kind:  Optional type filter: ``"pdf"``, ``"image"``, ``"video"``,
               ``"audio"``, ``"folder"``, ``"application"``, ``"code"``,
               or ``"document"``.
        limit: Maximum results to return (1–200, default 50).

    Returns:
        ``{"query", "scope", "kind", "count", "results"}`` — each result has
        ``path`` and ``name``.
    """
    if not IS_MACOS:
        return {"error": "search_files uses Spotlight (mdfind) which requires macOS."}
    if not query:
        return {"error": "query is required."}
    try:
        limit = max(1, min(int(limit), 200))
    except (TypeError, ValueError):
        limit = 50
    scope_path = str(pathlib.Path(scope).expanduser().resolve())
    # Build the mdfind query string.
    kind = kind.lower().strip()
    if kind and kind in _SPOTLIGHT_KIND_MAP:
        kind_pred   = _SPOTLIGHT_KIND_MAP[kind]
        safe_q       = query.replace('"', "").replace("*", "").replace("(", "").replace(")", "")
        mdfind_query = f'({kind_pred}) && kMDItemFSName == "*{safe_q}*"cd'
    else:
        kind         = ""
        mdfind_query = query
    try:
        proc = subprocess.run(
            ["mdfind", "-onlyin", scope_path, mdfind_query],
            capture_output=True, text=True, timeout=10,
        )
        paths   = [p for p in proc.stdout.splitlines() if p.strip()][:limit]
        results = [{"path": p, "name": pathlib.Path(p).name} for p in paths]
        return {
            "query":   query,
            "scope":   scope_path,
            "kind":    kind or "any",
            "count":   len(results),
            "results": results,
        }
    except subprocess.TimeoutExpired:
        return {"error": "Spotlight search timed out (>10s). Try narrowing the scope."}
    except FileNotFoundError:
        return {"error": "mdfind not found. Requires macOS with Spotlight enabled."}
    except Exception as e:
        return {"error": str(e)}


# ── Deep research ──────────────────────────────────────────────────────────────


def _list_agents() -> dict:
    """Return available sub-agents and their one-line descriptions."""
    try:
        from agent.agents import list_agents  # lazy import — avoids circular dep at load
        return {"agents": list_agents()}
    except Exception as exc:
        return {"error": f"Failed to list agents: {exc}"}


def _run_agent(agent_name: str, task: str) -> dict:
    """Delegate *task* to the named sub-agent and return its final output.

    The sub-agent runs in a separate MCPClient subprocess with
    ``SYSCONTROL_AGENT_DEPTH=1`` set in its environment, so nested
    ``run_agent`` calls are blocked at that level.
    """
    # Depth guard — block recursive spawning from within a sub-agent process.
    try:
        depth = int(os.environ.get("SYSCONTROL_AGENT_DEPTH", "0"))
    except (ValueError, TypeError):
        depth = 0
    if depth > 0:
        return {
            "error": "Nested agent spawning is not permitted.",
            "hint": "Sub-agents cannot call run_agent recursively.",
        }

    denied = _permission_check("allow_agents", "run_agent")
    if denied:
        return denied

    if not agent_name or not agent_name.strip():
        return {"error": "agent_name is required."}
    if not task or not task.strip():
        return {"error": "task is required."}

    api_key  = os.environ.get("SYSCONTROL_API_KEY", DEFAULT_API_KEY)
    base_url = os.environ.get("SYSCONTROL_BASE_URL", DEFAULT_BASE_URL)
    model    = os.environ.get("SYSCONTROL_MODEL", DEFAULT_LOCAL_MODEL)

    try:
        # Lazy imports keep server startup fast and avoid circular references.
        from agent.agents import AgentNotFoundError, get_agent, list_agents
        from agent.runner import run_subagent
    except ImportError as exc:
        return {"error": f"Agent module unavailable: {exc}"}

    try:
        spec = get_agent(agent_name.strip())
    except AgentNotFoundError:
        return {
            "error": f"Agent '{agent_name}' not found.",
            "available": list_agents(),
        }

    try:
        llm = OpenAI(api_key=api_key, base_url=base_url, timeout=120.0)
        result = run_subagent(spec, task.strip(), llm, model)
        return {"agent": spec.name, "task": task.strip(), "result": result}
    except Exception as exc:
        return {"error": f"Agent '{agent_name}' failed: {exc}"}


def _run_deep_research(
    question: str, max_sources: int = 15, max_loops: int = 5,
) -> dict:
    """Execute a deep research investigation and return a cited report."""
    denied = _permission_check("allow_deep_research", "deep_research")
    if denied:
        return denied
    if not question or not question.strip():
        return {"error": "question is required."}

    api_key = os.environ.get("SYSCONTROL_API_KEY", DEFAULT_API_KEY)
    base_url = os.environ.get("SYSCONTROL_BASE_URL", DEFAULT_BASE_URL)
    model = os.environ.get("SYSCONTROL_MODEL", DEFAULT_LOCAL_MODEL)

    try:
        from deep_research.orchestrator import orchestrate  # lazy import

        client = OpenAI(api_key=api_key, base_url=base_url, timeout=120.0)
        return orchestrate(
            question=question.strip(),
            search_fn=web_search,
            fetch_fn=web_fetch,
            llm_client=client,
            model=model,
            max_loops=max(1, min(max_loops, 10)),
            max_sources=max(3, min(max_sources, 25)),
        )
    except Exception as e:
        return {"error": f"Research failed: {e}"}


# ── Permission gate ────────────────────────────────────────────────────────────
#
# All sensitive tools are disabled until the user explicitly opts in via
# ~/.syscontrol/config.json. This prevents the agent from accessing private
# data or performing actions without the user's knowledge.
#
# Example config.json enabling all gates:
#   {
#     "allow_shell":           true,
#     "allow_messaging":       true,
#     "allow_message_history": true,
#     "allow_screenshot":      true,
#     "allow_file_read":       true,
#     "allow_file_write":      true,
#     "allow_calendar":        true,
#     "allow_contacts":        true,
#     "allow_accessibility":   true,
#     "allow_email":           true,
#     "allow_notes":           true,
#     "allow_brew":            true,
#     "allow_agents":          true
#   }

_SYSCONTROL_CONFIG_FILE = _REMINDER_DIR / "config.json"


_CONFIG_CACHE: dict = {}
_CONFIG_TTL: float = 5.0           # seconds; config changes take effect within one TTL window
_CONFIG_CACHE_TIME: float = float("-inf")  # force a disk read on the very first call
_CONFIG_LOCK = threading.Lock()


def _load_config() -> dict:
    """Load ~/.syscontrol/config.json, cached for _CONFIG_TTL seconds.

    Thread-safe: uses double-checked locking so concurrent callers
    never see a partially-updated cache.
    """
    global _CONFIG_CACHE, _CONFIG_CACHE_TIME
    now = time.monotonic()
    if now - _CONFIG_CACHE_TIME < _CONFIG_TTL:
        return _CONFIG_CACHE
    with _CONFIG_LOCK:
        # Re-check after acquiring; another thread may have refreshed.
        if now - _CONFIG_CACHE_TIME < _CONFIG_TTL:
            return _CONFIG_CACHE
        try:
            _CONFIG_CACHE = json.loads(_SYSCONTROL_CONFIG_FILE.read_text())
        except (FileNotFoundError, json.JSONDecodeError, OSError):
            _CONFIG_CACHE = {}
        _CONFIG_CACHE_TIME = now
        return _CONFIG_CACHE


def _permission_check(flag: str, tool_name: str) -> dict | None:
    """
    Returns None if *flag* is enabled in config (tool may proceed).
    Returns an error dict if the tool is disabled, describing how to enable it.
    """
    if _load_config().get(flag, False):
        return None  # permitted
    return {
        "error": f"{tool_name} is disabled by default for security.",
        "hint": (
            f'To enable it, add "{flag}": true to ~/.syscontrol/config.json.\n'
            f"Example: {{\"{ flag }\": true}}"
        ),
        "config_path": str(_SYSCONTROL_CONFIG_FILE),
    }


def run_shell_command(command: str, timeout: int = 30) -> dict:
    """
    Execute a shell command and return stdout, stderr, and exit code.
    Requires ``allow_shell: true`` in ~/.syscontrol/config.json.
    """
    denied = _permission_check("allow_shell", "run_shell_command")
    if denied:
        return denied
    if not command:
        return {"error": "command is required."}
    timeout = max(1, min(timeout, 120))
    try:
        proc = subprocess.run(
            ["bash", "-c", command],
            capture_output=True, text=True,
            timeout=timeout,
        )
        return {
            "command": command,
            "returncode": proc.returncode,
            "stdout": proc.stdout[:8000],
            "stderr": proc.stderr[:2000],
        }
    except subprocess.TimeoutExpired:
        return {"error": f"Command timed out after {timeout}s.", "command": command}
    except Exception as e:
        return {"error": str(e), "command": command}


# ── Calendar tool ─────────────────────────────────────────────────────────────

def _parse_calendar_items(raw_output: str) -> list[dict]:
    """Parse the pipe-delimited osascript output from the Calendar AppleScript."""
    events: list[dict] = []
    for item in raw_output.split("||"):
        item = item.strip()
        if not item:
            continue
        parts = item.split("|")
        if len(parts) >= 4:
            events.append({
                "calendar": parts[0],
                "title":    parts[1],
                "start":    parts[2],
                "end":      parts[3],
                "location": parts[4] if len(parts) > 4 else "",
            })
    return events


def get_calendar_events(lookahead_days: int = 7) -> dict:
    """Return upcoming calendar events from macOS Calendar.app via AppleScript."""
    denied = _permission_check("allow_calendar", "get_calendar_events")
    if denied:
        return denied
    if not IS_MACOS:
        return {"error": "get_calendar_events requires macOS."}
    lookahead_days = max(1, min(lookahead_days, 90))
    script = f"""
set resultList to {{}}
set startDate to current date
set endDate to startDate + ({lookahead_days} * days)

tell application "Calendar"
    repeat with theCalendar in calendars
        set calName to name of theCalendar
        set theEvents to (every event of theCalendar whose start date >= startDate and start date <= endDate)
        repeat with theEvent in theEvents
            set evtSummary to summary of theEvent
            set evtStart to start date of theEvent as string
            set evtEnd to end date of theEvent as string
            try
                set evtLocation to location of theEvent
            on error
                set evtLocation to ""
            end try
            set end of resultList to (calName & "|" & evtSummary & "|" & evtStart & "|" & evtEnd & "|" & evtLocation)
        end repeat
    end repeat
end tell

set AppleScript's text item delimiters to "||"
return resultList as string
"""
    try:
        proc = subprocess.run(
            ["osascript", "-e", script],
            capture_output=True, text=True, timeout=20,
        )
        if proc.returncode != 0:
            err = proc.stderr.strip()
            return {
                "error": err or "Calendar access denied.",
                "hint": "Grant Calendar access to Terminal in System Settings → Privacy & Security → Calendars.",
            }
        events = _parse_calendar_items(proc.stdout.strip())
        return {"lookahead_days": lookahead_days, "event_count": len(events), "events": events}
    except subprocess.TimeoutExpired:
        return {"error": "Calendar query timed out. Calendar.app may be unresponsive."}
    except Exception as e:
        return {"error": str(e)}


# ── Contacts tool ─────────────────────────────────────────────────────────────

def _parse_contact_items(raw_output: str) -> list[dict]:
    """Parse the pipe-delimited osascript output from the Contacts AppleScript."""
    contacts: list[dict] = []
    for item in raw_output.split("||"):
        item = item.strip()
        if not item:
            continue
        parts = item.split("|")
        person_name = parts[0] if parts else ""
        phones = [p for p in (parts[1].split(";") if len(parts) > 1 else []) if p]
        emails = [e for e in (parts[2].split(";") if len(parts) > 2 else []) if e]
        contacts.append({"name": person_name, "phones": phones, "emails": emails})
    return contacts


def get_contact(name: str) -> dict:
    """Search macOS Contacts.app for a person by name and return their details."""
    denied = _permission_check("allow_contacts", "get_contact")
    if denied:
        return denied
    if not IS_MACOS:
        return {"error": "get_contact requires macOS."}
    if not name:
        return {"error": "name is required."}
    script = f"""
set searchName to "{_escape_applescript(name)}"
set resultList to {{}}

tell application "Contacts"
    set matchedPeople to every person whose name contains searchName
    repeat with p in matchedPeople
        set personName to name of p
        -- phones
        set phoneStr to ""
        repeat with ph in phones of p
            set phoneStr to phoneStr & value of ph & ";"
        end repeat
        -- emails
        set emailStr to ""
        repeat with em in emails of p
            set emailStr to emailStr & value of em & ";"
        end repeat
        set end of resultList to (personName & "|" & phoneStr & "|" & emailStr)
    end repeat
end tell

set AppleScript's text item delimiters to "||"
return resultList as string
"""
    try:
        proc = subprocess.run(
            ["osascript", "-e", script],
            capture_output=True, text=True, timeout=15,
        )
        if proc.returncode != 0:
            err = proc.stderr.strip()
            return {
                "error": err or "Contacts access denied.",
                "hint": "Grant Contacts access to Terminal in System Settings → Privacy & Security → Contacts.",
            }
        contacts = _parse_contact_items(proc.stdout.strip())
        return {"query": name, "count": len(contacts), "contacts": contacts}
    except subprocess.TimeoutExpired:
        return {"error": "Contacts query timed out."}
    except Exception as e:
        return {"error": str(e)}


# ── Apple Notes tools ─────────────────────────────────────────────────────────


def list_notes(folder: str = "", count: int = 20) -> dict:
    """List notes from macOS Notes.app.

    Args:
        folder: Folder name to restrict results to (default: all folders).
        count:  Maximum notes to return (1–100, default 20).

    Returns:
        ``{"count", "notes"}`` — each note has ``id``, ``name``, ``folder``,
        ``created``, and ``modified``.
    """
    denied = _permission_check("allow_notes", "list_notes")
    if denied:
        return denied
    if not IS_MACOS:
        return {"error": "list_notes requires macOS (Notes.app)."}
    try:
        count = max(1, min(int(count), 100))
    except (TypeError, ValueError):
        count = 20
    if folder:
        safe_folder = _escape_applescript(folder)
        source      = f'notes of folder "{safe_folder}"'
    else:
        source = "notes"
    script = f"""
set output to ""
tell application "Notes"
    set noteList to {source}
    set total to count of noteList
    set lim to {count}
    if total < lim then set lim to total
    repeat with i from 1 to lim
        set n to item i of noteList
        set nId to id of n
        set nName to name of n
        set nCreated to (creation date of n) as string
        set nModified to (modification date of n) as string
        try
            set nFolder to name of container of n
        on error
            set nFolder to "Notes"
        end try
        set output to output & nId & "|||" & nName & "|||" & nFolder & "|||" & nCreated & "|||" & nModified & "|||||"
    end repeat
end tell
return output
"""
    try:
        proc = subprocess.run(
            ["osascript", "-e", script],
            capture_output=True, text=True, timeout=20,
        )
        if proc.returncode != 0:
            return {
                "error": proc.stderr.strip() or "Notes access denied.",
                "hint": "Grant Notes access in System Settings → Privacy & Security → Automation.",
            }
        notes: list[dict] = []
        for record in proc.stdout.strip().split("|||||"):
            record = record.strip()
            if not record:
                continue
            parts = record.split("|||")
            notes.append({
                "id":       parts[0] if len(parts) > 0 else "",
                "name":     parts[1] if len(parts) > 1 else "",
                "folder":   parts[2] if len(parts) > 2 else "",
                "created":  parts[3] if len(parts) > 3 else "",
                "modified": parts[4] if len(parts) > 4 else "",
            })
        return {"count": len(notes), "notes": notes}
    except subprocess.TimeoutExpired:
        return {"error": "Notes.app timed out."}
    except Exception as e:
        return {"error": str(e)}


def read_note(name: str) -> dict:
    """Read the full text of a note from macOS Notes.app.

    Searches for a note whose title contains *name* (case-insensitive) and
    returns the first match.

    Args:
        name: Full or partial note title.

    Returns:
        ``{"id", "name", "folder", "body", "created", "modified"}``
    """
    denied = _permission_check("allow_notes", "read_note")
    if denied:
        return denied
    if not IS_MACOS:
        return {"error": "read_note requires macOS (Notes.app)."}
    if not name:
        return {"error": "name is required."}
    safe_name = _escape_applescript(name)
    script = f"""
tell application "Notes"
    set matches to (every note whose name contains "{safe_name}")
    if (count of matches) = 0 then return "NOT_FOUND"
    set n to item 1 of matches
    try
        set nFolder to name of container of n
    on error
        set nFolder to "Notes"
    end try
    return (id of n) & "|" & (name of n) & "|" & nFolder & "|" & ((creation date of n) as string) & "|" & ((modification date of n) as string) & "|" & (body of n)
end tell
"""
    try:
        proc = subprocess.run(
            ["osascript", "-e", script],
            capture_output=True, text=True, timeout=15,
        )
        raw = proc.stdout.strip()
        if proc.returncode != 0:
            return {
                "error": proc.stderr.strip() or "Notes access denied.",
                "hint": "Grant Notes access in System Settings → Privacy & Security → Automation.",
            }
        if raw == "NOT_FOUND":
            return {"error": f"No note found matching: {name!r}"}
        # Body may contain pipe characters — split on first 5 delimiters only.
        parts = raw.split("|", 5)
        return {
            "id":       parts[0] if len(parts) > 0 else "",
            "name":     parts[1] if len(parts) > 1 else "",
            "folder":   parts[2] if len(parts) > 2 else "",
            "created":  parts[3] if len(parts) > 3 else "",
            "modified": parts[4] if len(parts) > 4 else "",
            "body":     parts[5] if len(parts) > 5 else "",
        }
    except subprocess.TimeoutExpired:
        return {"error": "Notes.app timed out reading note."}
    except Exception as e:
        return {"error": str(e)}


def create_note(title: str, body: str, folder: str = "") -> dict:
    """Create a new note in macOS Notes.app.

    Args:
        title:  Note title.
        body:   Note body text.
        folder: Folder name to create the note in (default: default Notes folder).

    Returns:
        ``{"status": "created", "id", "name", "folder"}``
    """
    denied = _permission_check("allow_notes", "create_note")
    if denied:
        return denied
    if not IS_MACOS:
        return {"error": "create_note requires macOS (Notes.app)."}
    if not title:
        return {"error": "title is required."}
    safe_title = _escape_applescript(title)
    safe_body  = _escape_applescript(body)
    if folder:
        safe_folder    = _escape_applescript(folder)
        container_open = f'tell folder "{safe_folder}"'
        container_key  = f'"{safe_folder}"'
    else:
        container_open = "tell default account"
        container_key  = '"Notes"'
    script = (
        f'tell application "Notes"\n'
        f'    {container_open}\n'
        f'        set newNote to make new note with properties '
        f'{{name:"{safe_title}", body:"{safe_body}"}}\n'
        f'        set nId to id of newNote\n'
        f'        try\n'
        f'            set nFolder to name of container of newNote\n'
        f'        on error\n'
        f'            set nFolder to {container_key}\n'
        f'        end try\n'
        f'        return nId & "|" & "{safe_title}" & "|" & nFolder\n'
        f'    end tell\n'
        f'end tell\n'
    )
    try:
        proc = subprocess.run(
            ["osascript", "-e", script],
            capture_output=True, text=True, timeout=15,
        )
        if proc.returncode != 0:
            return {
                "error": proc.stderr.strip() or "Could not create note.",
                "hint": "Grant Notes access in System Settings → Privacy & Security → Automation.",
            }
        parts = proc.stdout.strip().split("|", 2)
        return {
            "status": "created",
            "id":     parts[0] if len(parts) > 0 else "",
            "name":   parts[1] if len(parts) > 1 else title,
            "folder": parts[2] if len(parts) > 2 else (folder or "Notes"),
        }
    except subprocess.TimeoutExpired:
        return {"error": "Notes.app timed out creating note."}
    except Exception as e:
        return {"error": str(e)}


# ── macOS Shortcuts tool ──────────────────────────────────────────────────────

def run_shortcut(name: str, input_text: str = "") -> dict:
    """Run a macOS Shortcut by name (Shortcuts.app)."""
    if not IS_MACOS:
        return {"error": "run_shortcut requires macOS."}
    if not name:
        return {"error": "name is required."}
    cmd = ["shortcuts", "run", name]
    try:
        proc = subprocess.run(
            cmd,
            input=input_text or None, text=True,
            capture_output=True, timeout=60,
        )
        if proc.returncode != 0:
            stderr = proc.stderr.strip()
            return {
                "error": stderr or f"Shortcut '{name}' failed or does not exist.",
                "hint": "Check the shortcut name in Shortcuts.app — it's case-sensitive.",
            }
        return {
            "status": "ok",
            "shortcut": name,
            "output": proc.stdout.strip() or None,
        }
    except subprocess.TimeoutExpired:
        return {"error": f"Shortcut '{name}' timed out after 60s."}
    except FileNotFoundError:
        return {"error": "shortcuts CLI not found. Requires macOS 12+."}
    except Exception as e:
        return {"error": str(e)}


# ── Frontmost app tool ────────────────────────────────────────────────────────

def get_frontmost_app() -> dict:
    """Return the name of the application currently in focus."""
    denied = _permission_check("allow_accessibility", "get_frontmost_app")
    if denied:
        return denied
    if not IS_MACOS:
        return {"error": "get_frontmost_app requires macOS."}
    script = 'tell application "System Events" to get name of first process whose frontmost is true'
    try:
        proc = subprocess.run(
            ["osascript", "-e", script],
            capture_output=True, text=True, timeout=5,
        )
        if proc.returncode != 0:
            return {
                "error": proc.stderr.strip(),
                "hint": "Grant Terminal Accessibility access in System Settings → Privacy & Security → Accessibility.",
            }
        app_name = proc.stdout.strip()
        return {"app": app_name}
    except Exception as e:
        return {"error": str(e)}


# ── Do Not Disturb / Focus tool ───────────────────────────────────────────────

def toggle_do_not_disturb(enabled: bool) -> dict:
    """
    Enable or disable macOS Focus / Do Not Disturb.
    Uses the macOS `shortcuts` CLI to run the built-in Focus shortcuts.
    """
    if not IS_MACOS:
        return {"error": "toggle_do_not_disturb requires macOS."}
    # Attempt multiple known shortcut names for DnD/Focus
    if enabled:
        candidates = ["Turn On Do Not Disturb", "Enable Do Not Disturb", "Turn On Focus"]
    else:
        candidates = ["Turn Off Do Not Disturb", "Disable Do Not Disturb", "Turn Off Focus"]

    last_err = None
    for shortcut_name in candidates:
        try:
            proc = subprocess.run(
                ["shortcuts", "run", shortcut_name],
                capture_output=True, text=True, timeout=15,
            )
            if proc.returncode == 0:
                return {"status": "ok", "dnd_enabled": enabled, "shortcut_used": shortcut_name}
            last_err = proc.stderr.strip() or f"Shortcut '{shortcut_name}' not found."
        except subprocess.TimeoutExpired:
            last_err = f"Shortcut '{shortcut_name}' timed out."
        except FileNotFoundError:
            return {"error": "shortcuts CLI not found. Requires macOS 12+."}
        except Exception as e:
            last_err = str(e)

    # Fallback: try direct osascript Focus toggle (macOS 12+)
    try:
        script = 'do shell script "shortcuts run \'Focus\'"'
        proc2 = subprocess.run(
            ["osascript", "-e", script],
            capture_output=True, text=True, timeout=10,
        )
        if proc2.returncode == 0:
            return {"status": "ok", "dnd_enabled": enabled}
    except Exception:
        pass

    return {
        "error": last_err or "Could not toggle Focus mode.",
        "hint": (
            "Create a Shortcut named 'Turn On Do Not Disturb' or 'Turn Off Do Not Disturb' "
            "in Shortcuts.app, or check System Settings → Focus for the exact Focus name."
        ),
    }


# ── Eject disk tool ───────────────────────────────────────────────────────────

def eject_disk(mountpoint: str) -> dict:
    """Unmount and eject a disk by its mountpoint (e.g. '/Volumes/MyDrive')."""
    if not IS_MACOS:
        return {"error": "eject_disk requires macOS."}
    if not mountpoint:
        return {"error": "mountpoint is required."}
    p = pathlib.Path(mountpoint)
    if not p.exists():
        return {"error": f"Mountpoint does not exist: {mountpoint}"}
    try:
        proc = subprocess.run(
            ["diskutil", "eject", mountpoint],
            capture_output=True, text=True, timeout=15,
        )
        if proc.returncode != 0:
            return {"error": proc.stderr.strip() or proc.stdout.strip()}
        return {"status": "ejected", "mountpoint": mountpoint, "detail": proc.stdout.strip()}
    except subprocess.TimeoutExpired:
        return {"error": "diskutil timed out during eject."}
    except Exception as e:
        return {"error": str(e)}


# ── Tool self-extension ────────────────────────────────────────────────────────

def list_user_tools() -> dict:
    """Return all user-installed tools (created via create_tool)."""
    if _FROZEN:
        return {"count": 0, "user_tools": [], "note": "Tool creation is not available in the bundled app."}
    text  = _SERVER_FILE.read_text()
    names = [
        line.split(_USER_TOOL_FN_MARKER)[1].strip().rstrip("─").strip()
        for line in text.splitlines()
        # Only count actual comment lines, not the constant definition itself
        if line.strip().startswith(_USER_TOOL_FN_MARKER)
    ]
    return {
        "count":      len(names),
        "user_tools": names,
        "note":       "Restart the agent (syscontrol) for installed tools to appear.",
    }


# ── Memory tools ──────────────────────────────────────────────────────────────

def read_memory() -> dict:
    """Return the contents of the persistent memory file, or a notice if empty."""
    if not _MEMORY_FILE.exists():
        return {"memory": None, "note": "No memory file found. Use append_memory_note to start one."}
    text = _MEMORY_FILE.read_text(encoding="utf-8").strip()
    if not text:
        return {"memory": None, "note": "Memory file exists but is empty."}
    return {"memory": text}


def append_memory_note(note: str) -> dict:
    """Append a concise note to the persistent memory file."""
    if not note or not note.strip():
        return {"error": "Note is empty — nothing saved."}
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    entry = f"\n- [{timestamp}] {note.strip()}\n"
    with _MEMORY_LOCK, _MEMORY_FILE.open("a", encoding="utf-8") as fh:
        if _HAS_FCNTL:
            _fcntl.flock(fh, _fcntl.LOCK_EX)
        try:
            fh.seek(0, 2)
            if fh.tell() == 0:
                fh.write("# SysControl Memory\n\n")
            fh.write(entry)
        finally:
            if _HAS_FCNTL:
                _fcntl.flock(fh, _fcntl.LOCK_UN)
    return {"saved": note.strip(), "timestamp": timestamp}


# Names that cannot be called or imported, directly or indirectly.
_FORBIDDEN_CALL_NAMES: frozenset[str] = frozenset({
    "eval", "exec", "compile", "__import__", "open",
    "globals", "locals", "vars", "getattr", "setattr", "delattr",
})
_FORBIDDEN_IMPORT_MODULES: frozenset[str] = frozenset({
    "os", "subprocess", "shutil", "socket", "ctypes", "importlib",
    "pickle", "marshal", "pty", "popen2",
})
# Method/attribute names that are dangerous regardless of receiver.
_FORBIDDEN_ATTR_NAMES: frozenset[str] = frozenset({
    "system", "popen", "spawn", "spawnl", "spawnv", "spawnvp", "spawnve",
    "execv", "execve", "execvp", "execvpe", "execl", "execle", "execlp",
    "rmtree", "unlink", "remove", "rmdir", "fork", "kill",
})


def _ast_security_scan(tree: ast.AST) -> list[str]:
    """Walk *tree* and return a list of human-readable security violations.

    Catches:
      * Direct calls to forbidden builtins (eval, exec, getattr, …).
      * Imports of forbidden modules (os, subprocess, importlib, …).
      * Attribute references whose final element is a dangerous method name
        (e.g. ``foo.system``, ``foo.popen``, ``foo.rmtree``).
      * Subscription of ``__builtins__`` (e.g. ``__builtins__["eval"]``).
    """
    violations: list[str] = []

    def _attr_chain_root(node: ast.AST) -> str | None:
        cur = node
        while isinstance(cur, ast.Attribute):
            cur = cur.value
        return cur.id if isinstance(cur, ast.Name) else None

    for node in ast.walk(tree):
        # Forbidden imports.
        if isinstance(node, ast.Import):
            for alias in node.names:
                root = alias.name.split(".", 1)[0]
                if root in _FORBIDDEN_IMPORT_MODULES:
                    violations.append(f"import of '{alias.name}'")
        elif isinstance(node, ast.ImportFrom):
            root = (node.module or "").split(".", 1)[0]
            if root in _FORBIDDEN_IMPORT_MODULES:
                violations.append(f"import from '{node.module}'")

        # Forbidden direct calls (eval(...), exec(...), getattr(...), …).
        elif isinstance(node, ast.Call):
            func = node.func
            if isinstance(func, ast.Name) and func.id in _FORBIDDEN_CALL_NAMES:
                violations.append(f"call to '{func.id}()'")
            elif isinstance(func, ast.Attribute):
                if func.attr in _FORBIDDEN_ATTR_NAMES:
                    violations.append(f"call to '.{func.attr}()' on any receiver")
                chain_root = _attr_chain_root(func)
                if chain_root is not None and chain_root in _FORBIDDEN_IMPORT_MODULES:
                    violations.append(f"call into '{chain_root}' module")

        # Forbidden attribute references on builtins (e.g. __builtins__.eval).
        elif isinstance(node, ast.Attribute):
            if (
                isinstance(node.value, ast.Name)
                and node.value.id in {"__builtins__", "builtins"}
            ):
                violations.append(f"attribute '{node.attr}' on __builtins__")

        # __builtins__["eval"] style subscript.
        elif (
            isinstance(node, ast.Subscript)
            and isinstance(node.value, ast.Name)
            and node.value.id in {"__builtins__", "builtins"}
        ):
            violations.append("subscript of __builtins__")

    # De-duplicate while preserving order so the error message is stable.
    seen: set[str] = set()
    deduped: list[str] = []
    for v in violations:
        if v not in seen:
            seen.add(v)
            deduped.append(v)
    return deduped


def _validate_tool_code(
    name: str, description: str, implementation: str, server_text: str,
) -> tuple[str, str, list[str]] | dict:
    """Validate a new tool's name, description, and implementation source code.

    Returns:
        On success: ``(func_name, fn_lambda, security_warnings)`` triple.
        On failure: an error dict suitable for returning directly from the tool.
    """
    if not name or not re.match(r"^[a-z][a-z0-9_]*$", name):
        return {
            "error": (
                "Tool name must start with a lowercase letter and contain only "
                "lowercase letters, digits, and underscores (e.g. 'get_spotify_track')."
            )
        }
    if f'"{name}":' in server_text:
        return {"error": f"A tool named '{name}' already exists. Choose a different name."}
    if not description.strip():
        return {"error": "description is required."}
    if not implementation.strip():
        return {"error": "implementation is required."}

    try:
        tree = ast.parse(implementation)
    except SyntaxError as exc:
        return {"error": f"Syntax error in implementation: {exc}"}

    func_defs = [n for n in ast.walk(tree) if isinstance(n, ast.FunctionDef)]
    if not func_defs:
        return {"error": "implementation must define at least one function (def ...)."}

    func_name   = func_defs[0].name
    func_params = [a.arg for a in func_defs[0].args.args if a.arg != "self"]

    security_warnings = _ast_security_scan(tree)
    if security_warnings:
        return {
            "error": (
                "Implementation rejected by security scan: "
                + "; ".join(security_warnings)
                + ". User-defined tools cannot import os/subprocess/shutil/socket/ctypes/"
                "importlib, call eval/exec/compile, or reach builtins indirectly."
            )
        }

    if func_params:
        param_str = ", ".join(f'args.get("{p}")' for p in func_params)
        fn_lambda = f"lambda args: {func_name}({param_str})"
    else:
        fn_lambda = f"lambda _: {func_name}()"

    return func_name, fn_lambda, security_warnings


def _inject_tool(
    name: str, description: str, implementation: str,
    schema: dict, fn_lambda: str, server_text: str,
) -> str | dict:
    """Inject the function and TOOLS entry into *server_text*.

    Returns:
        Updated server source text on success, or an error dict on failure.
    """
    schema_str = json.dumps(schema)

    fn_section = (
        f"\n\n{_USER_TOOL_FN_MARKER} {name} "
        + "\u2500" * max(1, 74 - len(name))
        + f"\n\n{implementation.rstrip()}\n"
    )
    tools_dict_start = "\nTOOLS = {"
    if tools_dict_start not in server_text:
        return {"error": "Could not locate 'TOOLS = {' in server.py. The file may be malformed."}
    server_text = server_text.replace(tools_dict_start, fn_section + tools_dict_start, 1)

    tools_entry = (
        f'    "{name}": {{\n'
        f'        "description": {json.dumps(description)},\n'
        f'        "inputSchema": {schema_str},\n'
        f'        "fn": {fn_lambda},\n'
        f'    }},\n'
        f'    {_USER_TOOL_REG_MARKER}\n'
    )
    if _USER_TOOL_REG_MARKER not in server_text:
        return {
            "error": "Could not locate TOOLS insertion anchor. The registry marker may be missing."
        }
    server_text = server_text.replace(f"    {_USER_TOOL_REG_MARKER}\n", tools_entry, 1)

    try:
        compile(server_text, str(_SERVER_FILE), "exec")
    except SyntaxError as exc:
        return {"error": f"Generated code has a syntax error (not written): {exc}"}

    try:
        _SERVER_FILE.write_text(server_text)
    except Exception as exc:
        return {"error": f"Failed to write server.py: {exc}"}

    return server_text


def _update_prompt_json(name: str, description: str, prompt_doc: str) -> bool:
    """Append *name*/*description*/*prompt_doc* to prompt.json.  Best-effort only."""
    if not prompt_doc.strip():
        return False
    try:
        with open(_PROMPT_FILE) as fh:
            pdata = json.load(fh)
        p         = pdata["system_prompt"]["prompt"]
        qr_marker = "\u2550" * 55 + "\n## TOOL SELECTION QUICK-REFERENCE"
        if qr_marker in p:
            tool_doc = (
                f"\n**{name}** (user-defined)\n"
                f"  Description: {description}\n"
                f"  Usage: {prompt_doc}\n\n"
            )
            p = p.replace(qr_marker, tool_doc + qr_marker)
            list_row = "| List user-installed custom tools"
            if list_row in p:
                p = p.replace(
                    list_row,
                    f"| {description[:52]:<52} | {name:<27} |\n{list_row}",
                )
        pdata["system_prompt"]["prompt"] = p
        with open(_PROMPT_FILE, "w") as fh:
            json.dump(pdata, fh, indent=2, ensure_ascii=False)
        return True
    except Exception:
        return False  # prompt.json update is best-effort; don't fail the whole call


def create_tool(
    name:               str,
    description:        str,
    parameters_schema:  dict | None,
    implementation:     str,
    prompt_doc:         str = "",
) -> dict:
    """
    Generate, validate, and install a new MCP tool into server.py.

    Requires allow_tool_creation: true in ~/.syscontrol/config.json.
    """
    if _FROZEN:
        return {
            "error": (
                "create_tool is not available in the bundled .app — "
                "the bundle is read-only."
            )
        }

    denied = _permission_check("allow_tool_creation", "create_tool")
    if denied:
        return denied

    server_text = _SERVER_FILE.read_text()
    validation  = _validate_tool_code(name, description, implementation, server_text)
    if isinstance(validation, dict):
        return validation  # error dict
    func_name, fn_lambda, security_warnings = validation

    schema  = parameters_schema or {"type": "object", "properties": {}, "required": []}
    updated = _inject_tool(name, description, implementation, schema, fn_lambda, server_text)
    if isinstance(updated, dict):
        return updated  # error dict

    prompt_updated = _update_prompt_json(name, description, prompt_doc)

    return {
        "success":           True,
        "tool_name":         name,
        "function_name":     func_name,
        "security_warnings": security_warnings,
        "prompt_updated":    prompt_updated,
        "note": (
            f"Tool '{name}' installed in server.py. "
            "Restart the agent (syscontrol) for it to take effect."
        ),
        "code_written": implementation,
    }


# ── Tool registry ─────────────────────────────────────────────────────────────

TOOLS = {
    "get_cpu_usage": {
        "description": "Returns CPU usage percentage (total and per-core), core count, and frequency, with an inline bar chart.",
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: _cpu_with_chart(),
    },
    "get_ram_usage": {
        "description": "Returns RAM and swap memory usage (total, used, available, percent), with an inline stacked bar chart.",
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: _ram_with_chart(),
    },
    "get_gpu_usage": {
        "description": "Returns GPU load, VRAM usage, and temperature (requires nvidia-ml-py on NVIDIA hardware), with an inline grouped bar chart.",
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: _gpu_with_chart(),
    },
    "get_disk_usage": {
        "description": "Returns disk partition usage and I/O counters.",
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: get_disk_usage(),
    },
    "get_network_usage": {
        "description": "Returns total bytes sent/received and network interface status.",
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: get_network_usage(),
    },
    "get_realtime_io": {
        "description": "Measures actual disk and network I/O throughput by sampling twice over an interval. Returns disk read/write in MB/s and network download/upload in MB/s and Mbps. Call this instead of get_disk_usage or get_network_usage when the user asks about current speed or throughput.",
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "interval": {
                    "type": "integer",
                    "description": "Sampling interval in seconds (1–3). Default 1.",
                    "default": 1,
                    "minimum": 1,
                    "maximum": 3
                }
            },
            "required": []
        },
        "fn": lambda args: get_realtime_io(args.get("interval", 1)),
    },
    "get_top_processes": {
        "description": "Returns the top N resource-hungry processes sorted by CPU or memory usage.",
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "n": {"type": "integer", "description": "Number of processes to return (default 10)", "default": 10},
                "sort_by": {"type": "string", "enum": ["cpu", "memory"], "description": "Sort by 'cpu' or 'memory'", "default": "cpu"}
            },
            "required": []
        },
        "fn": lambda args: get_top_processes(args.get("n", 10), args.get("sort_by", "cpu")),
    },
    "get_full_snapshot": {
        "description": "Returns a full system snapshot: CPU, RAM, GPU, disk, network, and top processes.",
        "parallel": False,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: get_full_snapshot(),
    },
    "get_device_specs": {
        "description": "Returns static hardware specifications: CPU model, core count, total RAM, GPU model and VRAM, disk capacities, and OS details.",
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: get_device_specs(),
    },
    "get_battery_status": {
        "description": "Returns battery percentage, charging state, and estimated time remaining. Returns an error on desktops with no battery.",
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: get_battery_status(),
    },
    "get_temperature_sensors": {
        "description": "Returns CPU and motherboard temperature sensor readings. On macOS, returns a helpful message with alternatives (psutil cannot access kernel sensors on Darwin). On Linux/Windows, returns sensor groups with current, high, and critical thresholds.",
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: get_temperature_sensors(),
    },
    "get_system_uptime": {
        "description": "Returns how long the system has been running, the last boot time, and the 1/5/15-minute load averages.",
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: get_system_uptime(),
    },
    "get_system_alerts": {
        "description": "Scans all key system metrics (CPU, RAM, swap, disk partitions, GPU, battery) and returns a prioritized list of critical/warning alerts. Call this first for general 'why is my machine slow?' questions as a quick triage tool.",
        "parallel": False,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: get_system_alerts(),
    },
    "get_network_connections": {
        "description": "Returns all active TCP/UDP connections with local/remote addresses, status, and the owning process name.",
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: get_network_connections(),
    },
    "get_startup_items": {
        "description": "Lists applications and services configured to launch automatically at startup/login. macOS: scans ~/Library/LaunchAgents, /Library/LaunchAgents, /Library/LaunchDaemons. Windows: reads Run registry keys. Linux: scans ~/.config/autostart. Use when the user asks what runs at startup or wants to speed up boot times.",
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: get_startup_items(),
    },
    "get_process_details": {
        "description": "Returns detailed information about a specific process by PID: executable path, command line, user, memory breakdown, open file count, and more.",
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "pid": {"type": "integer", "description": "The process ID to inspect"}
            },
            "required": ["pid"]
        },
        "fn": lambda args: get_process_details(args["pid"]),
    },
    "search_process": {
        "description": "Searches for running processes by name (case-insensitive, partial match). Returns PID, CPU%, memory%, and status for each match.",
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "name": {"type": "string", "description": "Process name to search for (e.g. 'chrome', 'python')"}
            },
            "required": ["name"]
        },
        "fn": lambda args: search_process(args["name"]),
    },
    "kill_process": {
        "description": "Terminates a process by PID. Sends SIGTERM (graceful) by default; SIGKILL if force=True. Refuses to kill critical system processes (PID 1, launchd, systemd, init, kernel_task, core Windows services). Always confirm with the user before calling.",
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "pid": {"type": "integer", "description": "The PID of the process to terminate"},
                "force": {"type": "boolean", "description": "If true, send SIGKILL (immediate). Default false (SIGTERM, graceful).", "default": False}
            },
            "required": ["pid"]
        },
        "fn": lambda args: kill_process(args["pid"], args.get("force", False)),
    },
    "get_hardware_profile": {
        "description": "Returns a full hardware profile for a given use-case: specs, live pressure, overclocking capability (where supported), upgrade feasibility per component, and workload-specific bottleneck analysis. Use this when the user asks about speeding up a specific task, upgrading their machine, or overclocking.",
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "use_case": {
                    "type": "string",
                    "description": "The user's workload or goal, e.g. 'lightroom rendering', 'gaming', 'video editing', 'compiling code'"
                }
            },
            "required": []
        },
        "fn": lambda args: get_hardware_profile(args.get("use_case", "")),
    },
    "set_reminder": {
        "description": (
            "Schedule a reminder that fires a macOS notification at the specified time. "
            "Accepts natural-language time: 'in 2 hours', 'in 30 minutes', "
            "'at 9:00 am', 'at 3pm', 'tomorrow at 8am'. "
            "Returns a reminder ID that can be used with cancel_reminder."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "message": {
                    "type": "string",
                    "description": "The reminder text to display in the notification.",
                },
                "time": {
                    "type": "string",
                    "description": (
                        "When to fire the reminder. Examples: 'in 2 hours', "
                        "'in 30 minutes', 'at 9:00 am', 'at 3pm', 'tomorrow at 8am'."
                    ),
                },
            },
            "required": ["message", "time"],
        },
        "fn": lambda args: set_reminder(args["message"], args["time"]),
    },
    "list_reminders": {
        "description": "List all pending (unfired) reminders with their IDs, messages, and scheduled fire times.",
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: list_reminders(),
    },
    "cancel_reminder": {
        "description": "Cancel a pending reminder by its ID. Get the ID from set_reminder or list_reminders.",
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "id": {
                    "type": "string",
                    "description": "The reminder ID to cancel (8-character hex string).",
                }
            },
            "required": ["id"],
        },
        "fn": lambda args: cancel_reminder(args["id"]),
    },
    "get_weather": {
        "description": (
            "Returns current weather conditions and clothing suggestions. "
            "Auto-detects location from IP if no location is provided. "
            "Pass a city name for a specific location (e.g. 'Tokyo' or 'London, UK')."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "location": {
                    "type": "string",
                    "description": "City name (e.g. 'Tokyo', 'London, UK'). Leave empty to auto-detect from IP.",
                    "default": "",
                },
                "units": {
                    "type": "string",
                    "enum": ["imperial", "metric"],
                    "description": "Temperature units: 'imperial' (°F, mph) or 'metric' (°C, km/h). Defaults to imperial.",
                    "default": "imperial",
                },
            },
            "required": [],
        },
        "fn": lambda args: get_weather(args.get("location", ""), args.get("units", "imperial")),
    },
    "check_app_updates": {
        "description": (
            "macOS only: checks for outdated applications via Homebrew (formulae + casks), "
            "the Mac App Store (requires the 'mas' CLI — install with 'brew install mas'), "
            "and macOS system software updates. Returns lists of outdated apps with current vs available versions."
        ),
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: check_app_updates(),
    },
    # ── Homebrew management ───────────────────────────────────────────────────
    "brew_list": {
        "description": (
            "List all installed Homebrew packages (formulae and/or casks). "
            "Requires allow_brew in ~/.syscontrol/config.json. macOS only."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "kind": {
                    "type": "string",
                    "description": "'formulae', 'casks', or 'all' (default).",
                    "default": "all",
                },
            },
            "required": [],
        },
        "fn": lambda args: brew_list(args.get("kind", "all")),
    },
    "brew_install": {
        "description": (
            "Install a Homebrew formula or cask. "
            "Prefix with '--cask' for cask installs (e.g. '--cask firefox'). "
            "May take up to 5 minutes for large packages. "
            "Requires allow_brew in ~/.syscontrol/config.json. macOS only."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "package": {
                    "type": "string",
                    "description": "Package name to install (e.g. 'ripgrep', '--cask firefox').",
                },
            },
            "required": ["package"],
        },
        "fn": lambda args: brew_install(args["package"]),
    },
    "brew_upgrade": {
        "description": (
            "Upgrade one or all installed Homebrew packages. "
            "Leave package empty to upgrade everything. "
            "Requires allow_brew in ~/.syscontrol/config.json. macOS only."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "package": {
                    "type": "string",
                    "description": "Package name to upgrade, or empty to upgrade all (default: all).",
                    "default": "",
                },
            },
            "required": [],
        },
        "fn": lambda args: brew_upgrade(args.get("package", "")),
    },
    "brew_uninstall": {
        "description": (
            "Uninstall a Homebrew formula or cask. "
            "Requires allow_brew in ~/.syscontrol/config.json. macOS only."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "package": {
                    "type": "string",
                    "description": "Package name to uninstall.",
                },
            },
            "required": ["package"],
        },
        "fn": lambda args: brew_uninstall(args["package"]),
    },
    "track_package": {
        "description": (
            "Track a package by tracking number. Auto-detects the carrier (UPS, USPS, FedEx, DHL). "
            "Returns current status and recent tracking history. "
            "Note: Amazon TBA numbers must be tracked at amazon.com/orders."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "tracking_number": {
                    "type": "string",
                    "description": "The package tracking number (UPS, USPS, FedEx, or DHL).",
                }
            },
            "required": ["tracking_number"],
        },
        "fn": lambda args: track_package(args["tracking_number"]),
    },
    "find_large_files": {
        "description": (
            "Finds the top N largest files under a given directory path (default: home directory). "
            "Skips hidden directories, .git, __pycache__, node_modules, .venv, and Library. "
            "Use when the user asks what is using disk space or wants to free up storage."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Directory path to search (e.g. '/Users/you/Downloads'). Defaults to home directory if omitted.",
                    "default": "",
                },
                "n": {
                    "type": "integer",
                    "description": "Number of largest files to return (default 10, max 50).",
                    "default": 10,
                },
            },
            "required": [],
        },
        "fn": lambda args: find_large_files(args.get("path", ""), args.get("n", 10)),
    },
    "network_latency_check": {
        "description": (
            "Pings the local gateway, Cloudflare DNS (1.1.1.1), and Google DNS (8.8.8.8) "
            "concurrently and returns per-target latency and reachability. "
            "Includes an automatic diagnosis (router issue / ISP issue / congestion / normal). "
            "Use when the user asks if their internet is slow or to locate where latency is introduced."
        ),
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: network_latency_check(),
    },
    "get_docker_status": {
        "description": (
            "Returns all running Docker containers with their CPU%, memory usage, image, status, and ports. "
            "Also reports total container count (including stopped). "
            "Returns an actionable error if Docker is not installed or the daemon is not running."
        ),
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: get_docker_status(),
    },
    "get_time_machine_status": {
        "description": (
            "macOS only. Returns Time Machine backup status: whether a backup is currently running, "
            "last backup time and how long ago it was, backup destination name and kind. "
            "Uses tmutil status, latestbackup, and destinationinfo (run in parallel)."
        ),
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: get_time_machine_status(),
    },
    "tail_system_logs": {
        "description": (
            "Returns the last N lines from the system log. "
            "macOS: reads from the unified system log (last 5 minutes) via `log show`. "
            "Linux: reads from journalctl or /var/log/syslog. "
            "Optional filter_str narrows results to lines containing that keyword. "
            "Use to diagnose crashes, kernel panics, or application errors."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "lines": {
                    "type": "integer",
                    "description": "Number of log lines to return (default 50, max 500).",
                    "default": 50,
                },
                "filter_str": {
                    "type": "string",
                    "description": "Optional keyword to filter log lines (case-insensitive).",
                    "default": "",
                },
            },
            "required": [],
        },
        "fn": lambda args: tail_system_logs(args.get("lines", 50), args.get("filter_str", "")),
    },
    # ── Browser / Web tools ──────────────────────────────────────────────────
    "web_fetch": {
        "description": (
            "Fetch the plain-text content of any public web page. "
            "HTML is stripped. No browser needed, no permission required. "
            "Use this to read articles, docs, pricing pages, or any URL the user mentions."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "url": {"type": "string", "description": "Full URL to fetch (https:// assumed if omitted)."},
                "max_chars": {
                    "type": "integer", "default": 8000,
                    "description": "Max characters of plain text to return (500–32000).",
                },
            },
            "required": ["url"],
        },
        "fn": lambda args: web_fetch(args["url"], args.get("max_chars", 8000)),
    },
    "web_search": {
        "description": (
            "Search the web (DuckDuckGo) and return the top N results "
            "(title, URL, snippet). No API key. No browser permission required. "
            "Combine with web_fetch to read the full content of a result."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "Search query string."},
                "num_results": {
                    "type": "integer", "default": 5,
                    "description": "Number of results to return (1–10).",
                },
            },
            "required": ["query"],
        },
        "fn": lambda args: web_search(args["query"], args.get("num_results", 5)),
    },
    "grant_browser_access": {
        "description": (
            "Grants the agent permission to control the user's browser. "
            "ONLY call this tool after the user has explicitly said yes/granted/allow. "
            "Writes a permission flag; subsequent browser_* calls will then work."
        ),
        "parallel": False,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: grant_browser_access(),
    },
    "browser_open_url": {
        "description": (
            "Open a URL in the user's default browser as a new tab/window. "
            "Requires prior browser permission (grant_browser_access). "
            "macOS: uses `open` command. Linux/Windows: uses webbrowser module."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "url": {"type": "string", "description": "URL to open (https:// assumed if omitted)."},
            },
            "required": ["url"],
        },
        "fn": lambda args: browser_open_url(args["url"]),
    },
    "browser_navigate": {
        "description": (
            "Navigate the currently active browser tab to a different URL via AppleScript. "
            "macOS only (falls back to browser_open_url on other platforms). "
            "Requires browser permission. Supports Arc, Chrome, Brave, Edge, Safari."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "url": {"type": "string", "description": "URL to navigate to."},
            },
            "required": ["url"],
        },
        "fn": lambda args: browser_navigate(args["url"]),
    },
    "browser_get_page": {
        "description": (
            "Return the URL, title, and visible text of the currently active browser tab "
            "via AppleScript (macOS only). "
            "Requires browser permission. Use this to read what the user is currently looking at, "
            "summarise a page, or answer questions about its content."
        ),
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: browser_get_page(),
    },
    # ── iMessage ──────────────────────────────────────────────────────────────
    "send_imessage": {
        "description": (
            "Send an iMessage or SMS via macOS Messages.app. "
            "Accepts a phone number (e.g. '+14155551234') or Apple ID email. "
            "Requires Messages.app to be signed in and Terminal Automation permission. "
            "macOS only."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "recipient": {
                    "type": "string",
                    "description": "Phone number (e.g. '+14155551234') or Apple ID email of the recipient.",
                },
                "message": {
                    "type": "string",
                    "description": "Text message content to send.",
                },
            },
            "required": ["recipient", "message"],
        },
        "fn": lambda args: send_imessage(args["recipient"], args["message"]),
    },
    "get_imessage_history": {
        "description": (
            "Return recent iMessage/SMS messages matching a contact name, phone number, or email. "
            "Reads from ~/Library/Messages/chat.db. "
            "Requires Full Disk Access for Terminal in System Settings → Privacy & Security. "
            "macOS only."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "contact": {
                    "type": "string",
                    "description": "Name, phone, or email to filter messages (partial match).",
                },
                "limit": {
                    "type": "integer",
                    "description": "Max number of messages to return (default 20, max 200).",
                    "default": 20,
                },
            },
            "required": ["contact"],
        },
        "fn": lambda args: get_imessage_history(args["contact"], args.get("limit", 20)),
    },
    # ── Email ─────────────────────────────────────────────────────────────────
    "read_emails": {
        "description": (
            "Return recent emails from a Mail.app mailbox. "
            "Returns subject, sender, date, and a short preview for each message. "
            "Requires allow_email in ~/.syscontrol/config.json. macOS only."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "folder": {
                    "type": "string",
                    "description": "Mailbox name (default: 'INBOX').",
                    "default": "INBOX",
                },
                "count": {
                    "type": "integer",
                    "description": "Maximum messages to return (1–50, default 10).",
                    "default": 10,
                },
            },
            "required": [],
        },
        "fn": lambda args: read_emails(args.get("folder", "INBOX"), args.get("count", 10)),
    },
    "send_email": {
        "description": (
            "Send an email via Mail.app. "
            "Requires a configured sending account in Mail.app and allow_email in config. "
            "macOS only."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "to":      {"type": "string", "description": "Recipient email address."},
                "subject": {"type": "string", "description": "Email subject line."},
                "body":    {"type": "string", "description": "Plain-text message body."},
                "cc": {
                    "type": "string",
                    "description": "Optional CC email address.",
                    "default": "",
                },
            },
            "required": ["to", "subject", "body"],
        },
        "fn": lambda args: send_email(
            args["to"], args["subject"], args["body"], args.get("cc", ""),
        ),
    },
    "search_emails": {
        "description": (
            "Search emails across all Mail.app accounts and mailboxes. "
            "Matches against subject, sender, and body text. "
            "Requires allow_email in ~/.syscontrol/config.json. macOS only."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "Search term (subject, sender name, or body text).",
                },
                "count": {
                    "type": "integer",
                    "description": "Maximum results (1–50, default 20).",
                    "default": 20,
                },
            },
            "required": ["query"],
        },
        "fn": lambda args: search_emails(args["query"], args.get("count", 20)),
    },
    # ── Clipboard ─────────────────────────────────────────────────────────────
    "get_clipboard": {
        "description": (
            "Return the current text content of the system clipboard. "
            "macOS only (uses pbpaste)."
        ),
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: get_clipboard(),
    },
    "set_clipboard": {
        "description": (
            "Write text to the system clipboard. "
            "macOS only (uses pbcopy). "
            "Use to copy a result or command output so the user can paste it anywhere."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "text": {
                    "type": "string",
                    "description": "Text to place on the clipboard.",
                },
            },
            "required": ["text"],
        },
        "fn": lambda args: set_clipboard(args["text"]),
    },
    # ── Screenshot ────────────────────────────────────────────────────────────
    "take_screenshot": {
        "description": (
            "Capture a screenshot of the entire screen and return it as an inline image. "
            "Optionally saves to a file path. "
            "macOS only (uses screencapture -x, no shutter sound)."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Optional file path to save the PNG (e.g. '~/Desktop/screenshot.png'). Leave empty to skip saving.",
                    "default": "",
                },
            },
            "required": [],
        },
        "fn": lambda args: take_screenshot(args.get("path", "")),
    },
    # ── App Control ───────────────────────────────────────────────────────────
    "open_app": {
        "description": (
            "Open an application by name on macOS (uses 'open -a'). "
            "Works with any installed app, e.g. 'Calculator', 'Safari', 'Spotify'."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "name": {
                    "type": "string",
                    "description": "Application name as it appears in /Applications (e.g. 'Calculator', 'Spotify').",
                },
            },
            "required": ["name"],
        },
        "fn": lambda args: open_app(args["name"]),
    },
    "quit_app": {
        "description": (
            "Quit an application gracefully by name using AppleScript ('tell app to quit'). "
            "Pass force=true for immediate SIGKILL. macOS only."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "name": {
                    "type": "string",
                    "description": "Application name to quit (e.g. 'Safari', 'Spotify').",
                },
                "force": {
                    "type": "boolean",
                    "description": "If true, force-kill the process immediately (SIGKILL). Default false.",
                    "default": False,
                },
            },
            "required": ["name"],
        },
        "fn": lambda args: quit_app(args["name"], args.get("force", False)),
    },
    # ── Volume ────────────────────────────────────────────────────────────────
    "get_volume": {
        "description": (
            "Return the current macOS output volume level (0–100), input volume, alert volume, and mute state. "
            "macOS only."
        ),
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: get_volume(),
    },
    "set_volume": {
        "description": (
            "Set the macOS system output volume to a level between 0 (mute) and 100 (maximum). "
            "macOS only."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "level": {
                    "type": "integer",
                    "description": "Output volume level (0–100).",
                    "minimum": 0,
                    "maximum": 100,
                },
            },
            "required": ["level"],
        },
        "fn": lambda args: set_volume(args["level"]),
    },
    # ── Media ─────────────────────────────────────────────────────────────────
    "get_now_playing": {
        "description": (
            "Return the currently playing track in Music.app or Spotify. "
            "Shows track name, artist, album, playback position, and duration. "
            "Returns 'nothing_playing' if no media app is active. macOS only."
        ),
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: get_now_playing(),
    },
    "media_control": {
        "description": (
            "Control media playback in Music.app or Spotify. "
            "Actions: play, pause, play_pause, next, previous, stop. "
            "Auto-detects the active player if app is not specified. macOS only."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "action": {
                    "type": "string",
                    "description": "Playback action: 'play', 'pause', 'play_pause', 'next', 'previous', or 'stop'.",
                    "enum": ["play", "pause", "play_pause", "next", "previous", "stop"],
                },
                "app": {
                    "type": "string",
                    "description": "Target app: 'Music' or 'Spotify' (default: auto-detect).",
                    "default": "",
                },
            },
            "required": ["action"],
        },
        "fn": lambda args: media_control(args["action"], args.get("app", "")),
    },
    # ── Wi-Fi ─────────────────────────────────────────────────────────────────
    "get_wifi_networks": {
        "description": (
            "Scan for nearby Wi-Fi networks and return each network's SSID, BSSID, "
            "signal strength (RSSI in dBm), channel, and security type. "
            "Sorted strongest-first. macOS only (uses airport utility)."
        ),
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: get_wifi_networks(),
    },
    # ── File I/O ──────────────────────────────────────────────────────────────
    "read_file": {
        "description": (
            "Read the text contents of a file at the given path. "
            "Returns up to max_chars characters (default 16,000, max 32,000). "
            "Useful for reading config files, logs, scripts, notes, etc."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Absolute or home-relative path to the file (e.g. '~/.zshrc', '/etc/hosts').",
                },
                "max_chars": {
                    "type": "integer",
                    "description": "Maximum characters to return (default 16000, max 32000).",
                    "default": 16000,
                },
            },
            "required": ["path"],
        },
        "fn": lambda args: read_file(args["path"], args.get("max_chars", 16000)),
    },
    "write_file": {
        "description": (
            "Write text content to a file at the given path. "
            "Creates parent directories as needed. Overwrites by default. "
            "Use for saving notes, configs, scripts, or any text output."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Absolute or home-relative path where the file should be written.",
                },
                "content": {
                    "type": "string",
                    "description": "Text content to write.",
                },
                "overwrite": {
                    "type": "boolean",
                    "description": "If false, returns an error if the file already exists. Default true.",
                    "default": True,
                },
            },
            "required": ["path", "content"],
        },
        "fn": lambda args: write_file(args["path"], args["content"], args.get("overwrite", True)),
    },
    "read_file_lines": {
        "description": (
            "Read a file with line numbers, supporting offset and limit for "
            "navigating large files. Returns content formatted like 'cat -n'. "
            "Use this instead of read_file when you need line numbers or "
            "want to read a specific section. "
            "Requires allow_file_read in ~/.syscontrol/config.json."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Absolute or home-relative path to the file.",
                },
                "offset": {
                    "type": "integer",
                    "description": "1-based line number to start reading from (default 1).",
                    "default": 1,
                },
                "limit": {
                    "type": "integer",
                    "description": "Maximum number of lines to return (default 200, max 2000).",
                    "default": 200,
                },
            },
            "required": ["path"],
        },
        "fn": lambda args: read_file_lines(
            args["path"], args.get("offset", 1), args.get("limit", 200),
        ),
    },
    "edit_file": {
        "description": (
            "Make targeted edits to a file using find-and-replace. "
            "Provide the exact text to find (old_string) and its replacement (new_string). "
            "By default, old_string must appear exactly once (fails if ambiguous). "
            "Set replace_all=true to replace all occurrences. "
            "ALWAYS read the file first to get the exact text to replace. "
            "Requires allow_file_write in ~/.syscontrol/config.json."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Absolute or home-relative path to the file to edit.",
                },
                "old_string": {
                    "type": "string",
                    "description": "Exact text to find in the file. Must be unique unless replace_all=true.",
                },
                "new_string": {
                    "type": "string",
                    "description": "Replacement text. Use empty string to delete the matched text.",
                },
                "replace_all": {
                    "type": "boolean",
                    "description": "Replace all occurrences instead of requiring uniqueness (default false).",
                    "default": False,
                },
            },
            "required": ["path", "old_string", "new_string"],
        },
        "fn": lambda args: edit_file(
            args["path"], args["old_string"], args["new_string"],
            args.get("replace_all", False),
        ),
    },
    "glob_files": {
        "description": (
            "Find files matching a glob pattern (e.g. '**/*.py', 'src/**/*.ts'). "
            "Searches recursively from the given base path. "
            "Skips .git, node_modules, .venv, __pycache__ directories."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "pattern": {
                    "type": "string",
                    "description": "Glob pattern (e.g. '**/*.py', '*.json', 'src/**/*.ts').",
                },
                "path": {
                    "type": "string",
                    "description": "Base directory to search from (default: current directory).",
                    "default": ".",
                },
            },
            "required": ["pattern"],
        },
        "fn": lambda args: glob_files(args["pattern"], args.get("path", ".")),
    },
    "grep_files": {
        "description": (
            "Search file contents for a regex pattern, like grep/ripgrep. "
            "Returns matching lines with file paths and line numbers. "
            "Optionally filter by file glob (e.g. '*.py') and include context lines. "
            "Skips binary files and .git, node_modules, .venv, __pycache__."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "pattern": {
                    "type": "string",
                    "description": "Regex pattern to search for (Python re syntax).",
                },
                "path": {
                    "type": "string",
                    "description": "File or directory to search in (default: current directory).",
                    "default": ".",
                },
                "include": {
                    "type": "string",
                    "description": "Glob filter for file names (e.g. '*.py', '*.ts'). Empty = all files.",
                    "default": "",
                },
                "max_results": {
                    "type": "integer",
                    "description": "Maximum matching lines to return (default 50, max 500).",
                    "default": 50,
                },
                "context_lines": {
                    "type": "integer",
                    "description": "Lines of context before and after each match (default 0, max 10).",
                    "default": 0,
                },
            },
            "required": ["pattern"],
        },
        "fn": lambda args: grep_files(
            args["pattern"], args.get("path", "."), args.get("include", ""),
            args.get("max_results", 50), args.get("context_lines", 0),
        ),
    },
    "git_status": {
        "description": (
            "Show git repository status: current branch, staged/unstaged/untracked "
            "files, and the 5 most recent commits. Auto-detects the repo root."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Path inside the git repository (default: current directory).",
                    "default": ".",
                },
            },
            "required": [],
        },
        "fn": lambda args: git_status(args.get("path", ".")),
    },
    "git_diff": {
        "description": (
            "Show git diff output for a repository. By default shows unstaged changes; "
            "set staged=true for staged changes. Optionally scope to a specific file path."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Path inside the repo, or a specific file (default: repo root).",
                    "default": ".",
                },
                "staged": {
                    "type": "boolean",
                    "description": "Show staged (cached) changes instead of unstaged (default false).",
                    "default": False,
                },
            },
            "required": [],
        },
        "fn": lambda args: git_diff(args.get("path", "."), args.get("staged", False)),
    },
    "read_spreadsheet": {
        "description": (
            "Read rows and cells from a spreadsheet file (.xlsx or .csv). "
            "Returns headers and row data up to max_rows (default 200). "
            "For .xlsx files, optionally specify a sheet name and/or cell range (e.g. 'A1:D10'). "
            "Requires allow_file_read in ~/.syscontrol/config.json."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Path to the .xlsx or .csv file (e.g. '~/Desktop/budget.xlsx').",
                },
                "sheet": {
                    "type": "string",
                    "description": "Sheet name to read (.xlsx only). Defaults to the active sheet.",
                },
                "cell_range": {
                    "type": "string",
                    "description": "Cell range in A1 notation (e.g. 'A1:D10'). Omit to read all rows.",
                },
                "max_rows": {
                    "type": "integer",
                    "description": "Maximum rows to return (1–200, default 200).",
                    "default": 200,
                },
            },
            "required": ["path"],
        },
        "fn": lambda args: read_spreadsheet(
            args["path"],
            args.get("sheet"),
            args.get("cell_range"),
            args.get("max_rows", 200),
        ),
    },
    "edit_spreadsheet": {
        "description": (
            "Write cells or append rows to a spreadsheet file (.xlsx or .csv). "
            "'updates' is a list of {cell, value} objects using A1 notation (e.g. {\"cell\": \"B3\", \"value\": 500}). "
            "'append_rows' is a list of rows (each row is a list of values) to add at the end. "
            "Set create_if_missing=true to create the file if it does not exist (.xlsx only). "
            "Requires allow_file_write in ~/.syscontrol/config.json."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Path to the .xlsx or .csv file.",
                },
                "sheet": {
                    "type": "string",
                    "description": "Sheet name to edit (.xlsx only). Defaults to the active sheet.",
                },
                "updates": {
                    "type": "array",
                    "description": "List of cell updates: [{\"cell\": \"A1\", \"value\": \"Hello\"}, ...].",
                    "items": {
                        "type": "object",
                        "properties": {
                            "cell": {"type": "string"},
                            "value": {},
                        },
                        "required": ["cell", "value"],
                    },
                },
                "append_rows": {
                    "type": "array",
                    "description": "Rows to append at the end of the sheet: [[\"Q1\", 1000], [\"Q2\", 2000]].",
                    "items": {"type": "array"},
                },
                "create_if_missing": {
                    "type": "boolean",
                    "description": "Create the file if it does not exist (.xlsx only). Default false.",
                    "default": False,
                },
            },
            "required": ["path"],
        },
        "fn": lambda args: edit_spreadsheet(
            args["path"],
            args.get("sheet"),
            args.get("updates"),
            args.get("append_rows"),
            args.get("create_if_missing", False),
        ),
    },
    "read_document": {
        "description": (
            "Read paragraphs from a Word document (.docx) or plain text file (.txt, .md). "
            "Returns a list of non-empty paragraphs with their index, and a total word count. "
            "Requires allow_file_read in ~/.syscontrol/config.json."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Path to the .docx, .txt, or .md file.",
                },
            },
            "required": ["path"],
        },
        "fn": lambda args: read_document(args["path"]),
    },
    "edit_document": {
        "description": (
            "Edit a Word document (.docx): find-and-replace text, overwrite a paragraph by index, "
            "or append new paragraphs at the end. "
            "'replacements' is a list of {find, replace} objects. "
            "'set_paragraph' is {index, text} to overwrite one paragraph. "
            "'append_paragraphs' is a list of strings to add at the end. "
            "Requires allow_file_write in ~/.syscontrol/config.json."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Path to the .docx file.",
                },
                "replacements": {
                    "type": "array",
                    "description": "Find-and-replace pairs: [{\"find\": \"old text\", \"replace\": \"new text\"}, ...].",
                    "items": {
                        "type": "object",
                        "properties": {
                            "find": {"type": "string"},
                            "replace": {"type": "string"},
                        },
                        "required": ["find", "replace"],
                    },
                },
                "set_paragraph": {
                    "type": "object",
                    "description": "Overwrite a paragraph by index: {\"index\": 0, \"text\": \"New text\"}.",
                    "properties": {
                        "index": {"type": "integer"},
                        "text": {"type": "string"},
                    },
                    "required": ["index", "text"],
                },
                "append_paragraphs": {
                    "type": "array",
                    "description": "List of strings to append as new paragraphs at the end of the document.",
                    "items": {"type": "string"},
                },
            },
            "required": ["path"],
        },
        "fn": lambda args: edit_document(
            args["path"],
            args.get("replacements"),
            args.get("append_paragraphs"),
            args.get("set_paragraph"),
        ),
    },
    "read_pdf": {
        "description": (
            "Extract text from a PDF file, returned page by page. "
            "Use this when a user shares or mentions a PDF and wants you to read, summarise, "
            "or answer questions from its contents. "
            "Supports up to 200 pages per call (default 50); pass max_pages for larger documents. "
            "Encrypted/password-protected PDFs are not supported. "
            "Requires allow_file_read in ~/.syscontrol/config.json."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Absolute or home-relative path to the .pdf file.",
                },
                "max_pages": {
                    "type": "integer",
                    "description": f"Maximum number of pages to read (1-{_MAX_PDF_PAGES}, default {_DEFAULT_PDF_PAGES}).",
                    "default": _DEFAULT_PDF_PAGES,
                },
            },
            "required": ["path"],
        },
        "fn": lambda args: read_pdf(args["path"], args.get("max_pages", _DEFAULT_PDF_PAGES)),
    },
    # ── File management ───────────────────────────────────────────────────────
    "list_directory": {
        "description": (
            "List the contents of a directory. Returns each entry's name, type (file/dir), "
            "size, and modification time. Use to browse the filesystem or verify file presence. "
            "Requires allow_file_read in ~/.syscontrol/config.json."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Directory path to list (default: home directory).",
                    "default": "",
                },
                "show_hidden": {
                    "type": "boolean",
                    "description": "Include entries starting with '.' (default false).",
                    "default": False,
                },
            },
            "required": [],
        },
        "fn": lambda args: list_directory(args.get("path", ""), args.get("show_hidden", False)),
    },
    "move_file": {
        "description": (
            "Move or rename a file or directory. "
            "Use for renaming files, moving files to different folders, or reorganising directories. "
            "Requires allow_file_write in ~/.syscontrol/config.json."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "src": {"type": "string", "description": "Source path (absolute or home-relative)."},
                "dst": {"type": "string", "description": "Destination path or target directory."},
            },
            "required": ["src", "dst"],
        },
        "fn": lambda args: move_file(args["src"], args["dst"]),
    },
    "copy_file": {
        "description": (
            "Copy a file to a new location. "
            "If dst is a directory, the file is copied into it preserving its name. "
            "Requires allow_file_write in ~/.syscontrol/config.json."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "src":       {"type": "string", "description": "Source file path."},
                "dst":       {"type": "string", "description": "Destination file path or directory."},
                "overwrite": {
                    "type": "boolean",
                    "description": "Replace destination if it exists (default false).",
                    "default": False,
                },
            },
            "required": ["src", "dst"],
        },
        "fn": lambda args: copy_file(args["src"], args["dst"], args.get("overwrite", False)),
    },
    "delete_file": {
        "description": (
            "Delete a file or directory. "
            "On macOS, moves to Trash by default (recoverable); pass permanent=true to permanently delete. "
            "Safety: only allows deletion within the home directory. "
            "Requires allow_file_write in ~/.syscontrol/config.json."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "path":      {"type": "string", "description": "Path to the file or directory to remove."},
                "permanent": {
                    "type": "boolean",
                    "description": "Permanently delete instead of moving to Trash (default false).",
                    "default": False,
                },
            },
            "required": ["path"],
        },
        "fn": lambda args: delete_file(args["path"], args.get("permanent", False)),
    },
    "create_directory": {
        "description": (
            "Create a new directory (and any missing parent directories). "
            "Requires allow_file_write in ~/.syscontrol/config.json."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "Absolute or home-relative path for the new directory."},
            },
            "required": ["path"],
        },
        "fn": lambda args: create_directory(args["path"]),
    },
    "search_files": {
        "description": (
            "Search for files using macOS Spotlight (mdfind). "
            "Much faster than find — searches the entire system instantly. "
            "Use for finding files by name, content keywords, or type. "
            "macOS only."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "File name fragment or Spotlight query (e.g. 'budget', 'author:Alice').",
                },
                "scope": {
                    "type": "string",
                    "description": "Directory to restrict the search to (default: home directory).",
                    "default": "~",
                },
                "kind": {
                    "type": "string",
                    "description": (
                        "Optional type filter: 'pdf', 'image', 'video', 'audio', "
                        "'folder', 'application', 'code', or 'document'."
                    ),
                    "default": "",
                },
                "limit": {
                    "type": "integer",
                    "description": "Maximum results to return (1–200, default 50).",
                    "default": 50,
                },
            },
            "required": ["query"],
        },
        "fn": lambda args: search_files(
            args["query"],
            args.get("scope", "~"),
            args.get("kind", ""),
            args.get("limit", 50),
        ),
    },
    # ── Shell ─────────────────────────────────────────────────────────────────
    "run_shell_command": {
        "description": (
            "Execute an arbitrary shell (bash) command and return stdout, stderr, and exit code. "
            "DISABLED by default for safety. Enable by adding {\"allow_shell\": true} to ~/.syscontrol/config.json. "
            "Timeout is 30s by default (max 120s). "
            "Always confirm with the user before running destructive commands."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "command": {
                    "type": "string",
                    "description": "Bash command to run (e.g. 'ls -la ~/Desktop', 'git log -5').",
                },
                "timeout": {
                    "type": "integer",
                    "description": "Timeout in seconds (1–120, default 30).",
                    "default": 30,
                },
            },
            "required": ["command"],
        },
        "fn": lambda args: run_shell_command(args["command"], args.get("timeout", 30)),
    },
    # ── Calendar & Contacts ───────────────────────────────────────────────────
    "get_calendar_events": {
        "description": (
            "Return upcoming calendar events from macOS Calendar.app for the next N days. "
            "Includes title, calendar name, start/end time, and location. "
            "Requires Calendar access for Terminal in System Settings → Privacy & Security. "
            "macOS only."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "lookahead_days": {
                    "type": "integer",
                    "description": "Number of days ahead to look for events (1–90, default 7).",
                    "default": 7,
                },
            },
            "required": [],
        },
        "fn": lambda args: get_calendar_events(args.get("lookahead_days", 7)),
    },
    "get_contact": {
        "description": (
            "Search macOS Contacts.app for a person by name (partial match) "
            "and return their phone numbers and email addresses. "
            "Requires Contacts access for Terminal in System Settings → Privacy & Security. "
            "macOS only."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "name": {
                    "type": "string",
                    "description": "Name to search for (e.g. 'John', 'Appleseed'). Case-insensitive partial match.",
                },
            },
            "required": ["name"],
        },
        "fn": lambda args: get_contact(args["name"]),
    },
    # ── Notes ─────────────────────────────────────────────────────────────────
    "list_notes": {
        "description": (
            "List notes from macOS Notes.app. Returns title, folder, and timestamps. "
            "Optionally filter by folder name. "
            "Requires allow_notes in ~/.syscontrol/config.json. macOS only."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "folder": {
                    "type": "string",
                    "description": "Folder name to filter by (default: all folders).",
                    "default": "",
                },
                "count": {
                    "type": "integer",
                    "description": "Maximum notes to return (1–100, default 20).",
                    "default": 20,
                },
            },
            "required": [],
        },
        "fn": lambda args: list_notes(args.get("folder", ""), args.get("count", 20)),
    },
    "read_note": {
        "description": (
            "Read the full body of a note from macOS Notes.app by title (partial match). "
            "Requires allow_notes in ~/.syscontrol/config.json. macOS only."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "name": {
                    "type": "string",
                    "description": "Full or partial note title to search for (case-insensitive).",
                },
            },
            "required": ["name"],
        },
        "fn": lambda args: read_note(args["name"]),
    },
    "create_note": {
        "description": (
            "Create a new note in macOS Notes.app with a title and body. "
            "Optionally specify a folder. "
            "Requires allow_notes in ~/.syscontrol/config.json. macOS only."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "title":  {"type": "string", "description": "Note title."},
                "body":   {"type": "string", "description": "Note body text."},
                "folder": {
                    "type": "string",
                    "description": "Folder to create the note in (default: default Notes folder).",
                    "default": "",
                },
            },
            "required": ["title", "body"],
        },
        "fn": lambda args: create_note(args["title"], args["body"], args.get("folder", "")),
    },
    # ── Shortcuts & System ────────────────────────────────────────────────────
    "run_shortcut": {
        "description": (
            "Run a named macOS Shortcut from Shortcuts.app via the shortcuts CLI. "
            "Shortcut name is case-sensitive. Optionally pass input_text as stdin. "
            "Requires macOS 12+. "
            "Use to trigger user-defined automations (e.g. 'Send Daily Report', 'Resize Images')."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "name": {
                    "type": "string",
                    "description": "Exact name of the Shortcut to run (case-sensitive).",
                },
                "input_text": {
                    "type": "string",
                    "description": "Optional text input to pass to the Shortcut via stdin.",
                    "default": "",
                },
            },
            "required": ["name"],
        },
        "fn": lambda args: run_shortcut(args["name"], args.get("input_text", "")),
    },
    "get_frontmost_app": {
        "description": (
            "Return the name of the macOS application currently in focus (frontmost window). "
            "Requires Accessibility permission for Terminal in System Settings → Privacy & Security. "
            "macOS only."
        ),
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: get_frontmost_app(),
    },
    "toggle_do_not_disturb": {
        "description": (
            "Enable or disable macOS Focus / Do Not Disturb mode. "
            "Tries built-in Shortcut names: 'Turn On/Off Do Not Disturb' and 'Turn On/Off Focus'. "
            "If those don't exist, returns an error with setup instructions. "
            "macOS 12+ required."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "enabled": {
                    "type": "boolean",
                    "description": "True to enable Do Not Disturb / Focus, false to disable.",
                },
            },
            "required": ["enabled"],
        },
        "fn": lambda args: toggle_do_not_disturb(args["enabled"]),
    },
    # ── Disk ──────────────────────────────────────────────────────────────────
    "eject_disk": {
        "description": (
            "Unmount and eject an external disk by its mountpoint (e.g. '/Volumes/MyDrive'). "
            "Uses diskutil eject. macOS only. "
            "Use get_disk_usage to find available mountpoints."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "mountpoint": {
                    "type": "string",
                    "description": "Disk mountpoint path (e.g. '/Volumes/BackupDrive').",
                },
            },
            "required": ["mountpoint"],
        },
        "fn": lambda args: eject_disk(args["mountpoint"]),
    },
    # ── Tool self-extension ────────────────────────────────────────────────────
    "list_user_tools": {
        "description": "Lists all custom tools installed via create_tool.",
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: list_user_tools(),
    },
    "create_tool": {
        "description": (
            "Generates, validates, and installs a new MCP tool permanently into the server. "
            "Requires allow_tool_creation: true in ~/.syscontrol/config.json. "
            "The tool is available after restarting the agent."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "name": {
                    "type": "string",
                    "description": "snake_case tool identifier, e.g. 'get_spotify_track'.",
                },
                "description": {
                    "type": "string",
                    "description": "One-sentence description of what the tool does.",
                },
                "parameters_schema": {
                    "type": "object",
                    "description": "JSON Schema object for tool inputs. Omit for no-arg tools.",
                },
                "implementation": {
                    "type": "string",
                    "description": (
                        "Complete Python function definition(s). "
                        "stdlib is available; add 'import X' inside the function for extras."
                    ),
                },
                "prompt_doc": {
                    "type": "string",
                    "description": "Optional usage notes to insert into the system prompt.",
                },
            },
            "required": ["name", "description", "implementation"],
        },
        "fn": lambda args: create_tool(
            args.get("name", ""),
            args.get("description", ""),
            args.get("parameters_schema"),
            args.get("implementation", ""),
            args.get("prompt_doc", ""),
        ),
    },
    # ── Memory tools ───────────────────────────────────────────────────────────
    "read_memory": {
        "description": (
            "Read the persistent memory file that stores notes and key facts from past sessions. "
            "Call this when the user references something from a previous session, asks what you remember, "
            "or when prior context seems relevant to their request."
        ),
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: read_memory(),
    },
    "append_memory_note": {
        "description": (
            "Append a concise note or key fact to the persistent memory file. "
            "Use this proactively when the user shares a preference, important system fact, or decision "
            "that would be useful to recall in a future session. Keep notes brief (1-3 sentences)."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "note": {
                    "type": "string",
                    "description": "The note to save. Be concise — capture the key fact, not the full conversation.",
                },
            },
            "required": ["note"],
        },
        "fn": lambda args: append_memory_note(args.get("note", "")),
    },
    # ── Deep Research ─────────────────────────────────────────────────────────
    "deep_research": {
        "description": (
            "Conduct deep, multi-step web research on a topic. "
            "Plans subquestions, searches multiple sources, extracts claims, "
            "cross-verifies facts, and returns a citation-backed answer. "
            "Use for questions needing current information, evidence verification, "
            "or multiple perspectives. Takes 1-3 minutes. "
            "Requires allow_deep_research in ~/.syscontrol/config.json."
        ),
        "parallel": False,
        "inputSchema": {
            "type": "object",
            "properties": {
                "question": {
                    "type": "string",
                    "description": "The research question to investigate thoroughly.",
                },
                "max_sources": {
                    "type": "integer",
                    "description": "Maximum sources to consult (3-25, default 15).",
                    "default": 15,
                },
                "max_loops": {
                    "type": "integer",
                    "description": "Maximum research iterations (1-10, default 5).",
                    "default": 5,
                },
            },
            "required": ["question"],
        },
        "fn": lambda args: _run_deep_research(
            args["question"],
            args.get("max_sources", 15),
            args.get("max_loops", 5),
        ),
    },
    # ── Sub-agent orchestration ────────────────────────────────────────────────
    "list_agents": {
        "description": (
            "List all available sub-agents with their names and descriptions. "
            "Call this before run_agent if you are unsure which agent to delegate to."
        ),
        "parallel": True,
        "inputSchema": {"type": "object", "properties": {}, "required": []},
        "fn": lambda _: _list_agents(),
    },
    "run_agent": {
        "description": (
            "Delegate a focused task to a named sub-agent that runs in an isolated "
            "context with a restricted tool set. The sub-agent completes the task "
            "independently and returns only its final answer. "
            "Use list_agents to discover available agents. "
            "Do not delegate tasks you can complete directly with other tools. "
            "Requires allow_agents in ~/.syscontrol/config.json."
        ),
        "parallel": True,
        "inputSchema": {
            "type": "object",
            "properties": {
                "agent_name": {
                    "type": "string",
                    "description": (
                        "Name of the sub-agent (e.g. 'explorer', 'analyst', "
                        "'researcher', 'writer'). Use list_agents to see all options."
                    ),
                },
                "task": {
                    "type": "string",
                    "description": (
                        "Clear, self-contained description of the task. "
                        "Include all context the sub-agent needs — it has no "
                        "access to the parent conversation history."
                    ),
                },
            },
            "required": ["agent_name", "task"],
        },
        "fn": lambda args: _run_agent(
            args["agent_name"],
            args["task"],
        ),
    },
    # ── User-Defined Tools (registry) ──────────────────────────────────────────
    # (entries inserted here by create_tool — do not remove this comment)
}


# ── MCP request dispatcher ────────────────────────────────────────────────────

def _handle_tools_call(id_: int | None, params: dict) -> dict:
    """Execute a ``tools/call`` request and return the JSON-RPC response."""
    tool_name = params.get("name")
    args = params.get("arguments", {})
    if tool_name not in TOOLS:
        return make_error(id_, -32601, f"Unknown tool: {tool_name}")
    try:
        result = TOOLS[tool_name]["fn"](args)
        if isinstance(result, tuple):
            data, img_b64 = result
            content = [
                {"type": "text", "text": json.dumps(data, indent=2)},
                {"type": "image", "data": img_b64, "mimeType": "image/png"},
            ]
        else:
            content = [{"type": "text", "text": json.dumps(result, indent=2)}]
        return {"jsonrpc": "2.0", "id": id_, "result": {"content": content}}
    except Exception as e:
        return make_error(id_, -32603, str(e))


def handle_request(request: dict) -> dict | None:
    """Dispatch a JSON-RPC request to the appropriate handler.

    Supports ``initialize``, ``tools/list``, ``tools/call``, and ``ping``.
    Returns ``None`` for notifications (requests without an ``id``).
    """
    method = request.get("method")
    id_ = request.get("id")
    params = request.get("params", {})

    # Notifications have no "id" — must never be responded to
    if "id" not in request:
        return None

    if method == "initialize":
        return {
            "jsonrpc": "2.0",
            "id": id_,
            "result": {
                "protocolVersion": "2024-11-05",
                "capabilities": {"tools": {}},
                "serverInfo": {"name": "system-monitor", "version": "1.0.0"},
            }
        }

    if method == "tools/list":
        tools_list = [
            {
                "name": name,
                "description": meta["description"],
                "parallel": meta.get("parallel", True),
                "inputSchema": meta["inputSchema"],
            }
            for name, meta in TOOLS.items()
        ]
        return {"jsonrpc": "2.0", "id": id_, "result": {"tools": tools_list}}

    if method == "tools/call":
        return _handle_tools_call(id_, params)

    if method == "ping":
        return {"jsonrpc": "2.0", "id": id_, "result": {}}

    return make_error(id_, -32601, f"Method not found: {method}")


# ── stdio transport loop ──────────────────────────────────────────────────────

def main() -> None:
    """MCP stdio transport loop — read JSON-RPC requests from stdin, write responses to stdout."""
    _start_reminder_checker_once()
    for line in sys.stdin:
        line = line.strip()
        if not line:
            continue
        try:
            request = json.loads(line)
        except json.JSONDecodeError:
            try:
                sys.stdout.write(json.dumps(make_error(None, -32700, "Parse error")) + "\n")
                sys.stdout.flush()
            except BrokenPipeError:
                return
            continue

        response = handle_request(request)
        if response is not None:
            try:
                sys.stdout.write(json.dumps(response) + "\n")
                sys.stdout.flush()
            except BrokenPipeError:
                return


if __name__ == "__main__":
    main()
